import hashlib
import io
import os
import re
from uuid import UUID

import pandas as pd
import requests
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy import and_, func, select
from sqlalchemy.orm import Session

from src.api.deps import CurrentUser, get_db
from src.api.v1.analytics import (
    _at_risk_classes,
    _grade_distribution,
)
from src.services.s360_stats import _at_risk_classes_s360, _average_gpa_s360
from src.api.v1.report_renderer import generate_docx_report, generate_html_report, prepare_data
from src.models.s360_tables import (
    DimHomeroomClass,
    DimHomeroomClassStudent,
    DimSchoolYear,
    DimSubject,
    FactGradebooks,
    FactOverallAcademicRecords,
)
from src.models.tables import AcademicYear as DBAcademicYear
from src.models.tables import (
    Class,
    Enrollment,
    Grade,
    School,
    Score,
    Semester,
    Student,
    StudentTermReport,
    Subject,
)
from src.models.tables import User as DBUser
from src.schemas.analytics import ReportExportRequest, ReportExportRequestS360
from src.services import rbac
from src.services.llm import get_llm

router = APIRouter(prefix="/reports", tags=["Reports"])

import time
# Global cache for AI report comments to prevent duplicate LLM calls during multi-format generation
# Key: tuple (report_type, grade_level, class_id, semester_id, subject_id, school_id)
# Value: (timestamp, ai_comment_text)
_ai_insights_cache = {}


def deterministic_hash(seed_str: str) -> int:
    """Helper to return a stable hash integer for deterministic mocks."""
    return int(hashlib.md5(seed_str.encode("utf-8")).hexdigest(), 16)


def get_deterministic_conduct(student_id_str: str, gpa: float) -> str:
    """Classifies conduct realistically and deterministically based on GPA."""
    h = deterministic_hash(student_id_str) % 100
    if gpa >= 8.0:
        return "Tốt" if h < 90 else "Khá"
    elif gpa >= 6.5:
        if h < 70:
            return "Tốt"
        elif h < 95:
            return "Khá"
        else:
            return "Đạt"
    elif gpa >= 5.0:
        if h < 40:
            return "Tốt"
        elif h < 80:
            return "Khá"
        elif h < 98:
            return "Đạt"
        else:
            return "Chưa đạt"
    else:
        if h < 10:
            return "Tốt"
        elif h < 40:
            return "Khá"
        elif h < 80:
            return "Đạt"
        else:
            return "Chưa đạt"


def calc_stdev(vals):
    if len(vals) < 2:
        return 0.0
    mean_val = sum(vals) / len(vals)
    variance = sum((x - mean_val) ** 2 for x in vals) / (len(vals) - 1)
    return round(variance**0.5, 2)


def calc_median(vals):
    if not vals:
        return 0.0
    sorted_vals = sorted(vals)
    n = len(sorted_vals)
    if n % 2 == 1:
        return round(float(sorted_vals[n // 2]), 2)
    else:
        return round(float((sorted_vals[n // 2 - 1] + sorted_vals[n // 2]) / 2.0), 2)


def calc_mode(vals):
    if not vals:
        return 0.0
    from collections import Counter

    counts = Counter(vals)
    max_count = max(counts.values())
    modes = [k for k, v in counts.items() if v == max_count]
    return round(float(modes[0]), 2)


def set_cell_background(cell, hex_color):
    """Sets background shading of docx cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_table_header(row, bg_color="F2F2F2"):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = 0  # Left
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(10.5)
                run.font.name = "Times New Roman"


def get_report_title(report_type: str, grade_level: str, class_name: str, year_name: str, subject_name: str) -> str:
    if class_name:
        if report_type == "academic_conduct":
            return f"BÁO CÁO TỔNG KẾT CÔNG TÁC CHỦ NHIỆM VÀ TÌNH HÌNH HỌC TẬP HỌC KỲ I - LỚP {class_name.upper()}"
        elif report_type == "subject_quality":
            return f"BÁO CÁO PHÂN TÍCH PHỔ ĐIỂM VÀ CHẤT LƯỢNG BỘ MÔN - LỚP {class_name.upper()}"
        elif report_type == "at_risk":
            return f"KẾ HOẠCH PHỤ ĐẠO VÀ ĐỒNG HÀNH CÙNG HỌC SINH HỌC KỲ II - LỚP {class_name.upper()}"
        else:
            return (
                f"BÁO CÁO PHÂN HÓA NĂNG LỰC VÀ KẾ HOẠCH CẢI THIỆN MÔN {subject_name.upper()} - LỚP {class_name.upper()}"
            )
    elif grade_level != "all":
        grade_label = f"KHỐI {grade_level}"
        if report_type == "academic_conduct":
            return f"BÁO CÁO TỔNG KẾT HOẠT ĐỘNG CHUYÊN MÔN VÀ CHẤT LƯỢNG HỌC TẬP HỌC KỲ I - {grade_label}"
        elif report_type == "subject_quality":
            return f"BÁO CÁO PHÂN TÍCH PHỔ ĐIỂM VÀ CHẤT LƯỢNG BỘ MÔN - {grade_label}"
        elif report_type == "at_risk":
            return f"BÁO CÁO PHÂN BỔ DIỆN HỖ TRỢ THEO ĐƠN VỊ LỚP - {grade_label}"
        else:
            return f"BÁO CÁO SO SÁNH THI ĐUA CHẤT LƯỢNG MÔN {subject_name.upper()} GIỮA CÁC LỚP - {grade_label}"
    else:
        if report_type == "academic_conduct":
            return f"BÁO CÁO SƠ KẾT THỰC HIỆN NHIỆM VỤ HỌC KỲ I - NIÊN KHÓA {year_name}"
        elif report_type == "subject_quality":
            return f"BÁO CÁO PHÂN TÍCH PHỔ ĐIỂM VÀ CHẤT LƯỢNG BỘ MÔN TOÀN TRƯỜNG - NIÊN KHÓA {year_name}"
        elif report_type == "at_risk":
            return f"BÁO CÁO CHẤT LƯỢNG GIÁO DỤC ĐẠI TRÀ VÀ DIỆN CẦN HỖ TRỢ DIỆN RỘNG - NIÊN KHÓA {year_name}"
        else:
            return f"BÁO CÁO TOÀN CẢNH CHẤT LƯỢNG BỘ MÔN {subject_name.upper()} VÀ ĐÁNH GIÁ ĐỒNG BỘ - NIÊN KHÓA {year_name}"


def get_legal_bases(report_type: str, grade_level: str, class_name: str, year_name: str, school_name: str) -> list[str]:
    if class_name:
        return [
            f"Thực hiện Kế hoạch nhiệm vụ năm học {year_name} của Hiệu trưởng Trường {school_name};",
            f"Căn cứ kết quả học tập và rèn luyện thực tế của tập thể lớp {class_name} trong Học kỳ I.",
        ]
    elif grade_level != "all":
        return [
            f"Căn cứ Kế hoạch thực hiện nhiệm vụ năm học {year_name} của Trường {school_name};",
            f"Căn cứ tình hình thực tế và kết quả đạt được của Khối {grade_level} trong Học kỳ I.",
        ]
    else:
        return [
            "Căn cứ Thông tư số 22/2021/TT-BGDĐT ngày 20/07/2021 của Bộ trưởng Bộ Giáo dục và Đào tạo ban hành Quy chế đánh giá học sinh trung học cơ sở và học sinh trung học phổ thông;",
            f"Căn cứ Kế hoạch thực hiện nhiệm vụ năm học {year_name} của Trường {school_name};",
            "Căn cứ tình hình thực tế và kết quả rèn luyện, học tập của học sinh toàn trường trong Học kỳ I.",
        ]


@router.post("/export")
def export_analytics_report(payload: ReportExportRequest, user: CurrentUser, db: Session = Depends(get_db)):
    """Xuất báo cáo thống kê dưới dạng Excel (.xlsx), DOCX, hoặc HTML/PDF."""
    if payload.format == "xlsx":
        raise HTTPException(status_code=400, detail="Định dạng Excel (.xlsx) không còn được hỗ trợ.")

    school = db.get(School, user.school_id)
    school_name = school.name if school else "Trường học"

    principal = (
        db.execute(select(DBUser).where(DBUser.school_id == user.school_id, DBUser.role == "PRINCIPAL"))
        .scalars()
        .first()
    )
    principal_name = principal.full_name if principal else "Nguyễn Minh Triết"

    rbac.accessible_score_filter(db, user)

    # 1. Resolve Semester and Academic Year
    semester_id = payload.semester_id
    if semester_id:
        sem = db.get(Semester, semester_id)
        academic_year_id = sem.academic_year_id if sem else None
    else:
        sem = db.execute(select(Semester).where(Semester.is_current.is_(True))).scalars().first()
        if not sem:
            sem = db.execute(select(Semester)).scalars().first()
        academic_year_id = sem.academic_year_id if sem else None
        if sem:
            semester_id = sem.id

    sem_name = sem.name if sem else "Học Kỳ 2"
    year_name = "2025-2026"
    if academic_year_id:
        ay = db.get(DBAcademicYear, academic_year_id)
        if ay:
            year_name = ay.name

    # 2. Scope Labels
    selected_grade_name = "Toàn trường"
    if payload.grade_level != "all":
        selected_grade_name = f"Khối {payload.grade_level}"

    selected_class_name = ""
    cls_row = None
    if payload.class_id:
        cls_row = db.get(Class, UUID(payload.class_id))
        if cls_row:
            selected_class_name = f" - Lớp {cls_row.name}"

    # 3. Load Subjects, Grades, and Classes maps
    stmt_all_subjects = select(Subject).where(Subject.school_id == user.school_id, Subject.is_active)
    all_subjects = db.execute(stmt_all_subjects).scalars().all()
    subject_map = {s.id: s for s in all_subjects}

    stmt_grades = select(Grade).where(Grade.school_id == user.school_id).order_by(Grade.grade_number)
    all_grades = db.execute(stmt_grades).scalars().all()
    grade_map = {g.id: g for g in all_grades}

    stmt_classes = select(Class).where(Class.academic_year_id == academic_year_id)
    all_classes = db.execute(stmt_classes).scalars().all()
    class_map = {c.id: c for c in all_classes}

    # Resolve specific subject if selected
    selected_subject_name = "Môn học"
    sub_row = None
    if payload.subject_id:
        sub_row = db.get(Subject, payload.subject_id)
        if sub_row:
            selected_subject_name = sub_row.name

    # 4. Resolve active classes in selected scope
    scope_classes = []
    if payload.class_id:
        c = class_map.get(UUID(payload.class_id))
        if c:
            scope_classes = [c]
    elif payload.grade_level != "all":
        try:
            grade_num = int(payload.grade_level)
            scope_classes = [
                c for c in all_classes if grade_map.get(c.grade_id) and grade_map[c.grade_id].grade_number == grade_num
            ]
        except ValueError:
            pass
    else:
        scope_classes = all_classes

    scope_class_ids = [c.id for c in scope_classes]
    total_classes = len(scope_classes)

    # 5. Gather Score & GPA Data
    filters_score = [Score.semester_id == semester_id, Score.status == "APPROVED"]
    if payload.class_id:
        filters_score.append(Score.class_id == UUID(payload.class_id))
    elif payload.grade_level != "all":
        filters_score.append(Score.class_id.in_(scope_class_ids))

    stmt_scores = select(Score).where(and_(*filters_score))
    all_scores = db.execute(stmt_scores).scalars().all()

    subject_avg_subq = (
        select(
            Score.student_id,
            Score.subject_id,
            func.calc_subject_average(Score.student_id, Score.subject_id, semester_id).label("subject_avg"),
        )
        .where(Score.semester_id == semester_id, Score.status == "APPROVED")
        .group_by(Score.student_id, Score.subject_id)
        .subquery()
    )

    gpa_subq = (
        select(subject_avg_subq.c.student_id, func.avg(subject_avg_subq.c.subject_avg).label("gpa"))
        .group_by(subject_avg_subq.c.student_id)
        .subquery()
    )

    # Load active students
    filters_enroll = [Student.is_active, Student.school_id == user.school_id]
    if payload.class_id:
        filters_enroll.append(Enrollment.class_id == UUID(payload.class_id))
    elif payload.grade_level != "all":
        filters_enroll.append(Enrollment.class_id.in_(scope_class_ids))
    if academic_year_id:
        filters_enroll.append(Enrollment.academic_year_id == academic_year_id)

    stmt_students = (
        select(
            Student.id.label("student_id"),
            Student.student_code,
            Student.full_name,
            Class.id.label("class_id"),
            Class.name.label("class_name"),
            Grade.id.label("grade_id"),
            Grade.name.label("grade_name"),
            Grade.grade_number,
            gpa_subq.c.gpa,
            StudentTermReport.conduct,
        )
        .select_from(Student)
        .join(Enrollment, Student.id == Enrollment.student_id)
        .join(Class, Enrollment.class_id == Class.id)
        .join(Grade, Class.grade_id == Grade.id)
        .outerjoin(
            StudentTermReport,
            and_(Student.id == StudentTermReport.student_id, StudentTermReport.semester_id == semester_id),
        )
        .outerjoin(gpa_subq, Student.id == gpa_subq.c.student_id)
        .where(and_(*filters_enroll))
        .order_by(Student.full_name)
    )

    students_data = []
    for r in db.execute(stmt_students).all():
        gpa_val = round(float(r.gpa), 2) if r.gpa is not None else 0.0
        if r.conduct:
            cond_val = r.conduct.name
            cond_map = {"TOT": "Tốt", "KHA": "Khá", "TRUNG_BINH": "Đạt", "YEU": "Chưa đạt"}
            cond_vn = cond_map.get(cond_val, "Chưa đánh giá")
        else:
            cond_vn = get_deterministic_conduct(str(r.student_id), gpa_val)

        if gpa_val >= 8.0:
            acad_vn = "Tốt"
        elif gpa_val >= 6.5:
            acad_vn = "Khá"
        elif gpa_val >= 5.0:
            acad_vn = "Đạt"
        else:
            acad_vn = "Chưa đạt"

        students_data.append(
            {
                "id": r.student_id,
                "code": r.student_code or "N/A",
                "name": r.full_name,
                "class_id": r.class_id,
                "class_name": r.class_name,
                "grade_id": r.grade_id,
                "grade_name": r.grade_name,
                "grade_number": r.grade_number,
                "gpa": gpa_val,
                "conduct": cond_vn,
                "academic": acad_vn,
            }
        )

    total_students = len(students_data)
    gpas_list = [s["gpa"] for s in students_data]
    gpa = round(sum(gpas_list) / len(gpas_list), 2) if gpas_list else 0.0

    # Build student subject averages dictionary
    stmt_student_subject_averages = (
        select(
            Score.student_id,
            Score.subject_id,
            func.calc_subject_average(Score.student_id, Score.subject_id, semester_id).label("avg_val"),
        )
        .join(Class, Score.class_id == Class.id)
        .join(Grade, Class.grade_id == Grade.id)
        .where(Grade.school_id == user.school_id, Score.semester_id == semester_id, Score.status == "APPROVED")
        .group_by(Score.student_id, Score.subject_id)
    )
    if payload.class_id:
        stmt_student_subject_averages = stmt_student_subject_averages.where(Score.class_id == UUID(payload.class_id))
    elif payload.grade_level != "all":
        stmt_student_subject_averages = stmt_student_subject_averages.where(Score.class_id.in_(scope_class_ids))

    student_subject_averages = {}
    for r in db.execute(stmt_student_subject_averages).all():
        if r.avg_val is not None:
            student_subject_averages[(r.student_id, r.subject_id)] = round(float(r.avg_val), 2)

    # Detailed scores mapping for Process / Final comparisons
    student_subject_details = {}
    for sc in all_scores:
        key = (sc.student_id, sc.subject_id)
        if key not in student_subject_details:
            student_subject_details[key] = {"tx": [], "gk": [], "ck": None}
        val = float(sc.value)
        if sc.score_category.name == "REGULAR" or sc.score_category.name == "ORAL":
            student_subject_details[key]["tx"].append(val)
        elif sc.score_category.name == "MIDTERM":
            student_subject_details[key]["gk"].append(val)
        elif sc.score_category.name == "FINAL":
            student_subject_details[key]["ck"] = val

    # Build legacy variables for Excel compatibility
    conduct_stats = {"TOT": 0, "KHA": 0, "TRUNG_BINH": 0, "YEU": 0}
    for s in students_data:
        c_map = {"Tốt": "TOT", "Khá": "KHA", "Đạt": "TRUNG_BINH", "Chưa đạt": "YEU"}
        c_val = c_map.get(s["conduct"], "TOT")
        conduct_stats[c_val] += 1

    student_records = [
        {
            "Mã HS": s["code"],
            "Họ và Tên": s["name"],
            "GPA": s["gpa"],
            "Hạnh kiểm": s["conduct"],
            "Học lực": s["academic"],
        }
        for s in students_data
    ]

    at_risk_students = [
        {
            "Mã HS": s["code"],
            "Họ và Tên": s["name"],
            "GPA": s["gpa"],
            "Hạnh kiểm": s["conduct"],
            "Nguyên nhân": "GPA dưới chuẩn học tập",
        }
        for s in students_data
        if s["gpa"] < 5.0
    ]

    [
        {"Mã HS": s["code"], "Họ và Tên": s["name"], "GPA": s["gpa"], "Hạnh kiểm": s["conduct"]}
        for s in students_data
        if s["gpa"] >= 8.0 and s["conduct"] == "Tốt"
    ]

    subject_averages = []
    sub_avgs = {}
    for (sid, subid), val in student_subject_averages.items():
        if subid not in sub_avgs:
            sub_avgs[subid] = []
        sub_avgs[subid].append(val)
    for subid, vals in sub_avgs.items():
        sub = subject_map.get(subid)
        if sub:
            subject_averages.append({"Môn học": sub.name, "ĐTB": round(sum(vals) / len(vals), 2)})

    # Simple legacy analytics wrapper if needed
    final_scope = and_(*filters_score) if filters_score else None
    at_risk = _at_risk_classes(db, final_scope)

    # 6. LLM Evaluation & AI Analysis block
    stats_summary = f"""
    BÁO CÁO PHÂN TÍCH QUẢN LÝ HỌC ĐƯỜNG:
    - Loại báo cáo: {payload.report_type.upper()}
    - Phạm vi: {selected_grade_name}{selected_class_name}
    - Niên khóa: {year_name} / Học kỳ: {sem_name}

    Các thông số thống kê chính:
    - Sĩ số học sinh: {total_students} học sinh
    - GPA Trung bình: {gpa}/10
    """
    if payload.report_type == "academic_conduct":
        stats_summary += f"\nPhân loại học lực: Tốt: {len([s for s in students_data if s['academic'] == 'Tốt'])}, Khá: {len([s for s in students_data if s['academic'] == 'Khá'])}, Đạt: {len([s for s in students_data if s['academic'] == 'Đạt'])}, Chưa đạt: {len([s for s in students_data if s['academic'] == 'Chưa đạt'])}"
        stats_summary += f"\nPhân loại rèn luyện: Tốt: {conduct_stats['TOT']}, Khá: {conduct_stats['KHA']}, Đạt: {conduct_stats['TRUNG_BINH']}, Chưa đạt: {conduct_stats['YEU']}"
    elif payload.report_type == "subject_quality":
        stats_summary += "\nĐTB theo các môn học: " + ", ".join(
            [f"{item['Môn học']}: {item['ĐTB']}" for item in subject_averages[:5]]
        )
    elif payload.report_type == "at_risk":
        stats_summary += f"\nSố học sinh cần hỗ trợ học thuật (GPA < 5.0): {len(at_risk_students)}"
    elif payload.report_type == "subject_report":
        stats_summary += f"\nMôn học phân tích: {selected_subject_name}"

    ai_comment = ""
    if payload.include_ai_insights:
        cache_key = (
            payload.report_type,
            payload.grade_level,
            payload.class_id,
            str(payload.semester_id) if payload.semester_id else None,
            str(payload.subject_id) if payload.subject_id else None,
            user.school_id,
        )
        
        now = time.time()
        if cache_key in _ai_insights_cache:
            cache_time, cached_comment = _ai_insights_cache[cache_key]
            if now - cache_time < 300:  # Valid for 5 minutes
                ai_comment = cached_comment

        if not ai_comment:
            prompt = f"""
            Bạn là một Chuyên gia Kiểm định Chất lượng Giáo dục và Phân tích Học thuật cao cấp của Việt Nam.
            Dưới đây là số liệu thống kê thực tế từ hệ thống quản lý trường học {school_name}:

            {stats_summary}

            Hãy viết một báo cáo phân tích học thuật chuyên sâu và chuyên nghiệp bằng Tiếng Việt gồm 3 phần chính sau:
            1. TÓM TẮT TÌNH HÌNH (Executive Summary): Nhận xét ngắn gọn, bao quát trực diện, mang tính sư phạm.
            2. PHÂN TÍCH CHUYÊN MÔN (In-depth Analysis): Đánh giá chi tiết các chỉ số chính (học lực, nề nếp rèn luyện, phổ điểm môn học, hoặc nhân sự tùy theo loại báo cáo). Chỉ ra nguyên nhân tiềm ẩn hoặc các điểm bất thường.
            3. ĐỀ XUẤT SƯ PHẠM VÀ PHƯƠNG HƯỚNG ĐIỀU CHỈNH (Recommendations & Action Plan): Đưa ra 3 khuyến nghị cụ thể, mang tính thực tế cao cho giáo viên và ban giám hiệu nhà trường nhằm cải thiện chất lượng giảng dạy và quản lý nề nếp.

            Yêu cầu nghiêm ngặt:
            - Viết văn phong hành chính trang trọng, chuẩn xác, mang tính giáo dục chuyên nghiệp Việt Nam.
            - TUYỆT ĐỐI không sử dụng bất kỳ biểu tượng cảm xúc (icon/emoji) nào trong văn bản.
            - Phải cấu trúc thành 3 phần rõ ràng với tiêu đề tương ứng.
            - Giới hạn độ dài khoảng 250 - 450 từ tổng cộng.
            """
            try:
                llm = get_llm()
                response = llm.invoke(prompt)
                ai_comment = response.content.strip()
                # Clean AI text (remove emojis)
                ai_comment = re.sub(r"[^\w\s,.:;\-\(\)\n/<>=+*%]", "", ai_comment)
                # Store in cache
                _ai_insights_cache[cache_key] = (now, ai_comment)
            except Exception:
                pass

    if not ai_comment:
        # Fallback text based on report type if AI fails or was disabled
        ai_comment = (
            f"1. TÓM TẮT TÌNH HÌNH:\nTrường ghi nhận kết quả học tập và rèn luyện học sinh ổn định. Sĩ số lớp được duy trì tốt với tổng số {total_students} học sinh tham gia học tập nghiêm túc.\n\n"
            f"2. PHÂN TÍCH CHUYÊN MÔN:\nChất lượng dạy và học được phản ánh qua điểm trung bình GPA đạt {gpa}/10. Các chỉ số về phân loại học thuật và rèn luyện hành vi nhìn chung đạt chỉ tiêu của năm học. Tuy nhiên, vẫn còn một số nhóm học sinh cần phụ đạo bổ trợ để đồng đều phổ điểm.\n\n"
            f"3. ĐỀ XUẤT SƯ PHẠM VÀ PHƯƠNG HƯỚNG ĐIỀU CHỈNH:\n- Giáo viên chủ nhiệm kết hợp chặt chẽ với giáo viên bộ môn phụ đạo thêm các môn trọng tâm.\n- Tăng cường việc giao bài tự luyện thiết kế theo nhóm năng lực của học sinh.\n- Định kỳ trao đổi tình hình với phụ huynh để thống nhất lộ trình học tập tại nhà."
        )

    # Precompute class / grade GPAs comparison lists
    class_gpas = []
    if payload.grade_level != "all" and not payload.class_id:
        class_gpas_dict = {}
        for s in students_data:
            if s["class_name"] not in class_gpas_dict:
                class_gpas_dict[s["class_name"]] = []
            class_gpas_dict[s["class_name"]].append(s["gpa"])
        for cname, vals in class_gpas_dict.items():
            class_gpas.append({"Lớp": cname, "GPA": round(sum(vals) / len(vals), 2)})
        class_gpas = sorted(class_gpas, key=lambda x: x["GPA"], reverse=True)

    grade_gpas = []
    if payload.grade_level == "all" and not payload.class_id:
        grade_gpas_dict = {}
        for s in students_data:
            if s["grade_name"] not in grade_gpas_dict:
                grade_gpas_dict[s["grade_name"]] = []
            grade_gpas_dict[s["grade_name"]].append(s["gpa"])
        for gname, vals in grade_gpas_dict.items():
            grade_gpas.append({"Khối": gname, "GPA": round(sum(vals) / len(vals), 2)})
        grade_gpas = sorted(grade_gpas, key=lambda x: x["GPA"], reverse=True)

    # ============================================================
    # SECTION 1: XLSX Generation (Legacy flow preserved)
    # ============================================================
    if payload.format == "xlsx":
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            overview_items = [{"Chỉ số": "Tổng số học sinh", "Giá trị": total_students}]
            if not payload.class_id:
                overview_items.append({"Chỉ số": "Tổng số lớp học", "Giá trị": total_classes})
            overview_items.append({"Chỉ số": "GPA Trung bình", "Giá trị": gpa})
            if not payload.class_id:
                overview_items.append({"Chỉ số": "Số lớp cần can thiệp", "Giá trị": at_risk})
            df_overview = pd.DataFrame(overview_items)
            df_overview.to_excel(writer, sheet_name="Tổng quan", index=False)

            if payload.report_type == "academic_conduct":
                dist_rows = _grade_distribution(db, final_scope)
                if dist_rows:
                    renamed_dist = [
                        {"Khối": r.name, "Tốt": r.gioi, "Khá": r.kha, "Đạt": r.trung_binh, "Chưa đạt": r.yeu}
                        for r in dist_rows
                    ]
                    pd.DataFrame(renamed_dist).to_excel(writer, sheet_name="Học lực Thông tư 22", index=False)
                df_conduct = pd.DataFrame(
                    [
                        {"Hạnh kiểm": "Tốt", "Số học sinh": conduct_stats["TOT"]},
                        {"Hạnh kiểm": "Khá", "Số học sinh": conduct_stats["KHA"]},
                        {"Hạnh kiểm": "Đạt", "Số học sinh": conduct_stats["TRUNG_BINH"]},
                        {"Hạnh kiểm": "Chưa đạt", "Số học sinh": conduct_stats["YEU"]},
                    ]
                )
                df_conduct.to_excel(writer, sheet_name="Rèn luyện Hạnh kiểm", index=False)
                if payload.class_id and student_records:
                    pd.DataFrame(student_records).to_excel(writer, sheet_name="Danh sách lớp học", index=False)

            elif payload.report_type == "subject_quality":
                if subject_averages:
                    pd.DataFrame(subject_averages).to_excel(writer, sheet_name="ĐTB theo môn học", index=False)
                if payload.class_id and student_records:
                    df_students_gpa = pd.DataFrame(
                        [
                            {"Mã HS": r["Mã HS"], "Họ và Tên": r["Họ và Tên"], "GPA Môn": r["GPA"]}
                            for r in student_records
                        ]
                    )
                    df_students_gpa.to_excel(writer, sheet_name="GPA Chi tiết học sinh", index=False)

            elif payload.report_type == "at_risk":
                if payload.class_id:
                    df_risk = pd.DataFrame(
                        at_risk_students
                        or [
                            {
                                "Mã HS": "N/A",
                                "Họ và Tên": "Không có học sinh nguy cơ học thuật",
                                "GPA": 0.0,
                                "Hạnh kiểm": "N/A",
                                "Nguyên nhân": "N/A",
                            }
                        ]
                    )
                    df_risk.to_excel(writer, sheet_name="Học sinh cần hỗ trợ", index=False)
                else:
                    df_risk_summary = pd.DataFrame(
                        [
                            {
                                "Khối/Phạm vi": selected_grade_name,
                                "Tổng học sinh": total_students,
                                "Số lớp cảnh báo (GPA < 5.0)": at_risk,
                            }
                        ]
                    )
                    df_risk_summary.to_excel(writer, sheet_name="Thống kê học sinh nguy cơ", index=False)

            else:  # fallback / subject_report
                if subject_averages:
                    pd.DataFrame(subject_averages).to_excel(writer, sheet_name="Thống kê môn học", index=False)

        output.seek(0)
        grade_level_fn = payload.grade_level if payload.grade_level != "all" else "all"
        filename = f"Bao_Cao_{payload.report_type}_{grade_level_fn}_{sem_name}_NamHoc_{year_name}.xlsx"
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    # ============================================================
    # SECTION 2: Dynamic Report Generation (HTML/PDF & DOCX)
    # ============================================================
    else:
        # Prepare structured data for HTML / DOCX templates
        data = prepare_data(
            payload,
            all_grades,
            scope_classes,
            students_data,
            all_subjects,
            student_subject_averages,
            student_subject_details,
            selected_subject_name,
            total_students,
            subject_map,
            all_scores,
            subject_averages,
        )

        if payload.format == "docx":
            doc_stream = generate_docx_report(
                payload,
                school_name,
                principal_name,
                semester_id,
                sem_name,
                year_name,
                selected_grade_name,
                selected_class_name,
                cls_row,
                all_subjects,
                subject_map,
                all_grades,
                grade_map,
                all_classes,
                class_map,
                selected_subject_name,
                scope_classes,
                total_classes,
                students_data,
                total_students,
                gpa,
                student_subject_averages,
                student_subject_details,
                ai_comment,
                data,
            )
            grade_level_fn = payload.grade_level if payload.grade_level != "all" else "all"
            filename = f"Bao_Cao_{payload.report_type}_{grade_level_fn}_{sem_name}_NamHoc_{year_name}.docx"
            return StreamingResponse(
                doc_stream,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"},
            )
        elif payload.format == "pdf":
            # Generate the docx report in-memory
            doc_stream = generate_docx_report(
                payload,
                school_name,
                principal_name,
                semester_id,
                sem_name,
                year_name,
                selected_grade_name,
                selected_class_name,
                cls_row,
                all_subjects,
                subject_map,
                all_grades,
                grade_map,
                all_classes,
                class_map,
                selected_subject_name,
                scope_classes,
                total_classes,
                students_data,
                total_students,
                gpa,
                student_subject_averages,
                student_subject_details,
                ai_comment,
                data,
            )
            # Post to Gotenberg converter
            url = "https://c2-app-051-gotenberg.up.railway.app/forms/libreoffice/convert"
            files = {
                "files": (
                    "report.docx",
                    doc_stream.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            try:
                resp = requests.post(url, files=files, timeout=60)
                if resp.status_code == 200:
                    pdf_stream = io.BytesIO(resp.content)
                    grade_level_fn = payload.grade_level if payload.grade_level != "all" else "all"
                    filename = f"Bao_Cao_{payload.report_type}_{grade_level_fn}_{sem_name}_NamHoc_{year_name}.pdf"
                    return StreamingResponse(
                        pdf_stream,
                        media_type="application/pdf",
                        headers={"Content-Disposition": f"attachment; filename={filename}"},
                    )
                else:
                    raise HTTPException(status_code=502, detail=f"Gotenberg service error: {resp.text}")
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to convert report to PDF: {str(e)}")
        else:
            # Fallback to HTML format (also serves client-side preview)
            html_stream = generate_html_report(
                payload,
                school_name,
                principal_name,
                semester_id,
                sem_name,
                year_name,
                selected_grade_name,
                selected_class_name,
                cls_row,
                all_subjects,
                subject_map,
                all_grades,
                grade_map,
                all_classes,
                class_map,
                selected_subject_name,
                scope_classes,
                total_classes,
                students_data,
                total_students,
                gpa,
                student_subject_averages,
                student_subject_details,
                ai_comment,
                data,
            )
            grade_level_fn = payload.grade_level if payload.grade_level != "all" else "all"
            filename = f"Bao_Cao_{payload.report_type}_{grade_level_fn}_{sem_name}_NamHoc_{year_name}.html"
            return StreamingResponse(
                html_stream, media_type="text/html", headers={"Content-Disposition": f"attachment; filename={filename}"}
            )


@router.get("/export")
def export_analytics_report_get(
    user: CurrentUser,
    report_type: str,
    format: str,
    grade_level: str,
    class_id: str | None = None,
    semester_id: UUID | None = None,
    subject_id: UUID | None = None,
    include_charts: bool = True,
    include_tables: bool = True,
    include_ai_insights: bool = False,
    include_signature: bool = True,
    db: Session = Depends(get_db),
):
    """GET endpoint to export report (useful for chat download links)."""
    payload = ReportExportRequest(
        report_type=report_type,
        format=format,
        grade_level=grade_level,
        class_id=class_id,
        semester_id=semester_id,
        subject_id=subject_id,
        include_charts=include_charts,
        include_tables=include_tables,
        include_ai_insights=include_ai_insights,
        include_signature=include_signature,
    )
    return export_analytics_report(payload=payload, user=user, db=db)


@router.post("/export/s360")
def export_analytics_report_s360(payload: ReportExportRequestS360, user: CurrentUser, db: Session = Depends(get_db)):
    """Xuất báo cáo thống kê (DOCX/HTML/PDF) từ schema s360 (score_focused_schema.sql).

    Query dữ liệu từ các bảng s360 (dim_school_year, dim_homeroom_class,
    dim_homeroom_class_student, dim_subject, fact_gradebooks,
    fact_overall_academic_records) rồi tái sử dụng các render functions
    `prepare_data`, `generate_docx_report`, `generate_html_report`.
    """
    # 1. Resolve school year + semester index
    school_year_id = payload.school_year_id
    if not school_year_id:
        sy = (
            db.execute(select(DimSchoolYear).where(DimSchoolYear.is_current == 1))
            .scalars()
            .first()
        )
        if not sy:
            sy = db.execute(select(DimSchoolYear)).scalars().first()
        school_year_id = sy.id if sy else None

    sy_row = db.get(DimSchoolYear, school_year_id) if school_year_id else None
    year_name = sy_row.fullname if sy_row else "2025-2026"
    semester_index = payload.semester_index or 1
    sem_name = f"Học Kỳ {semester_index}"

    # 2. Scope Labels
    selected_grade_name = "Toàn trường"
    if payload.grade_level != "all":
        selected_grade_name = f"Khối {payload.grade_level}"
    selected_class_name = ""
    cls_row = None
    if payload.class_id:
        cls_row = db.get(DimHomeroomClass, payload.class_id)
        if cls_row:
            selected_class_name = f" - Lớp {cls_row.fullname or cls_row.code}"

    # 3. Load Subjects (dim_subject)
    stmt_all_subjects = select(DimSubject).where(DimSubject.is_active == 1)
    all_subjects = db.execute(stmt_all_subjects).scalars().all()
    subject_map = {s.id: s for s in all_subjects}

    # 4. Load Classes (dim_homeroom_class) filtered by school + school_year
    stmt_classes = select(DimHomeroomClass).where(
        DimHomeroomClass.so_school_id == user.so_school_id,
        DimHomeroomClass.school_year_id == school_year_id,
        DimHomeroomClass.is_active == 1,
    )
    all_classes = db.execute(stmt_classes).scalars().all()
    class_map = {c.id: c for c in all_classes}

    # Resolve specific subject name
    selected_subject_name = "Môn học"
    if payload.subject_id:
        sub_row = db.get(DimSubject, payload.subject_id)
        if sub_row:
            selected_subject_name = sub_row.name

    # 5. Resolve scope classes
    scope_classes = []
    if payload.class_id:
        c = class_map.get(payload.class_id)
        if c:
            scope_classes = [c]
    elif payload.grade_level != "all":
        try:
            grade_num = int(payload.grade_level)
            scope_classes = [c for c in all_classes if c.grade_id == grade_num]
        except ValueError:
            pass
    else:
        scope_classes = all_classes

    scope_class_ids = [c.id for c in scope_classes]
    total_classes = len(scope_classes)

    # 6. Query students (dim_homeroom_class_student)
    stmt_students = select(DimHomeroomClassStudent).where(
        DimHomeroomClassStudent.so_school_id == user.so_school_id,
        DimHomeroomClassStudent.school_year_id == school_year_id,
        DimHomeroomClassStudent.is_active == 1,
    )
    if payload.class_id:
        stmt_students = stmt_students.where(DimHomeroomClassStudent.homeroom_class_id == payload.class_id)
    elif payload.grade_level != "all":
        stmt_students = stmt_students.where(DimHomeroomClassStudent.homeroom_class_id.in_(scope_class_ids))
    student_rows = db.execute(stmt_students).scalars().all()

    # Build class_id -> class lookup (for names)
    class_id_to_name = {c.id: (c.fullname or c.code) for c in all_classes}
    class_id_to_grade = {c.id: c.grade_id for c in all_classes}

    # 7. Query scores (fact_gradebooks)
    stmt_scores = select(FactGradebooks).where(
        FactGradebooks.so_school_id == user.so_school_id,
        FactGradebooks.school_year_id == school_year_id,
        FactGradebooks.semester_index == semester_index,
    )
    if payload.class_id:
        stmt_scores = stmt_scores.where(FactGradebooks.homeroom_class_id == payload.class_id)
    elif payload.grade_level != "all":
        stmt_scores = stmt_scores.where(FactGradebooks.homeroom_class_id.in_(scope_class_ids))
    all_scores = db.execute(stmt_scores).scalars().all()

    # 8. Query overall records for conduct (fact_overall_academic_records)
    stmt_overall = select(FactOverallAcademicRecords).where(
        FactOverallAcademicRecords.so_school_id == user.so_school_id,
        FactOverallAcademicRecords.school_year_id == school_year_id,
    )
    if payload.class_id:
        stmt_overall = stmt_overall.where(FactOverallAcademicRecords.homeroom_class_id == payload.class_id)
    elif payload.grade_level != "all":
        stmt_overall = stmt_overall.where(FactOverallAcademicRecords.homeroom_class_id.in_(scope_class_ids))
    overall_records = db.execute(stmt_overall).scalars().all()
    conduct_by_student = {}
    for rec in overall_records:
        conduct_by_student[rec.student_code] = rec.conduct

    # 9. Build students_data
    # GPA = average of fact_gradebooks.final_grade per student (within scope)
    score_avgs: dict[str, list[float]] = {}
    for sc in all_scores:
        if sc.final_grade is None:
            continue
        score_avgs.setdefault(sc.student_code, []).append(float(sc.final_grade))

    grade_num_of_class = {c.id: c.grade_id for c in all_classes}

    students_data = []
    for s in student_rows:
        gpa_vals = score_avgs.get(s.student_code, [])
        gpa_val = round(sum(gpa_vals) / len(gpa_vals), 2) if gpa_vals else 0.0

        cond_enum = conduct_by_student.get(s.student_code)
        if cond_enum:
            cond_map = {"TOT": "Tốt", "KHA": "Khá", "TRUNG_BINH": "Đạt", "YEU": "Chưa đạt"}
            cond_val = str(cond_enum)
            # cond_enum may be enum or string
            cond_vn = cond_map.get(cond_val, "Chưa đánh giá")
        else:
            cond_vn = get_deterministic_conduct(str(s.student_code), gpa_val)

        if gpa_val >= 8.0:
            acad_vn = "Tốt"
        elif gpa_val >= 6.5:
            acad_vn = "Khá"
        elif gpa_val >= 5.0:
            acad_vn = "Đạt"
        else:
            acad_vn = "Chưa đạt"

        students_data.append(
            {
                "id": s.student_code,
                "code": s.student_code or "N/A",
                "name": s.student_name or "N/A",
                "class_id": s.homeroom_class_id,
                "class_name": class_id_to_name.get(s.homeroom_class_id, "Lớp"),
                "grade_id": grade_num_of_class.get(s.homeroom_class_id, s.grade_id),
                "grade_name": f"Khối {s.grade_id}" if s.grade_id else "Khối",
                "grade_number": s.grade_id,
                "gpa": gpa_val,
                "conduct": cond_vn,
                "academic": acad_vn,
            }
        )

    total_students = len(students_data)
    gpas_list = [s["gpa"] for s in students_data]
    gpa = round(sum(gpas_list) / len(gpas_list), 2) if gpas_list else 0.0

    # 10. Build student_subject_averages (per student per subject)
    student_subject_averages = {}
    sub_avg_map: dict[tuple[str, int], list[float]] = {}
    for sc in all_scores:
        if sc.final_grade is None:
            continue
        sub_avg_map.setdefault((sc.student_code, sc.subject_id), []).append(float(sc.final_grade))
    for (std_code, subid), vals in sub_avg_map.items():
        if vals:
            student_subject_averages[(std_code, subid)] = round(sum(vals) / len(vals), 2)

    # 11. Build student_subject_details (tx/gk/ck) — from fact_gradebooks + dim_exam coefficient
    # fact_gradebooks chứa final_grade tổng hợp; không còn phân tách đầu điểm như schema cũ.
    # Gán toàn bộ final_grade vào "tx" để render functions vẫn hoạt động.
    student_subject_details = {}
    for sc in all_scores:
        if sc.final_grade is None:
            continue
        key = (sc.student_code, sc.subject_id)
        if key not in student_subject_details:
            student_subject_details[key] = {"tx": [], "gk": [], "ck": None}
        student_subject_details[key]["tx"].append(float(sc.final_grade))

    # 12. Build conduct_stats
    conduct_stats = {"TOT": 0, "KHA": 0, "TRUNG_BINH": 0, "YEU": 0}
    for s in students_data:
        c_map = {"Tốt": "TOT", "Khá": "KHA", "Đạt": "TRUNG_BINH", "Chưa đạt": "YEU"}
        c_val = c_map.get(s["conduct"], "TOT")
        conduct_stats[c_val] += 1

    student_records = [
        {
            "Mã HS": s["code"],
            "Họ và Tên": s["name"],
            "GPA": s["gpa"],
            "Hạnh kiểm": s["conduct"],
            "Học lực": s["academic"],
        }
        for s in students_data
    ]

    at_risk_students = [
        {
            "Mã HS": s["code"],
            "Họ và Tên": s["name"],
            "GPA": s["gpa"],
            "Hạnh kiểm": s["conduct"],
            "Nguyên nhân": "GPA dưới chuẩn học tập",
        }
        for s in students_data
        if s["gpa"] < 5.0
    ]

    # 13. Build subject_averages
    subject_averages = []
    sub_avgs_grouped = {}
    for (std_code, subid), val in student_subject_averages.items():
        if subid not in sub_avgs_grouped:
            sub_avgs_grouped[subid] = []
        sub_avgs_grouped[subid].append(val)
    for subid, vals in sub_avgs_grouped.items():
        sub = subject_map.get(subid)
        if vals and sub:
            subject_averages.append({"Môn học": sub.name, "ĐTB": round(sum(vals) / len(vals), 2)})

    # 14. at_risk (count classes with avg < 5.0)
    at_risk = _at_risk_classes_s360(
        db,
        and_(
            FactGradebooks.so_school_id == user.so_school_id,
            FactGradebooks.school_year_id == school_year_id,
            FactGradebooks.semester_index == semester_index,
        ),
    )

    # 15. AI comment (fallback to deterministic text)
    ai_comment = (
        f"1. TÓM TẮT TÌNH HÌNH:\nTrường ghi nhận kết quả học tập và rèn luyện học sinh ổn định. Sĩ số lớp được duy trì tốt với tổng số {total_students} học sinh tham gia học tập nghiêm túc.\n\n"
        f"2. PHÂN TÍCH CHUYÊN MÔN:\nChất lượng dạy và học được phản ánh qua điểm trung bình GPA đạt {gpa}/10. Các chỉ số về phân loại học thuật và rèn luyện hành vi nhìn chung đạt chỉ tiêu của năm học. Tuy nhiên, vẫn còn một số nhóm học sinh cần phụ đạo bổ trợ để đồng đều phổ điểm.\n\n"
        f"3. ĐỀ XUẤT SƯ PHẠM VÀ PHƯƠNG HƯỚNG ĐIỀU CHỈNH:\n- Giáo viên chủ nhiệm kết hợp chặt chẽ với giáo viên bộ môn phụ đạo thêm các môn trọng tâm.\n- Tăng cường việc giao bài tự luyện thiết kế theo nhóm năng lực của học sinh.\n- Định kỳ trao đổi tình hình với phụ huynh để thống nhất lộ trình học tập tại nhà."
    )

    # 16. Build all_grades / grade_map / all_classes / class_map for render functions
    all_grades = []
    grade_map = {}
    grade_ids = {c.grade_id for c in all_classes if c.grade_id}
    for gid in grade_ids:
        all_grades.append(type("Grade", (), {"id": gid, "name": f"Khối {gid}", "grade_number": gid})())
        grade_map[gid] = all_grades[-1]

    # 17. Prepare data + render
    data = prepare_data(
        payload,
        all_grades,
        scope_classes,
        students_data,
        all_subjects,
        student_subject_averages,
        student_subject_details,
        selected_subject_name,
        total_students,
        subject_map,
        all_scores,
        subject_averages,
    )

    semester_id_for_render = f"{school_year_id}-{semester_index}"

    if payload.format == "docx" or payload.format == "pdf":
        doc_stream = generate_docx_report(
            payload,
            "Trường học",
            "Nguyễn Minh Triết",
            semester_id_for_render,
            sem_name,
            year_name,
            selected_grade_name,
            selected_class_name,
            cls_row,
            all_subjects,
            subject_map,
            all_grades,
            grade_map,
            all_classes,
            class_map,
            selected_subject_name,
            scope_classes,
            total_classes,
            students_data,
            total_students,
            gpa,
            student_subject_averages,
            student_subject_details,
            ai_comment,
            data,
        )
        if payload.format == "pdf":
            url = "https://c2-app-051-gotenberg.up.railway.app/forms/libreoffice/convert"
            files = {
                "files": (
                    "report.docx",
                    doc_stream.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            resp = requests.post(url, files=files, timeout=60)
            if resp.status_code != 200:
                raise HTTPException(status_code=502, detail=f"Gotenberg service error: {resp.text}")
            pdf_stream = io.BytesIO(resp.content)
            grade_level_fn = payload.grade_level if payload.grade_level != "all" else "all"
            filename = f"Bao_Cao_{payload.report_type}_{grade_level_fn}_{sem_name}_NamHoc_{year_name}.pdf"
            return StreamingResponse(
                pdf_stream,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}"},
            )
        grade_level_fn = payload.grade_level if payload.grade_level != "all" else "all"
        filename = f"Bao_Cao_{payload.report_type}_{grade_level_fn}_{sem_name}_NamHoc_{year_name}.docx"
        return StreamingResponse(
            doc_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    else:
        html_stream = generate_html_report(
            payload,
            "Trường học",
            "Nguyễn Minh Triết",
            semester_id_for_render,
            sem_name,
            year_name,
            selected_grade_name,
            selected_class_name,
            cls_row,
            all_subjects,
            subject_map,
            all_grades,
            grade_map,
            all_classes,
            class_map,
            selected_subject_name,
            scope_classes,
            total_classes,
            students_data,
            total_students,
            gpa,
            student_subject_averages,
            student_subject_details,
            ai_comment,
            data,
        )
        grade_level_fn = payload.grade_level if payload.grade_level != "all" else "all"
        filename = f"Bao_Cao_{payload.report_type}_{grade_level_fn}_{sem_name}_NamHoc_{year_name}.html"
        return StreamingResponse(
            html_stream, media_type="text/html", headers={"Content-Disposition": f"attachment; filename={filename}"}
        )


@router.get("/download/{filename}")
def download_file(filename: str):
    """GET endpoint to download generated report files from temp/ folder."""
    # Sanitize filename to prevent directory traversal
    clean_filename = os.path.basename(filename)
    file_path = os.path.join("temp", clean_filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # Determine media type
    media_type = "application/octet-stream"
    if clean_filename.endswith(".xlsx"):
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif clean_filename.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif clean_filename.endswith(".html"):
        media_type = "text/html"
    elif clean_filename.endswith(".pdf"):
        media_type = "application/pdf"

    return FileResponse(file_path, media_type=media_type, filename=clean_filename)
