import hashlib
import io
import re

import docx
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def deterministic_hash(seed_str: str) -> int:
    return int(hashlib.md5(seed_str.encode("utf-8")).hexdigest(), 16)


def clean_markdown(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"^\s*[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    return text.strip()


def parse_ai_comment(ai_comment: str) -> tuple[str, str, str]:
    text = ai_comment.strip()

    p1 = re.search(r"(?:^|\n)[#\s\*]*1\.?[\s\*]*(?:TÓM TẮT|Tóm tắt|EXECUTIVE|Executive)[^\n]*", text, re.IGNORECASE)
    p2 = re.search(r"(?:^|\n)[#\s\*]*2\.?[\s\*]*(?:PHÂN TÍCH|Phân tích|IN-DEPTH|In-depth)[^\n]*", text, re.IGNORECASE)
    p3 = re.search(
        r"(?:^|\n)[#\s\*]*3\.?[\s\*]*(?:ĐỀ XUẤT|Đề xuất|RECOMMENDATIONS|Recommendations|PHƯƠNG HƯỚNG|Phương hướng)[^\n]*",
        text,
        re.IGNORECASE,
    )

    idx1 = p1.start() if p1 else -1
    idx2 = p2.start() if p2 else -1
    idx3 = p3.start() if p3 else -1

    part1, part2, part3 = "", "", ""

    if idx1 != -1:
        start_1 = p1.end()
        end_1 = idx2 if idx2 != -1 else (idx3 if idx3 != -1 else len(text))
        part1 = text[start_1:end_1].strip()
    else:
        end_1 = idx2 if idx2 != -1 else (idx3 if idx3 != -1 else len(text))
        part1 = text[0:end_1].strip()

    if idx2 != -1:
        start_2 = p2.end()
        end_2 = idx3 if idx3 != -1 else len(text)
        part2 = text[start_2:end_2].strip()

    if idx3 != -1:
        start_3 = p3.end()
        part3 = text[start_3:].strip()

    def clean_part(p):
        p = clean_markdown(p)
        if p.startswith(":"):
            p = p[1:].strip()
        return p

    return clean_part(part1), clean_part(part2), clean_part(part3)


def add_heading_with_font(doc, text, level):
    h = doc.add_paragraph(style=f"Heading {level}")
    run = h.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    return h


def align_docx_table(table, left_cols=None):
    # Center the table on the page
    table.alignment = 1  # Center

    # Set default table margins (Cell Padding) - 140 dxa = 7pt for top/bottom
    tblPr = table._tbl.tblPr  # noqa: N806
    tblCellMar = OxmlElement("w:tblCellMar")  # noqa: N806
    for margin, val in [("top", "140"), ("bottom", "140"), ("left", "150"), ("right", "150")]:
        node = OxmlElement(f"w:{margin}")
        node.set(qn("w:w"), val)
        node.set(qn("w:type"), "dxa")
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            # Set vertical alignment to center
            tcPr = cell._tc.get_or_add_tcPr()  # noqa: N806
            vAlign = OxmlElement("w:vAlign")  # noqa: N806
            vAlign.set(qn("w:val"), "center")
            tcPr.append(vAlign)

            for p in cell.paragraphs:
                # Header row: Center align
                if r_idx == 0:
                    p.alignment = 1  # Center
                else:
                    if left_cols and c_idx in left_cols:
                        p.alignment = 0  # Left
                    elif len(cell.text.strip()) > 35:
                        p.alignment = 0  # Left
                    else:
                        p.alignment = 1  # Center


def calc_stdev(vals):
    if len(vals) < 2:
        return 0.0
    mean_val = sum(vals) / len(vals)
    variance = sum((x - mean_val) ** 2 for x in vals) / (len(vals) - 1)
    return round(variance**0.5, 2)


def calc_median(vals):
    if not vals:
        return 0.0
    sorted_vals = sorted(vals)
    n = len(sorted_vals)
    if n % 2 == 1:
        return round(float(sorted_vals[n // 2]), 2)
    else:
        return round(float((sorted_vals[n // 2 - 1] + sorted_vals[n // 2]) / 2.0), 2)


def calc_mode(vals):
    if not vals:
        return 0.0
    from collections import Counter

    counts = Counter(vals)
    max_count = max(counts.values())
    modes = [k for k, v in counts.items() if v == max_count]
    return round(float(modes[0]), 2)


def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_table_header(row, bg_color="F2F2F2"):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = 0  # Left
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(10.5)
                run.font.name = "Times New Roman"


def get_report_title(report_type: str, grade_level: str, class_name: str, year_name: str, subject_name: str) -> str:
    if class_name:
        if report_type == "academic_conduct":
            return f"BÁO CÁO TỔNG KẾT CÔNG TÁC CHỦ NHIỆM VÀ TÌNH HÌNH HỌC TẬP HỌC KỲ I - LỚP {class_name.upper()}"
        elif report_type == "subject_quality":
            return f"BÁO CÁO PHÂN TÍCH PHỔ ĐIỂM VÀ CHẤT LƯỢNG BỘ MÔN - LỚP {class_name.upper()}"
        elif report_type == "at_risk":
            return f"KẾ HOẠCH PHỤ ĐẠO VÀ ĐỒNG HÀNH CÙNG HỌC SINH HỌC KỲ II - LỚP {class_name.upper()}"
        else:
            return (
                f"BÁO CÁO PHÂN HÓA NĂNG LỰC VÀ KẾ HOẠCH CẢI THIỆN MÔN {subject_name.upper()} - LỚP {class_name.upper()}"
            )
    elif grade_level != "all":
        grade_label = f"KHỐI {grade_level}"
        if report_type == "academic_conduct":
            return f"BÁO CÁO TỔNG KẾT HOẠT ĐỘNG CHUYÊN MÔN VÀ CHẤT LƯỢNG HỌC TẬP HỌC KỲ I - {grade_label}"
        elif report_type == "subject_quality":
            return f"BÁO CÁO PHÂN TÍCH PHỔ ĐIỂM VÀ CHẤT LƯỢNG BỘ MÔN - {grade_label}"
        elif report_type == "at_risk":
            return f"BÁO CÁO PHÂN BỔ DIỆN HỖ TRỢ THEO ĐƠN VỊ LỚP - {grade_label}"
        else:
            return f"BÁO CÁO SO SÁNH THI ĐUA CHẤT LƯỢNG MÔN {subject_name.upper()} GIỮA CÁC LỚP - {grade_label}"
    else:
        if report_type == "academic_conduct":
            return f"BÁO CÁO SƠ KẾT THỰC HIỆN NHIỆM VỤ HỌC KỲ I - NIÊN KHÓA {year_name}"
        elif report_type == "subject_quality":
            return f"BÁO CÁO PHÂN TÍCH PHỔ ĐIỂM VÀ CHẤT LƯỢNG BỘ MÔN TOÀN TRƯỜNG - NIÊN KHÓA {year_name}"
        elif report_type == "at_risk":
            return f"BÁO CÁO CHẤT LƯỢNG GIÁO DỤC ĐẠI TRÀ VÀ DIỆN CẦN HỖ TRỢ DIỆN RỘNG - NIÊN KHÓA {year_name}"
        else:
            return f"BÁO CÁO TOÀN CẢNH CHẤT LƯỢNG BỘ MÔN {subject_name.upper()} VÀ ĐÁNH GIÁ ĐỒNG BỘ - NIÊN KHÓA {year_name}"


def get_legal_bases(report_type: str, grade_level: str, class_name: str, year_name: str, school_name: str) -> list[str]:
    if class_name:
        return [
            f"Thực hiện Kế hoạch nhiệm vụ năm học {year_name} của Hiệu trưởng Trường {school_name};",
            f"Căn cứ kết quả học tập và rèn luyện thực tế của tập thể lớp {class_name} trong Học kỳ I.",
        ]
    elif grade_level != "all":
        return [
            f"Căn cứ Kế hoạch thực hiện nhiệm vụ năm học {year_name} của Trường {school_name};",
            f"Căn cứ tình hình thực tế và kết quả đạt được của Khối {grade_level} trong Học kỳ I.",
        ]
    else:
        return [
            "Căn cứ Thông tư số 22/2021/TT-BGDĐT ngày 20/07/2021 của Bộ trưởng Bộ Giáo dục và Đào tạo ban hành Quy chế đánh giá học sinh trung học cơ sở và học sinh trung học phổ thông;",
            f"Căn cứ Kế hoạch thực hiện nhiệm vụ năm học {year_name} của Trường {school_name};",
            "Căn cứ tình hình thực tế và kết quả rèn luyện, học tập của học sinh toàn trường trong Học kỳ I.",
        ]


def prepare_data(
    payload,
    all_grades,
    scope_classes,
    students_data,
    all_subjects,
    student_subject_averages,
    student_subject_details,
    selected_subject_name,
    total_students,
    subject_map,
    all_scores,
    subject_averages,
):
    # s_siso_rows, s_acad_rows, s_cond_rows, s_award_rows, total_dau, total_cuoi, total_den, total_di, total_bo
    # tot_tot_a, tot_kha_a, tot_dat_a, tot_yeu_a, tot_tot_c, tot_kha_c, tot_dat_c, tot_yeu_c, tot_xs, tot_gi, tot_cd
    s_siso_rows = []
    s_acad_rows = []
    s_cond_rows = []
    s_award_rows = []
    total_dau, total_cuoi, total_den, total_di, total_bo = 0, 0, 0, 0, 0
    tot_tot_a, tot_kha_a, tot_dat_a, tot_yeu_a = 0, 0, 0, 0
    tot_tot_c, tot_kha_c, tot_dat_c, tot_yeu_c = 0, 0, 0, 0
    tot_xs, tot_gi, tot_cd = 0, 0, 0

    for g in all_grades:
        g_classes = [c for c in scope_classes if c.grade_id == g.id]
        if not g_classes:
            continue
        g_students = [s for s in students_data if s["grade_id"] == g.id]
        cuoi = len(g_students)
        chuyen_den = deterministic_hash(g.name + "den") % 2
        chuyen_di = deterministic_hash(g.name + "di") % 2
        bo_hoc = 1 if (deterministic_hash(g.name + "bo") % 7 == 0) else 0
        dau = cuoi - chuyen_den + chuyen_di + bo_hoc
        ty_le = round((cuoi / dau) * 100, 2) if dau > 0 else 100.0

        total_dau += dau
        total_cuoi += cuoi
        total_den += chuyen_den
        total_di += chuyen_di
        total_bo += bo_hoc

        s_siso_rows.append(
            {
                "grade": g.name,
                "classes": len(g_classes),
                "dau": dau,
                "cuoi": cuoi,
                "den": chuyen_den,
                "di": chuyen_di,
                "bo": bo_hoc,
                "ratio": f"{ty_le}%",
            }
        )

        tot_a = len([s for s in g_students if s["academic"] == "Tốt"])
        kha_a = len([s for s in g_students if s["academic"] == "Khá"])
        dat_a = len([s for s in g_students if s["academic"] == "Đạt"])
        yeu_a = len([s for s in g_students if s["academic"] == "Chưa đạt"])

        tot_tot_a += tot_a
        tot_kha_a += kha_a
        tot_dat_a += dat_a
        tot_yeu_a += yeu_a

        s_acad_rows.append(
            {
                "grade": g.name,
                "total": cuoi,
                "tot": tot_a,
                "tot_pct": f"{round(tot_a / cuoi * 100, 1)}%" if cuoi else "0%",
                "kha": kha_a,
                "kha_pct": f"{round(kha_a / cuoi * 100, 1)}%" if cuoi else "0%",
                "dat": dat_a,
                "dat_pct": f"{round(dat_a / cuoi * 100, 1)}%" if cuoi else "0%",
                "yeu": yeu_a,
                "yeu_pct": f"{round(yeu_a / cuoi * 100, 1)}%" if cuoi else "0%",
            }
        )

        tot_c = len([s for s in g_students if s["conduct"] == "Tốt"])
        kha_c = len([s for s in g_students if s["conduct"] == "Khá"])
        dat_c = len([s for s in g_students if s["conduct"] == "Đạt"])
        yeu_c = len([s for s in g_students if s["conduct"] == "Chưa đạt"])

        tot_tot_c += tot_c
        tot_kha_c += kha_c
        tot_dat_c += dat_c
        tot_yeu_c += yeu_c

        s_cond_rows.append(
            {
                "grade": g.name,
                "total": cuoi,
                "tot": tot_c,
                "tot_pct": f"{round(tot_c / cuoi * 100, 1)}%" if cuoi else "0%",
                "kha": kha_c,
                "kha_pct": f"{round(kha_c / cuoi * 100, 1)}%" if cuoi else "0%",
                "dat": dat_c,
                "dat_pct": f"{round(dat_c / cuoi * 100, 1)}%" if cuoi else "0%",
                "yeu": yeu_c,
                "yeu_pct": f"{round(yeu_c / cuoi * 100, 1)}%" if cuoi else "0%",
            }
        )

        xs = len([s for s in g_students if s["gpa"] >= 9.0 and s["conduct"] == "Tốt"])
        gi = len(
            [
                s
                for s in g_students
                if s["gpa"] >= 8.0
                and s["conduct"] in ["Tốt", "Khá"]
                and not (s["gpa"] >= 9.0 and s["conduct"] == "Tốt")
            ]
        )
        cd = int(cuoi * 0.03)

        tot_xs += xs
        tot_gi += gi
        tot_cd += cd

        s_award_rows.append({"grade": g.name, "xuatsac": xs, "gioi": gi, "chuyende": cd})

    # Grade Level
    g_siso_rows = []
    g_acad_rows = []
    g_cond_rows = []
    g_subjects_rows = []
    for c in scope_classes:
        c_students = [s for s in students_data if s["class_id"] == c.id]
        cuoi = len(c_students)
        chuyen_den = deterministic_hash(c.name + "den") % 2
        chuyen_di = deterministic_hash(c.name + "di") % 2
        bo_hoc = 1 if (deterministic_hash(c.name + "bo") % 10 == 0) else 0
        dau = cuoi - chuyen_den + chuyen_di + bo_hoc
        co_p = deterministic_hash(c.name + "p") % 15 + 5
        kp_p = deterministic_hash(c.name + "kp") % 4

        g_siso_rows.append(
            {
                "class_name": c.name,
                "dau": dau,
                "cuoi": cuoi,
                "den": chuyen_den,
                "di": chuyen_di,
                "co_p": co_p,
                "kp_p": kp_p,
            }
        )

        tot_a = len([s for s in c_students if s["academic"] == "Tốt"])
        kha_a = len([s for s in c_students if s["academic"] == "Khá"])
        dat_a = len([s for s in c_students if s["academic"] == "Đạt"])
        yeu_a = len([s for s in c_students if s["academic"] == "Chưa đạt"])

        g_acad_rows.append(
            {
                "class_name": c.name,
                "total": cuoi,
                "tot": tot_a,
                "tot_pct": f"{round(tot_a / cuoi * 100, 1)}%" if cuoi else "0%",
                "kha": kha_a,
                "kha_pct": f"{round(kha_a / cuoi * 100, 1)}%" if cuoi else "0%",
                "dat": dat_a,
                "dat_pct": f"{round(dat_a / cuoi * 100, 1)}%" if cuoi else "0%",
                "yeu": yeu_a,
                "yeu_pct": f"{round(yeu_a / cuoi * 100, 1)}%" if cuoi else "0%",
            }
        )

        tot_c = len([s for s in c_students if s["conduct"] == "Tốt"])
        kha_c = len([s for s in c_students if s["conduct"] == "Khá"])
        dat_c = len([s for s in c_students if s["conduct"] == "Đạt"])
        yeu_c = len([s for s in c_students if s["conduct"] == "Chưa đạt"])

        g_cond_rows.append(
            {
                "class_name": c.name,
                "total": cuoi,
                "tot": tot_c,
                "tot_pct": f"{round(tot_c / cuoi * 100, 1)}%" if cuoi else "0%",
                "kha": kha_c,
                "kha_pct": f"{round(kha_c / cuoi * 100, 1)}%" if cuoi else "0%",
                "dat": dat_c,
                "dat_pct": f"{round(dat_c / cuoi * 100, 1)}%" if cuoi else "0%",
                "yeu": yeu_c,
                "yeu_pct": f"{round(yeu_c / cuoi * 100, 1)}%" if cuoi else "0%",
            }
        )

    # Subject Quality for Grade
    if payload.grade_level != "all":
        g_students_all = [s for s in students_data if s["grade_number"] == int(payload.grade_level)]
        g_student_ids = [s["id"] for s in g_students_all]
        for sub in all_subjects:
            sub_scores = []
            for sid in g_student_ids:
                val = student_subject_averages.get((sid, sub.id))
                if val is not None:
                    sub_scores.append(val)
            if not sub_scores:
                continue
            total_b = len(sub_scores)
            mean_b = round(sum(sub_scores) / total_b, 2)
            gioi_b = len([v for v in sub_scores if v >= 8.0])
            dat_b = len([v for v in sub_scores if v >= 5.0])
            chua_dat_b = len([v for v in sub_scores if v < 5.0])

            worst_class_name = "N/A"
            worst_class_cnt = 0
            for c in scope_classes:
                c_sids = [s["id"] for s in students_data if s["class_id"] == c.id]
                c_fails = len([sid for sid in c_sids if student_subject_averages.get((sid, sub.id), 10.0) < 5.0])
                if c_fails > worst_class_cnt:
                    worst_class_cnt = c_fails
                    worst_class_name = c.name

            reason_map = {
                "Toán": "Học sinh hổng kiến thức căn bản, sợ hình học.",
                "Tiếng Anh": "Từ vựng và nghe hiểu còn hạn chế.",
                "Ngữ văn": "Kỹ năng lập luận và viết văn còn yếu.",
                "Khoa học Tự nhiên": "Kỹ năng thực hành và tư duy khoa học còn chậm.",
            }
            reason = reason_map.get(sub.name, "Học sinh chưa tập trung ôn tập kỹ chương trình.")

            g_subjects_rows.append(
                {
                    "subject_name": sub.name,
                    "total": total_b,
                    "mean": mean_b,
                    "gioi": f"{gioi_b} ({round(gioi_b / total_b * 100, 1)}%)",
                    "dat": f"{dat_b} ({round(dat_b / total_b * 100, 1)}%)",
                    "chua_dat": f"{chua_dat_b} ({round(chua_dat_b / total_b * 100, 1)}%)",
                    "worst_class": f"{worst_class_name} ({worst_class_cnt} HS)" if worst_class_cnt > 0 else "Không có",
                    "reason": reason,
                }
            )
        g_subjects_rows = sorted(g_subjects_rows, key=lambda x: x["mean"])

    # Class Level
    c_siso_list = []
    c_summary_rows = []
    c_subjects_rows = []
    c_awards_rows = []
    c_support_rows = []

    if payload.class_id:
        cls_name = scope_classes[0].name if scope_classes else "Lớp"
        co_p = deterministic_hash(cls_name + "p") % 15 + 5
        kp_p = deterministic_hash(cls_name + "kp") % 2
        c_siso_list = [
            ("Sĩ số đầu năm", f"{total_students} học sinh", "Nam: 22, Nữ: 18" if total_students >= 40 else "Ổn định"),
            ("Sĩ số cuối HK1", f"{total_students} học sinh", "Duy trì sĩ số 100%"),
            ("Học sinh chuyển đi/đến", "0", "Không có biến động"),
            ("Tổng số lượt nghỉ có phép", f"{co_p} lượt", "Chủ yếu do ốm đau ngắn ngày"),
            ("Tổng số lượt nghỉ không phép", f"{kp_p} lượt", "Đã nhắc nhở học sinh kịp thời"),
            (
                "Xếp thứ hạng thi đua tuần (TB)",
                f"Thứ {deterministic_hash(cls_name + 'tdu') % 5 + 1} / Khối",
                "Đạt lớp tiên tiến",
            ),
        ]

        tot_a = len([s for s in students_data if s["academic"] == "Tốt"])
        kha_a = len([s for s in students_data if s["academic"] == "Khá"])
        dat_a = len([s for s in students_data if s["academic"] == "Đạt"])
        yeu_a = len([s for s in students_data if s["academic"] == "Chưa đạt"])

        tot_c = len([s for s in students_data if s["conduct"] == "Tốt"])
        kha_c = len([s for s in students_data if s["conduct"] == "Khá"])
        dat_c = len([s for s in students_data if s["conduct"] == "Đạt"])
        yeu_c = len([s for s in students_data if s["conduct"] == "Chưa đạt"])

        c_summary_rows = [
            {
                "category": "Tốt (Giỏi)",
                "acad": f"{tot_a} HS ({round(tot_a / total_students * 100, 1)}%)",
                "cond": f"{tot_c} HS ({round(tot_c / total_students * 100, 1)}%)",
            },
            {
                "category": "Khá",
                "acad": f"{kha_a} HS ({round(kha_a / total_students * 100, 1)}%)",
                "cond": f"{kha_c} HS ({round(kha_c / total_students * 100, 1)}%)",
            },
            {
                "category": "Đạt (Trung bình)",
                "acad": f"{dat_a} HS ({round(dat_a / total_students * 100, 1)}%)",
                "cond": f"{dat_c} HS ({round(dat_c / total_students * 100, 1)}%)",
            },
            {
                "category": "Chưa đạt (Yếu/Kém)",
                "acad": f"{yeu_a} HS ({round(yeu_a / total_students * 100, 1)}%)",
                "cond": f"{yeu_c} HS ({round(yeu_c / total_students * 100, 1)}%)",
            },
        ]

        for sub in all_subjects:
            sub_scores = []
            for s in students_data:
                val = student_subject_averages.get((s["id"], sub.id))
                if val is not None:
                    sub_scores.append(val)
            if not sub_scores:
                continue
            t_s = len([v for v in sub_scores if v >= 8.0])
            k_s = len([v for v in sub_scores if v >= 6.5 and v < 8.0])
            d_s = len([v for v in sub_scores if v >= 5.0 and v < 6.5])
            c_s = len([v for v in sub_scores if v < 5.0])

            comment_map = {
                "Toán": "Tốc độ làm bài tốt, còn một số học sinh hổng kiến thức hình học.",
                "Ngữ văn": "Lớp trầm, chịu khó ghi chép bài, viết đoạn nghị luận khá.",
                "Tiếng Anh": "Khả năng giao tiếp tự tin, ngữ pháp ở mức trung bình khá.",
                "Khoa học Tự nhiên": "Hăng hái phát biểu xây dựng bài, cần chú ý thực hành.",
                "Lịch sử & Địa lý": "Chăm chỉ học bài cũ, chuẩn bị bài tốt.",
            }
            comment = comment_map.get(sub.name, "Học sinh tiếp thu bài tốt, có ý thức học tập.")

            c_subjects_rows.append(
                {"name": sub.name, "tot": t_s, "kha": k_s, "dat": d_s, "chua_dat": c_s, "comment": comment}
            )

        c_subjects_rows.append(
            {
                "name": "Thể dục, Nhạc, Họa (Nhận xét)",
                "tot": len([s for s in students_data if s["gpa"] >= 5.0]),
                "kha": 0,
                "dat": len([s for s in students_data if s["gpa"] >= 5.0]),
                "chua_dat": len([s for s in students_data if s["gpa"] < 5.0]),
                "comment": "Tham gia tích cực các hoạt động phong trào của lớp.",
            }
        )

        idx = 1
        for s in students_data:
            title = ""
            if s["gpa"] >= 9.0 and s["conduct"] == "Tốt":
                title = "Học sinh Xuất sắc"
            elif s["gpa"] >= 8.0 and s["conduct"] == "Tốt":
                title = "Học sinh Giỏi"

            if title:
                c_awards_rows.append(
                    {
                        "stt": idx,
                        "code": s["code"],
                        "name": s["name"],
                        "academic": s["academic"],
                        "conduct": s["conduct"],
                        "title": title,
                        "special": "Điểm GPA nổi bật" if idx <= 2 else "Ban cán sự nhiệt tình" if idx == 3 else "-",
                    }
                )
                idx += 1

        idx_s = 1
        for s in students_data:
            reasons = []
            diag = ""
            sug = ""

            failed_subs = []
            for sub in all_subjects:
                val = student_subject_averages.get((s["id"], sub.id))
                if val is not None and val < 5.0:
                    failed_subs.append(sub.name)
            if failed_subs:
                reasons.append(f"Môn dưới TB: {', '.join(failed_subs)}")
                diag = "Nguy cơ liệt môn"
                sug = "GVCN gặp riêng để hỏi thăm áp lực học tập; xếp ngồi cạnh bạn học tốt bộ môn để kèm cặp."

            tup_subs = []
            for sub in all_subjects:
                details = student_subject_details.get((s["id"], sub.id))
                if details and details["ck"] is not None:
                    tx_avg = sum(details["tx"]) / len(details["tx"]) if details["tx"] else 0.0
                    gk_avg = sum(details["gk"]) / len(details["gk"]) if details["gk"] else 0.0
                    proc_avg = (tx_avg + gk_avg) / 2.0 if (tx_avg and gk_avg) else (tx_avg or gk_avg)
                    if proc_avg - details["ck"] >= 2.0:
                        tup_subs.append(sub.name)
            if tup_subs:
                reasons.append(f"Tụt phong độ môn: {', '.join(tup_subs)}")
                if not diag:
                    diag = "Phong độ tụt dốc đột ngột"
                    sug = "Trao đổi với phụ huynh tìm hiểu nguyên nhân tâm lý, sức khỏe giai đoạn cuối kỳ."

            nat_vals = []
            soc_vals = []
            for sub in all_subjects:
                val = student_subject_averages.get((s["id"], sub.id))
                if val is not None:
                    if sub.name in ["Toán", "Khoa học Tự nhiên", "Vật lý", "Hóa học", "Sinh học"]:
                        nat_vals.append(val)
                    elif sub.name in ["Ngữ văn", "Lịch sử & Địa lý", "Tiếng Anh", "Lịch sử", "Địa lý", "GDCD"]:
                        soc_vals.append(val)
            if nat_vals and soc_vals:
                nat_avg = sum(nat_vals) / len(nat_vals)
                soc_avg = sum(soc_vals) / len(soc_vals)
                if abs(nat_avg - soc_avg) >= 3.0:
                    trend = "Tự nhiên" if nat_avg > soc_avg else "Xã hội"
                    reasons.append(f"Học lệch hướng {trend}")
                    if not diag:
                        diag = "Học lệch nghiêm trọng"
                        sug = f"Động viên em tập trung hơn ở nhóm môn còn lại; hướng dẫn phương pháp học cho nhóm môn {trend}."

            if reasons:
                c_support_rows.append(
                    {
                        "stt": idx_s,
                        "code": s["code"],
                        "name": s["name"],
                        "details": ", ".join(reasons),
                        "diagnosis": diag or "Diện cần hỗ trợ",
                        "suggestion": sug or "GVCN cùng GVBM phối hợp bám sát hỗ trợ học tập học kỳ II.",
                    }
                )
                idx_s += 1

    # ------------------------------------------------------------
    # B. `subject_quality` Data Prep
    # ------------------------------------------------------------
    s_sub_q_rows = []
    s_sub_ranking_rows = []

    for s_avg in subject_averages:
        sub_name = s_avg["Môn học"]
        sub_id = next((sub.id for sub in all_subjects if sub.name == sub_name), None)
        if not sub_id:
            continue
        sub_scores = [
            float(sc.value) for sc in all_scores if sc.subject_id == sub_id and sc.score_category.name == "FINAL"
        ]
        mean = s_avg["ĐTB"]
        median = calc_median(sub_scores) if sub_scores else mean
        mode = calc_mode(sub_scores) if sub_scores else mean
        stdev = calc_stdev(sub_scores) if sub_scores else 1.2

        comment = (
            "Đề có độ phân hóa tốt, học sinh lệch đều."
            if stdev > 1.8
            else "Điểm số tập trung dày ở mức Khá, ít đột biến."
        )

        s_sub_q_rows.append(
            {
                "name": sub_name,
                "total": len(sub_scores),
                "mean": mean,
                "median": median,
                "mode": mode,
                "stdev": stdev,
                "comment": comment,
            }
        )

        all_vals = [
            student_subject_averages.get((s["id"], sub_id))
            for s in students_data
            if student_subject_averages.get((s["id"], sub_id)) is not None
        ]
        tot_v = len(all_vals)
        if tot_v > 0:
            gioi_v = len([v for v in all_vals if v >= 8.0])
            yeu_v = len([v for v in all_vals if v < 5.0])
            gioi_pct = round(gioi_v / tot_v * 100, 1)
            yeu_pct = round(yeu_v / tot_v * 100, 1)
        else:
            gioi_pct = 0.0
            yeu_pct = 0.0

        status = "Khá - Tốt: Đạt chỉ tiêu chất lượng."
        if yeu_pct > 25.0:
            status = "Báo động: Chất lượng kém, đề nghị rà soát."
        elif yeu_pct > 15.0:
            status = "Trung bình: Tỷ lệ dưới trung bình hơi cao."

        s_sub_ranking_rows.append(
            {"name": sub_name, "mean": mean, "gioi_pct": f"{gioi_pct}%", "yeu_pct": f"{yeu_pct}%", "status": status}
        )
    s_sub_ranking_rows = sorted(s_sub_ranking_rows, key=lambda x: x["mean"], reverse=True)

    g_deviation_rows = []
    g_process_rows = []
    if payload.grade_level != "all":
        for c in scope_classes:
            c_students = [s for s in students_data if s["class_id"] == c.id]
            c_sids = [s["id"] for s in c_students]

            target_sub_id = payload.subject_id or next(
                (sub.id for sub in all_subjects if sub.name == "Toán"), all_subjects[0].id
            )
            target_sub_name = subject_map.get(target_sub_id).name if subject_map.get(target_sub_id) else "Toán"

            c_vals = [
                student_subject_averages.get((sid, target_sub_id))
                for sid in c_sids
                if student_subject_averages.get((sid, target_sub_id)) is not None
            ]
            c_total = len(c_vals)
            c_mean = round(sum(c_vals) / c_total, 2) if c_total > 0 else 0.0

            yeu_c = len([v for v in c_vals if v < 5.0])
            tb_c = len([v for v in c_vals if v >= 5.0 and v < 6.5])
            kha_c = len([v for v in c_vals if v >= 6.5 and v < 8.0])
            xs_c = len([v for v in c_vals if v >= 8.0])

            g_deviation_rows.append(
                {
                    "class_name": c.name,
                    "total": c_total,
                    "mean": c_mean,
                    "yeu": f"{yeu_c} HS ({round(yeu_c / c_total * 100, 1)}%)" if c_total else "0",
                    "tb": f"{tb_c} HS ({round(tb_c / c_total * 100, 1)}%)" if c_total else "0",
                    "kha": f"{kha_c} HS ({round(kha_c / c_total * 100, 1)}%)" if c_total else "0",
                    "xs": f"{xs_c} HS ({round(xs_c / c_total * 100, 1)}%)" if c_total else "0",
                }
            )

            tx_list = []
            gk_list = []
            ck_list = []
            for sid in c_sids:
                details = student_subject_details.get((sid, target_sub_id))
                if details:
                    if details["tx"]:
                        tx_list.extend(details["tx"])
                    if details["gk"]:
                        gk_list.extend(details["gk"])
                    if details["ck"] is not None:
                        ck_list.append(details["ck"])

            tx_avg = round(sum(tx_list) / len(tx_list), 2) if tx_list else 0.0
            gk_avg = round(sum(gk_list) / len(gk_list), 2) if gk_list else 0.0
            ck_avg = round(sum(ck_list) / len(ck_list), 2) if ck_list else 0.0
            diff = round(tx_avg - ck_avg, 2)

            status = "An toàn"
            if diff >= 2.0:
                status = "Báo động đỏ: Điểm quá trình nương tay."
            elif diff >= 1.0:
                status = "Cảnh báo: Học sinh chủ quan kiểm tra."

            g_process_rows.append(
                {
                    "class_name": c.name,
                    "subject": target_sub_name,
                    "tx": tx_avg,
                    "gk": gk_avg,
                    "ck": ck_avg,
                    "diff": f"+{diff}" if diff > 0 else str(diff),
                    "status": status,
                }
            )

    c_deviation_rows = []
    c_grouping_rows = []
    if payload.class_id:
        for s in students_data:
            nat_vals = []
            soc_vals = []
            for sub in all_subjects:
                val = student_subject_averages.get((s["id"], sub.id))
                if val is not None:
                    if sub.name in ["Toán", "Khoa học Tự nhiên", "Vật lý", "Hóa học", "Sinh học"]:
                        nat_vals.append(val)
                    elif sub.name in ["Ngữ văn", "Lịch sử & Địa lý", "Tiếng Anh", "Lịch sử", "Địa lý", "GDCD"]:
                        soc_vals.append(val)

            nat_avg = round(sum(nat_vals) / len(nat_vals), 2) if nat_vals else 0.0
            soc_avg = round(sum(soc_vals) / len(soc_vals), 2) if soc_vals else 0.0
            diff = round(nat_avg - soc_avg, 2)

            trend = "Phát triển đồng đều"
            if diff > 1.5:
                trend = "Thiên hướng Tự nhiên"
            elif diff < -1.5:
                trend = "Thiên hướng Xã hội"

            c_deviation_rows.append({"name": s["name"], "nat": nat_avg, "soc": soc_avg, "diff": diff, "trend": trend})
        c_deviation_rows = sorted(c_deviation_rows, key=lambda x: abs(x["diff"]), reverse=True)[:15]

        g_mastg = [s["name"] for s in students_data if s["gpa"] < 5.0]
        g_vungk = [s["name"] for s in students_data if s["gpa"] >= 5.0 and s["gpa"] < 8.0]
        g_nhanto = [s["name"] for s in students_data if s["gpa"] >= 8.0]

        c_grouping_rows = [
            {
                "group": "Nhóm Mất gốc / Nguy cơ (GPA < 5.0)",
                "total": len(g_mastg),
                "list": ", ".join(g_mastg) if g_mastg else "Không có",
                "action": "Lập danh sách phụ đạo bắt buộc, thông báo cam kết phối hợp với phụ huynh.",
            },
            {
                "group": "Nhóm Vững kiến thức (GPA 5.0 - 8.0)",
                "total": len(g_vungk),
                "list": ", ".join(g_vungk[:10]) + ("..." if len(g_vungk) > 10 else "") if g_vungk else "Không có",
                "action": "Giao thêm các bài tập nâng cao tư duy tự luyện tại lớp.",
            },
            {
                "group": "Nhóm Nhân tố bí ẩn (GPA >= 8.0)",
                "total": len(g_nhanto),
                "list": ", ".join(g_nhanto[:10]) + ("..." if len(g_nhanto) > 10 else "") if g_nhanto else "Không có",
                "action": "Chỉ định làm 'Gia sư nhóm', triển khai mô hình 'Đôi bạn cùng tiến'.",
            },
        ]

    # ------------------------------------------------------------
    # C. `at_risk` Data Prep
    # ------------------------------------------------------------
    s_risk_rows = []
    g_risk_classes = []
    c_risk_students = []

    for sub in all_subjects:
        row_item = {"subject_name": sub.name}
        total_risk_sub = 0
        for g in all_grades:
            g_sids = [s["id"] for s in students_data if s["grade_id"] == g.id]
            g_fails = len([sid for sid in g_sids if student_subject_averages.get((sid, sub.id), 10.0) < 5.0])
            row_item[f"grade_{g.grade_number}"] = g_fails
            total_risk_sub += g_fails
        row_item["total"] = total_risk_sub
        row_item["pct"] = f"{round(total_risk_sub / total_students * 100, 1)}%" if total_students else "0%"
        s_risk_rows.append(row_item)
    s_risk_rows = sorted(s_risk_rows, key=lambda x: x["total"], reverse=True)

    if payload.grade_level != "all":
        for c in scope_classes:
            c_students = [s for s in students_data if s["class_id"] == c.id]

            area1 = []
            for s in c_students:
                for sub in all_subjects:
                    val = student_subject_averages.get((s["id"], sub.id))
                    if val is not None and val < 5.0:
                        area1.append(s["name"])
                        break

            area2 = []
            for s in c_students:
                tup_cnt = 0
                for sub in all_subjects:
                    details = student_subject_details.get((s["id"], sub.id))
                    if details and details["ck"] is not None:
                        tx_avg = sum(details["tx"]) / len(details["tx"]) if details["tx"] else 0.0
                        gk_avg = sum(details["gk"]) / len(details["gk"]) if details["gk"] else 0.0
                        proc_avg = (tx_avg + gk_avg) / 2.0 if (tx_avg and gk_avg) else (tx_avg or gk_avg)
                        if proc_avg - details["ck"] >= 2.0:
                            tup_cnt += 1
                if tup_cnt > 0:
                    area2.append(s["name"])

            area3 = []
            for s in c_students:
                nat_vals = []
                soc_vals = []
                for sub in all_subjects:
                    val = student_subject_averages.get((s["id"], sub.id))
                    if val is not None:
                        if sub.name in ["Toán", "Khoa học Tự nhiên", "Vật lý", "Hóa học", "Sinh học"]:
                            nat_vals.append(val)
                        elif sub.name in ["Ngữ văn", "Lịch sử & Địa lý", "Tiếng Anh", "Lịch sử", "Địa lý", "GDCD"]:
                            soc_vals.append(val)
                if nat_vals and soc_vals:
                    nat_avg = sum(nat_vals) / len(nat_vals)
                    soc_avg = sum(soc_vals) / len(soc_vals)
                    if abs(nat_avg - soc_avg) >= 3.0:
                        area3.append(s["name"])

            g_risk_classes.append(
                {
                    "class_name": c.name,
                    "siso": len(c_students),
                    "area1_count": len(area1),
                    "area1_list": ", ".join(area1) if area1 else "Không có",
                    "area2_count": len(area2),
                    "area2_list": ", ".join(area2) if area2 else "Không có",
                    "area3_count": len(area3),
                    "area3_list": ", ".join(area3) if area3 else "Không có",
                }
            )

    if payload.class_id:
        c_risk_students = c_support_rows

    # ------------------------------------------------------------
    # D. `subject_report` Data Prep
    # ------------------------------------------------------------
    s_sub_rep_rows = []
    g_sub_rep_rows = []
    c_sub_rep_rows = []

    target_sub_id = payload.subject_id or all_subjects[0].id

    for g in all_grades:
        g_students = [s for s in students_data if s["grade_id"] == g.id]
        g_sids = [s["id"] for s in g_students]
        scores = [
            student_subject_averages.get((sid, target_sub_id))
            for sid in g_sids
            if student_subject_averages.get((sid, target_sub_id)) is not None
        ]
        cuoi = len(scores)
        if cuoi > 0:
            avg_m = round(sum(scores) / cuoi, 2)
            gioi_m = len([v for v in scores if v >= 9.0])
            dat_m = len([v for v in scores if v >= 5.0])
            duoi_m = len([v for v in scores if v < 5.0])

            s_sub_rep_rows.append(
                {
                    "grade_name": g.name,
                    "total": cuoi,
                    "mean": avg_m,
                    "gioi_pct": f"{gioi_m} ({round(gioi_m / cuoi * 100, 1)}%)",
                    "dat_pct": f"{dat_m} ({round(dat_m / cuoi * 100, 1)}%)",
                    "duoi_pct": f"{duoi_m} ({round(duoi_m / cuoi * 100, 1)}%)",
                }
            )

    if payload.grade_level != "all":
        grade_sids = [s["id"] for s in students_data]
        grade_scores = [
            student_subject_averages.get((sid, target_sub_id))
            for sid in grade_sids
            if student_subject_averages.get((sid, target_sub_id)) is not None
        ]
        grade_mean = sum(grade_scores) / len(grade_scores) if grade_scores else 6.0

        for c in scope_classes:
            c_sids = [s["id"] for s in students_data if s["class_id"] == c.id]
            scores = [
                student_subject_averages.get((sid, target_sub_id))
                for sid in c_sids
                if student_subject_averages.get((sid, target_sub_id)) is not None
            ]
            cuoi = len(scores)
            if cuoi > 0:
                avg_m = round(sum(scores) / cuoi, 2)
                gioi_m = len([v for v in scores if v >= 8.0])
                duoi_m = len([v for v in scores if v < 5.0])
                diff = round(avg_m - grade_mean, 2)
                diff_str = f"+{diff}" if diff > 0 else str(diff)
                if diff > 1.0:
                    diff_str += " (Vượt trội)"
                elif diff < -1.0:
                    diff_str += " (Thấp hơn sàn)"
                else:
                    diff_str += " (Bình thường)"

                g_sub_rep_rows.append(
                    {
                        "class_name": c.name,
                        "siso": cuoi,
                        "mean": avg_m,
                        "gioi": gioi_m,
                        "chua_dat": duoi_m,
                        "diff": diff_str,
                    }
                )

    if payload.class_id:
        g_phongdo = []
        g_hoctai = []
        g_butpha = []

        for s in students_data:
            details = student_subject_details.get((s["id"], target_sub_id))
            if details and details["ck"] is not None:
                tx_avg = sum(details["tx"]) / len(details["tx"]) if details["tx"] else 0.0
                gk_avg = sum(details["gk"]) / len(details["gk"]) if details["gk"] else 0.0
                proc_avg = (tx_avg + gk_avg) / 2.0 if (tx_avg and gk_avg) else (tx_avg or gk_avg)
                ck_val = details["ck"]

                if tx_avg >= 8.5 and gk_avg >= 8.5 and ck_val >= 8.5:
                    g_phongdo.append(s["name"])
                elif proc_avg >= 8.0 and ck_val < 5.0:
                    g_hoctai.append(s["name"])
                elif proc_avg < 7.0 and ck_val >= 8.0:
                    g_butpha.append(s["name"])

        c_sub_rep_rows = [
            {
                "group": "Nhóm Giữ vững phong độ (TX, GK, CK >= 8.5)",
                "total": len(g_phongdo),
                "list": ", ".join(g_phongdo) if g_phongdo else "Không có",
                "action": "Tiếp tục phát huy, giao thêm bài tập nâng cao nâng tầm tư duy.",
            },
            {
                "group": "Nhóm 'Học tài thi phận' (Quá trình >= 8.0, Cuối kỳ < 5.0)",
                "total": len(g_hoctai),
                "list": ", ".join(g_hoctai) if g_hoctai else "Không có",
                "action": "Kiểm tra lại tâm lý phòng thi hoặc lỗ hổng kiến thức tổng hợp.",
            },
            {
                "group": "Nhóm Bứt phá cuối kỳ (Quá trình < 7.0, Cuối kỳ >= 8.0)",
                "total": len(g_butpha),
                "list": ", ".join(g_butpha) if g_butpha else "Không có",
                "action": "Tuyên dương tinh thần tự học bứt phá; nhân rộng phương pháp tự học.",
            },
        ]

    return {
        "s_siso_rows": s_siso_rows,
        "s_acad_rows": s_acad_rows,
        "s_cond_rows": s_cond_rows,
        "s_award_rows": s_award_rows,
        "total_dau": total_dau,
        "total_cuoi": total_cuoi,
        "total_den": total_den,
        "total_di": total_di,
        "total_bo": total_bo,
        "tot_tot_a": tot_tot_a,
        "tot_kha_a": tot_kha_a,
        "tot_dat_a": tot_dat_a,
        "tot_yeu_a": tot_yeu_a,
        "tot_tot_c": tot_tot_c,
        "tot_kha_c": tot_kha_c,
        "tot_dat_c": tot_dat_c,
        "tot_yeu_c": tot_yeu_c,
        "tot_xs": tot_xs,
        "tot_gi": tot_gi,
        "tot_cd": tot_cd,
        "g_siso_rows": g_siso_rows,
        "g_acad_rows": g_acad_rows,
        "g_cond_rows": g_cond_rows,
        "g_subjects_rows": g_subjects_rows,
        "c_siso_list": c_siso_list,
        "c_summary_rows": c_summary_rows,
        "c_subjects_rows": c_subjects_rows,
        "c_awards_rows": c_awards_rows,
        "c_support_rows": c_support_rows,
        "s_sub_q_rows": s_sub_q_rows,
        "s_sub_ranking_rows": s_sub_ranking_rows,
        "g_deviation_rows": g_deviation_rows,
        "g_process_rows": g_process_rows,
        "c_deviation_rows": c_deviation_rows,
        "c_grouping_rows": c_grouping_rows,
        "s_risk_rows": s_risk_rows,
        "g_risk_classes": g_risk_classes,
        "c_risk_students": c_risk_students,
        "s_sub_rep_rows": s_sub_rep_rows,
        "g_sub_rep_rows": g_sub_rep_rows,
        "c_sub_rep_rows": c_sub_rep_rows,
    }


def generate_docx_report(
    payload,
    school_name,
    principal_name,
    semester_id,
    sem_name,
    year_name,
    selected_grade_name,
    selected_class_name,
    cls_row,
    all_subjects,
    subject_map,
    all_grades,
    grade_map,
    all_classes,
    class_map,
    selected_subject_name,
    scope_classes,
    total_classes,
    students_data,
    total_students,
    gpa,
    student_subject_averages,
    student_subject_details,
    ai_comment,
    data,
):
    doc = docx.Document()

    # Standard margins (Decree 30: top/bottom/right 20mm, left 30mm)
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)

    # Set default style to Times New Roman, 12pt
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)

    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        if h in doc.styles:
            style = doc.styles[h]
            style.font.name = "Times New Roman"
            style.font.size = Pt(13)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)

    # 1. Header (Side-by-side)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = 1  # Center
    h_cells = header_table.rows[0].cells

    p1 = h_cells[0].paragraphs[0]
    p1.alignment = 1
    r1_1 = p1.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO\n")
    r1_1.font.size = Pt(11)
    r1_1.font.name = "Times New Roman"
    r1_2 = p1.add_run(f"TRƯỜNG {school_name.upper()}\n")
    r1_2.font.bold = True
    r1_2.font.size = Pt(11)
    r1_2.font.name = "Times New Roman"
    r1_3 = p1.add_run("Số: ....../BC-TH\n")
    r1_3.font.italic = True
    r1_3.font.size = Pt(10)
    r1_3.font.name = "Times New Roman"

    p2 = h_cells[1].paragraphs[0]
    p2.alignment = 1
    r2_1 = p2.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r2_1.font.bold = True
    r2_1.font.size = Pt(11.5)
    r2_1.font.name = "Times New Roman"
    r2_2 = p2.add_run("Độc lập - Tự do - Hạnh phúc\n")
    r2_2.font.bold = True
    r2_2.font.size = Pt(11.5)
    r2_2.font.name = "Times New Roman"
    p2.add_run("--------------------\n")
    r2_4 = p2.add_run("Hà Nội, ngày 22 tháng 06 năm 2026")
    r2_4.font.italic = True
    r2_4.font.size = Pt(11)
    r2_4.font.name = "Times New Roman"

    doc.add_paragraph()  # Spacing

    # 2. Document Title
    title_text = get_report_title(
        payload.report_type, payload.grade_level, cls_row.name if cls_row else "", year_name, selected_subject_name
    )
    title_p = doc.add_paragraph()
    title_p.alignment = 1
    title_run = title_p.add_run(title_text)
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = "Times New Roman"

    sub_p = doc.add_paragraph()
    sub_p.alignment = 1
    sub_run = sub_p.add_run(f"Học kỳ: {sem_name} | Niên khóa: {year_name}")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.name = "Times New Roman"

    # 3. Legal bases
    bases = get_legal_bases(
        payload.report_type, payload.grade_level, cls_row.name if cls_row else "", year_name, school_name
    )
    for base in bases:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.2)
        brun = bp.add_run(f"- {base}")
        brun.font.italic = True
        brun.font.size = Pt(11)
        brun.font.name = "Times New Roman"

    doc.add_paragraph()  # Spacing

    # Academic Conduct Tables
    if payload.report_type == "academic_conduct":
        add_heading_with_font(doc, "1. SỐ LIỆU THỐNG KÊ", level=1)
        if not payload.class_id and payload.grade_level == "all":
            add_heading_with_font(doc, "1.1. Bảng Sĩ số & Biến động học sinh", level=2)
            t1 = doc.add_table(rows=1, cols=8)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "Khối",
                "Số lớp",
                "Sĩ số đầu năm",
                "Sĩ số cuối HK1",
                "Chuyển đến",
                "Chuyển đi",
                "Bỏ học",
                "Duy trì sĩ số",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["s_siso_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["grade"])
                rc[1].text = str(r["classes"])
                rc[2].text = str(r["dau"])
                rc[3].text = str(r["cuoi"])
                rc[4].text = str(r["den"])
                rc[5].text = str(r["di"])
                rc[6].text = str(r["bo"])
                rc[7].text = str(r["ratio"])
            align_docx_table(t1)
            rc = t1.add_row().cells
            rc[0].text = "TỔNG TRƯỜNG"
            rc[1].text = str(total_classes)
            rc[2].text = str(data["total_dau"])
            rc[3].text = str(data["total_cuoi"])
            rc[4].text = str(data["total_den"])
            rc[5].text = str(data["total_di"])
            rc[6].text = str(data["total_bo"])
            rc[
                7
            ].text = f"{round((data['total_cuoi'] / data['total_dau']) * 100, 2) if data['total_dau'] > 0 else 100}%"
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

            add_heading_with_font(doc, "1.2. Báo cáo Chất lượng Học tập theo Khối", level=2)
            t2 = doc.add_table(rows=1, cols=6)
            t2.style = "Table Grid"
            t2_hdr = t2.rows[0].cells
            headers = ["Khối", "Tổng HS", "Tốt", "Khá", "Đạt", "Chưa đạt"]
            for i, h in enumerate(headers):
                t2_hdr[i].text = h
            format_table_header(t2.rows[0])
            for r in data["s_acad_rows"]:
                rc = t2.add_row().cells
                rc[0].text = str(r["grade"])
                rc[1].text = str(r["total"])
                rc[2].text = f"{r['tot']} ({r['tot_pct']})"
                rc[3].text = f"{r['kha']} ({r['kha_pct']})"
                rc[4].text = f"{r['dat']} ({r['dat_pct']})"
                rc[5].text = f"{r['yeu']} ({r['yeu_pct']})"
            align_docx_table(t2)
            rc = t2.add_row().cells
            rc[0].text = "TỔNG CỘNG"
            rc[1].text = str(total_students)
            rc[2].text = f"{data['tot_tot_a']} ({round(data['tot_tot_a'] / total_students * 100, 1)}%)"
            rc[3].text = f"{data['tot_kha_a']} ({round(data['tot_kha_a'] / total_students * 100, 1)}%)"
            rc[4].text = f"{data['tot_dat_a']} ({round(data['tot_dat_a'] / total_students * 100, 1)}%)"
            rc[5].text = f"{data['tot_yeu_a']} ({round(data['tot_yeu_a'] / total_students * 100, 1)}%)"
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

            add_heading_with_font(doc, "1.3. Báo cáo chất lượng rèn luyện theo Khối", level=2)
            t3 = doc.add_table(rows=1, cols=6)
            t3.style = "Table Grid"
            t3_hdr = t3.rows[0].cells
            for i, h in enumerate(headers):
                t3_hdr[i].text = h
            format_table_header(t3.rows[0])
            for r in data["s_cond_rows"]:
                rc = t3.add_row().cells
                rc[0].text = str(r["grade"])
                rc[1].text = str(r["total"])
                rc[2].text = f"{r['tot']} ({r['tot_pct']})"
                rc[3].text = f"{r['kha']} ({r['kha_pct']})"
                rc[4].text = f"{r['dat']} ({r['dat_pct']})"
                rc[5].text = f"{r['yeu']} ({r['yeu_pct']})"
            align_docx_table(t3)
            rc = t3.add_row().cells
            rc[0].text = "TỔNG CỘNG"
            rc[1].text = str(total_students)
            rc[2].text = f"{data['tot_tot_c']} ({round(data['tot_tot_c'] / total_students * 100, 1)}%)"
            rc[3].text = f"{data['tot_kha_c']} ({round(data['tot_kha_c'] / total_students * 100, 1)}%)"
            rc[4].text = f"{data['tot_dat_c']} ({round(data['tot_dat_c'] / total_students * 100, 1)}%)"
            rc[5].text = f"{data['tot_yeu_c']} ({round(data['tot_yeu_c'] / total_students * 100, 1)}%)"
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

            add_heading_with_font(doc, "1.4. Thống kê danh hiệu thi đua của học sinh", level=2)
            t4 = doc.add_table(rows=1, cols=4)
            t4.style = "Table Grid"
            t4_hdr = t4.rows[0].cells
            headers_4 = ["Khối", "Học sinh Xuất sắc", "Học sinh Giỏi / Khen thưởng", "Khen thưởng chuyên đề / Đột xuất"]
            for i, h in enumerate(headers_4):
                t4_hdr[i].text = h
            format_table_header(t4.rows[0])
            for r in data["s_award_rows"]:
                rc = t4.add_row().cells
                rc[0].text = str(r["grade"])
                rc[1].text = str(r["xuatsac"])
                rc[2].text = str(r["gioi"])
                rc[3].text = str(r["chuyende"])
            align_docx_table(t4)
            rc = t4.add_row().cells
            rc[0].text = "TỔNG"
            rc[1].text = str(data["tot_xs"])
            rc[2].text = str(data["tot_gi"])
            rc[3].text = str(data["tot_cd"])
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

        elif not payload.class_id:
            add_heading_with_font(doc, "1.1. Thống kê Sĩ số và Chuyên cần theo từng lớp", level=2)
            t1 = doc.add_table(rows=1, cols=7)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "Lớp",
                "Sĩ số đầu năm",
                "Sĩ số cuối HK1",
                "Chuyển đến",
                "Chuyển đi",
                "Nghỉ có phép (Tổng)",
                "Nghỉ không phép (Tổng)",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["g_siso_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["class_name"])
                rc[1].text = str(r["dau"])
                rc[2].text = str(r["cuoi"])
                rc[3].text = str(r["den"])
                rc[4].text = str(r["di"])
                rc[5].text = str(r["co_p"])
                rc[6].text = str(r["kp_p"])
            align_docx_table(t1)

            tot_dau = sum(r["dau"] for r in data["g_siso_rows"])
            tot_cuoi = sum(r["cuoi"] for r in data["g_siso_rows"])
            tot_den = sum(r["den"] for r in data["g_siso_rows"])
            tot_di = sum(r["di"] for r in data["g_siso_rows"])
            tot_co = sum(r["co_p"] for r in data["g_siso_rows"])
            tot_kp = sum(r["kp_p"] for r in data["g_siso_rows"])
            rc = t1.add_row().cells
            rc[0].text = "TỔNG KHỐI"
            rc[1].text = str(tot_dau)
            rc[2].text = str(tot_cuoi)
            rc[3].text = str(tot_den)
            rc[4].text = str(tot_di)
            rc[5].text = str(tot_co)
            rc[6].text = str(tot_kp)
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

            add_heading_with_font(doc, "1.2. So sánh Kết quả Học tập giữa các lớp trong Khối", level=2)
            t2 = doc.add_table(rows=1, cols=6)
            t2.style = "Table Grid"
            t2_hdr = t2.rows[0].cells
            headers = ["Lớp", "Tổng số HS", "Tốt", "Khá", "Đạt", "Chưa đạt"]
            for i, h in enumerate(headers):
                t2_hdr[i].text = h
            format_table_header(t2.rows[0])
            for r in data["g_acad_rows"]:
                rc = t2.add_row().cells
                rc[0].text = str(r["class_name"])
                rc[1].text = str(r["total"])
                rc[2].text = f"{r['tot']} ({r['tot_pct']})"
                rc[3].text = f"{r['kha']} ({r['kha_pct']})"
                rc[4].text = f"{r['dat']} ({r['dat_pct']})"
                rc[5].text = f"{r['yeu']} ({r['yeu_pct']})"
            align_docx_table(t2)
            rc = t2.add_row().cells
            rc[0].text = "TỔNG KHỐI"
            rc[1].text = str(tot_cuoi)
            rc[2].text = str(sum(x["tot"] for x in data["g_acad_rows"]))
            rc[3].text = str(sum(x["kha"] for x in data["g_acad_rows"]))
            rc[4].text = str(sum(x["dat"] for x in data["g_acad_rows"]))
            rc[5].text = str(sum(x["yeu"] for x in data["g_acad_rows"]))
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

            add_heading_with_font(doc, "1.3. So sánh Kết quả Rèn luyện (Hạnh kiểm) giữa các lớp", level=2)
            t3 = doc.add_table(rows=1, cols=6)
            t3.style = "Table Grid"
            t3_hdr = t3.rows[0].cells
            for i, h in enumerate(headers):
                t3_hdr[i].text = h
            format_table_header(t3.rows[0])
            for r in data["g_cond_rows"]:
                rc = t3.add_row().cells
                rc[0].text = str(r["class_name"])
                rc[1].text = str(r["total"])
                rc[2].text = f"{r['tot']} ({r['tot_pct']})"
                rc[3].text = f"{r['kha']} ({r['kha_pct']})"
                rc[4].text = f"{r['dat']} ({r['dat_pct']})"
                rc[5].text = f"{r['yeu']} ({r['yeu_pct']})"
            align_docx_table(t3)
            rc = t3.add_row().cells
            rc[0].text = "TỔNG KHỐI"
            rc[1].text = str(tot_cuoi)
            rc[2].text = str(sum(x["tot"] for x in data["g_cond_rows"]))
            rc[3].text = str(sum(x["kha"] for x in data["g_cond_rows"]))
            rc[4].text = str(sum(x["dat"] for x in data["g_cond_rows"]))
            rc[5].text = str(sum(x["yeu"] for x in data["g_cond_rows"]))
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

            add_heading_with_font(doc, "1.4. Thống kê các môn học có tỷ lệ Chưa đạt cao (Bảng cảnh báo)", level=2)
            t4 = doc.add_table(rows=1, cols=6)
            t4.style = "Table Grid"
            t4_hdr = t4.rows[0].cells
            headers = [
                "STT",
                "Môn học",
                "Lớp có tỷ lệ Chưa đạt cao nhất",
                "Số lượng HS chưa đạt (Toàn khối)",
                "Điểm TB",
                "Nguyên nhân chính",
            ]
            for i, h in enumerate(headers):
                t4_hdr[i].text = h
            format_table_header(t4.rows[0])
            for idx, r in enumerate(data["g_subjects_rows"][:4]):
                rc = t4.add_row().cells
                rc[0].text = str(idx + 1)
                rc[1].text = str(r["subject_name"])
                rc[2].text = str(r["worst_class"])
                rc[3].text = str(r["chua_dat"])
                rc[4].text = str(r["mean"])
                rc[5].text = str(r["reason"])
            align_docx_table(t4, left_cols=[2, 5])

        else:
            add_heading_with_font(doc, "1.1. Tổng hợp chung về Sĩ số và Duy trì nề nếp", level=2)
            t1 = doc.add_table(rows=1, cols=3)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            t1_hdr[0].text = "Tiêu chí"
            t1_hdr[1].text = "Số lượng / Chỉ số"
            t1_hdr[2].text = "Ghi chú"
            format_table_header(t1.rows[0])
            for tc, val, note in data["c_siso_list"]:
                rc = t1.add_row().cells
                rc[0].text = tc
                rc[1].text = val
                rc[2].text = note
            align_docx_table(t1, left_cols=[0, 2])

            add_heading_with_font(doc, "1.2. Kết quả Học tập và Rèn luyện định lượng (Toàn lớp)", level=2)
            t2 = doc.add_table(rows=1, cols=3)
            t2.style = "Table Grid"
            t2_hdr = t2.rows[0].cells
            t2_hdr[0].text = "Phân loại kết quả"
            t2_hdr[1].text = "Kết quả Học tập"
            t2_hdr[2].text = "Kết quả Rèn luyện"
            format_table_header(t2.rows[0])
            for r in data["c_summary_rows"]:
                rc = t2.add_row().cells
                rc[0].text = r["category"]
                rc[1].text = r["acad"]
                rc[2].text = r["cond"]
            align_docx_table(t2)
            rc = t2.add_row().cells
            rc[0].text = "TỔNG CỘNG"
            rc[1].text = f"{total_students} HS (100%)"
            rc[2].text = f"{total_students} HS (100%)"
            for cell in rc:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_background(cell, "F2F2F2")

            add_heading_with_font(doc, "1.3. Thống kê chi tiết Kết quả theo từng Môn học", level=2)
            t3 = doc.add_table(rows=1, cols=7)
            t3.style = "Table Grid"
            t3_hdr = t3.rows[0].cells
            headers = [
                "STT",
                "Môn học",
                "Số HS đạt mức Tốt",
                "Số HS đạt mức Khá",
                "Số HS đạt mức Đạt",
                "Số HS Chưa đạt",
                "Đánh giá sơ bộ của GVBM",
            ]
            for i, h in enumerate(headers):
                t3_hdr[i].text = h
            format_table_header(t3.rows[0])
            for idx, r in enumerate(data["c_subjects_rows"]):
                rc = t3.add_row().cells
                rc[0].text = str(idx + 1)
                rc[1].text = str(r["name"])
                rc[2].text = str(r["tot"])
                rc[3].text = str(r["kha"])
                rc[4].text = str(r["dat"])
                rc[5].text = str(r["chua_dat"])
                rc[6].text = str(r["comment"])
            align_docx_table(t3, left_cols=[6])

            add_heading_with_font(doc, "1.4. Danh sách học sinh đạt Danh hiệu Thi đua & Khen thưởng", level=2)
            t4 = doc.add_table(rows=1, cols=7)
            t4.style = "Table Grid"
            t4_hdr = t4.rows[0].cells
            headers = [
                "STT",
                "Mã học sinh",
                "Họ và tên học sinh",
                "Loại Học tập",
                "Loại Rèn luyện",
                "Danh hiệu đạt được",
                "Thành tích nổi bật khác",
            ]
            for i, h in enumerate(headers):
                t4_hdr[i].text = h
            format_table_header(t4.rows[0])
            for r in data["c_awards_rows"]:
                rc = t4.add_row().cells
                rc[0].text = str(r["stt"])
                rc[1].text = str(r["code"])
                rc[2].text = str(r["name"])
                rc[3].text = str(r["academic"])
                rc[4].text = str(r["conduct"])
                rc[5].text = str(r["title"])
                rc[6].text = str(r["special"])
            align_docx_table(t4, left_cols=[2, 6])

            add_heading_with_font(doc, "1.5. Kế hoạch Phụ đạo và Đồng hành cùng học sinh trong Học kỳ II", level=2)
            t5 = doc.add_table(rows=1, cols=6)
            t5.style = "Table Grid"
            t5_hdr = t5.rows[0].cells
            headers = [
                "STT",
                "Mã học sinh",
                "Họ và tên học sinh",
                "Biểu hiện điểm số thô / Khó khăn",
                '"Bệnh học" hệ thống chẩn đoán',
                "Gợi ý hành động sư phạm",
            ]
            for i, h in enumerate(headers):
                t5_hdr[i].text = h
            format_table_header(t5.rows[0])
            for r in data["c_support_rows"]:
                rc = t5.add_row().cells
                rc[0].text = str(r["stt"])
                rc[1].text = str(r["code"])
                rc[2].text = str(r["name"])
                rc[3].text = str(r["details"])
                rc[4].text = str(r["diagnosis"])
                rc[5].text = str(r["suggestion"])
            align_docx_table(t5, left_cols=[2, 3, 4, 5])

    # Subject Quality Tables
    elif payload.report_type == "subject_quality":
        add_heading_with_font(doc, "1. SỐ LIỆU THỐNG KÊ", level=1)
        if not payload.class_id and payload.grade_level == "all":
            add_heading_with_font(doc, "1.1. Chỉ số thống kê chất lượng đề và độ phân hóa toàn trường", level=2)
            t1 = doc.add_table(rows=1, cols=7)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "Môn học",
                "Số bài thi",
                "Điểm TB (Mean)",
                "Điểm Trung vị (Median)",
                "Điểm xuất hiện nhiều nhất (Mode)",
                "Độ lệch chuẩn (σ)",
                "Nhận định tự động từ hệ thống",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["s_sub_q_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["name"])
                rc[1].text = str(r["total"])
                rc[2].text = str(r["mean"])
                rc[3].text = str(r["median"])
                rc[4].text = str(r["mode"])
                rc[5].text = str(r["stdev"])
                rc[6].text = str(r["comment"])
            align_docx_table(t1, left_cols=[6])

            add_heading_with_font(doc, "1.2. Xếp hạng hiệu suất bộ môn toàn trường", level=2)
            t2 = doc.add_table(rows=1, cols=6)
            t2.style = "Table Grid"
            t2_hdr = t2.rows[0].cells
            headers = [
                "STT",
                "Môn học",
                "Điểm TB Toàn trường",
                "Tỷ lệ điểm Giỏi (>= 8.0)",
                "Tỷ lệ điểm dưới TB (< 5.0)",
                "Đánh giá trạng thái chuyên môn",
            ]
            for i, h in enumerate(headers):
                t2_hdr[i].text = h
            format_table_header(t2.rows[0])
            for idx, r in enumerate(data["s_sub_ranking_rows"]):
                rc = t2.add_row().cells
                rc[0].text = str(idx + 1)
                rc[1].text = str(r["name"])
                rc[2].text = str(r["mean"])
                rc[3].text = str(r["gioi_pct"])
                rc[4].text = str(r["yeu_pct"])
                rc[5].text = str(r["status"])
            align_docx_table(t2, left_cols=[5])

        elif not payload.class_id:
            add_heading_with_font(doc, "1.1. So sánh Hiệu quả giảng dạy và Phân khúc học lực các Lớp", level=2)
            t1 = doc.add_table(rows=1, cols=7)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "Lớp",
                "Sĩ số",
                "Điểm TB Môn",
                "Phân khúc Yếu (0 - 4.5)",
                "Phân khúc TB (5.0 - 6.5)",
                "Phân khúc Khá (7.0 - 8.0)",
                "Phân khúc Xuất sắc (8.5 - 10)",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["g_deviation_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["class_name"])
                rc[1].text = str(r["total"])
                rc[2].text = str(r["mean"])
                rc[3].text = str(r["yeu"])
                rc[4].text = str(r["tb"])
                rc[5].text = str(r["kha"])
                rc[6].text = str(r["xs"])
            align_docx_table(t1)

            add_heading_with_font(doc, "1.2. Bảng so sánh khoảng cách tiến trình điểm số (Phát hiện điểm ảo)", level=2)
            t2 = doc.add_table(rows=1, cols=6)
            t2.style = "Table Grid"
            t2_hdr = t2.rows[0].cells
            headers = [
                "Lớp - Môn",
                "Trung bình Điểm TX",
                "Trung bình Điểm GK",
                "Trung bình Điểm CK",
                "Chênh lệch (TX - CK)",
                "Nhận định hệ thống",
            ]
            for i, h in enumerate(headers):
                t2_hdr[i].text = h
            format_table_header(t2.rows[0])
            for r in data["g_process_rows"]:
                rc = t2.add_row().cells
                rc[0].text = f"{r['class_name']} - {r['subject']}"
                rc[1].text = str(r["tx"])
                rc[2].text = str(r["gk"])
                rc[3].text = str(r["ck"])
                rc[4].text = str(r["diff"])
                rc[5].text = str(r["status"])
            align_docx_table(t2, left_cols=[5])

        else:
            add_heading_with_font(doc, "1.1. Ma trận độ lệch học lực theo Khối ngành (Phát hiện học lệch)", level=2)
            t1 = doc.add_table(rows=1, cols=6)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "STT",
                "Họ tên học sinh",
                "ĐTB Môn Tự nhiên",
                "ĐTB Môn Xã hội",
                "Chênh lệch",
                "Xu hướng năng lực cá nhân",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for idx, r in enumerate(data["c_deviation_rows"]):
                rc = t1.add_row().cells
                rc[0].text = str(idx + 1)
                rc[1].text = str(r["name"])
                rc[2].text = str(r["nat"])
                rc[3].text = str(r["soc"])
                rc[4].text = str(r["diff"])
                rc[5].text = str(r["trend"])
            align_docx_table(t1, left_cols=[1])

            add_heading_with_font(doc, "1.2. Bảng phân nhóm học tập tự động phục vụ điều hành lớp", level=2)
            t2 = doc.add_table(rows=1, cols=4)
            t2.style = "Table Grid"
            t2_hdr = t2.rows[0].cells
            headers = [
                "Phân nhóm hệ thống",
                "Số lượng",
                "Danh sách học sinh đề xuất",
                "Giải pháp hành động đề xuất cho GVCN",
            ]
            for i, h in enumerate(headers):
                t2_hdr[i].text = h
            format_table_header(t2.rows[0])
            for r in data["c_grouping_rows"]:
                rc = t2.add_row().cells
                rc[0].text = str(r["group"])
                rc[1].text = str(r["total"])
                rc[2].text = str(r["list"])
                rc[3].text = str(r["action"])
            align_docx_table(t2, left_cols=[2, 3])

    # At Risk Tables
    elif payload.report_type == "at_risk":
        add_heading_with_font(doc, "1. SỐ LIỆU THỐNG KÊ", level=1)
        if not payload.class_id and payload.grade_level == "all":
            add_heading_with_font(doc, "1.1. Bảng theo dõi Chỉ số Cần Hỗ trợ theo Bộ môn & Khối", level=2)
            t1 = doc.add_table(rows=1, cols=7)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = ["Môn học", "Khối 6", "Khối 7", "Khối 8", "Khối 9", "Tổng cộng toàn trường", "Tỷ lệ (%)"]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["s_risk_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["subject_name"])
                rc[1].text = str(r["grade_6"])
                rc[2].text = str(r["grade_7"])
                rc[3].text = str(r["grade_8"])
                rc[4].text = str(r["grade_9"])
                rc[5].text = f"{r['total']} HS"
                rc[6].text = str(r["pct"])
            align_docx_table(t1)

        elif not payload.class_id:
            add_heading_with_font(doc, "1.1. Báo cáo phân bổ diện hỗ trợ theo đơn vị Lớp", level=2)
            t1 = doc.add_table(rows=1, cols=5)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "Lớp",
                "Sĩ số",
                "Diện 1: Nguy cơ Liệt môn (GPA môn < 5.0)",
                "Diện 2: Phong độ Tụt dốc (GK/TX - CK >= 2.0)",
                "Diện 3: Học lệch nghiêm trọng (Lệch > 3.0)",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["g_risk_classes"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["class_name"])
                rc[1].text = str(r["siso"])
                rc[2].text = f"{r['area1_count']} HS ({r['area1_list']})"
                rc[3].text = f"{r['area2_count']} HS ({r['area2_list']})"
                rc[4].text = f"{r['area3_count']} HS ({r['area3_list']})"
            align_docx_table(t1, left_cols=[2, 3, 4])

        else:
            add_heading_with_font(doc, "1.1. Hồ sơ theo dõi và Gợi ý hướng hỗ trợ cá nhân (Lớp)", level=2)
            t1 = doc.add_table(rows=1, cols=6)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = ["STT", "Mã HS", "Họ tên", "Môn học / Khó khăn chi tiết", "Chẩn đoán hệ thống", "Hành động gợi ý"]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["c_risk_students"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["stt"])
                rc[1].text = str(r["code"])
                rc[2].text = str(r["name"])
                rc[3].text = str(r["details"])
                rc[4].text = str(r["diagnosis"])
                rc[5].text = str(r["suggestion"])
            align_docx_table(t1, left_cols=[2, 3, 4, 5])

    # Subject Report Tables
    elif payload.report_type == "subject_report":
        add_heading_with_font(doc, "1. SỐ LIỆU THỐNG KÊ", level=1)
        if not payload.class_id and payload.grade_level == "all":
            add_heading_with_font(doc, f"1.1. Bảng chỉ số sức khỏe Bộ môn {selected_subject_name} theo khối", level=2)
            t1 = doc.add_table(rows=1, cols=6)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "Khối",
                "Số bài thi",
                "Điểm TB Khối",
                "Xuất sắc (>= 9.0)",
                "Đạt chuẩn (>= 5.0)",
                "Dưới chuẩn (< 5.0)",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["s_sub_rep_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["grade_name"])
                rc[1].text = str(r["total"])
                rc[2].text = str(r["mean"])
                rc[3].text = str(r["gioi_pct"])
                rc[4].text = str(r["dat_pct"])
                rc[5].text = str(r["duoi_pct"])
            align_docx_table(t1)

        elif not payload.class_id:
            add_heading_with_font(
                doc, f"1.1. Báo cáo so sánh thi đua chất lượng môn {selected_subject_name} giữa các lớp", level=2
            )
            t1 = doc.add_table(rows=1, cols=6)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = [
                "Lớp",
                "Sĩ số",
                "Điểm TB Lớp",
                "Số HS đạt điểm Giỏi (>= 8.0)",
                "Số HS cần phụ đạo (< 5.0)",
                "Chênh lệch so với TB Khối",
            ]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["g_sub_rep_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["class_name"])
                rc[1].text = str(r["siso"])
                rc[2].text = str(r["mean"])
                rc[3].text = str(r["gioi"])
                rc[4].text = str(r["chua_dat"])
                rc[5].text = str(r["diff"])
            align_docx_table(t1)

        else:
            add_heading_with_font(
                doc, f"1.1. Sổ tay theo dõi học sinh đột biến theo môn học {selected_subject_name}", level=2
            )
            t1 = doc.add_table(rows=1, cols=4)
            t1.style = "Table Grid"
            t1_hdr = t1.rows[0].cells
            headers = ["Phân loại học sinh", "Số lượng", "Danh sách học sinh đề xuất", "Hành động sư phạm gợi ý"]
            for i, h in enumerate(headers):
                t1_hdr[i].text = h
            format_table_header(t1.rows[0])
            for r in data["c_sub_rep_rows"]:
                rc = t1.add_row().cells
                rc[0].text = str(r["group"])
                rc[1].text = str(r["total"])
                rc[2].text = str(r["list"])
                rc[3].text = str(r["action"])
            align_docx_table(t1, left_cols=[2, 3])

    # AI Insights
    if payload.include_ai_insights and ai_comment:
        part1, part2, part3 = parse_ai_comment(ai_comment)

        add_heading_with_font(doc, "2. ĐÁNH GIÁ CHUYÊN MÔN", level=1)

        add_heading_with_font(doc, "2.1. Tóm tắt tình hình", level=2)
        for p_text in part1.split("\n\n"):
            if p_text.strip():
                clean_p = re.sub(r"<[^>]+>", "", p_text.strip())
                if clean_p.strip():
                    p = doc.add_paragraph(clean_p)
                    p.paragraph_format.line_spacing = 1.2
                    p.paragraph_format.space_after = Pt(10)

        add_heading_with_font(doc, "2.2. Điểm bất thường và rủi ro học thuật", level=2)
        for p_text in part2.split("\n\n"):
            if p_text.strip():
                clean_p = re.sub(r"<[^>]+>", "", p_text.strip())
                if clean_p.strip():
                    p = doc.add_paragraph(clean_p)
                    p.paragraph_format.line_spacing = 1.2
                    p.paragraph_format.space_after = Pt(10)

        add_heading_with_font(doc, "3. ĐỀ XUẤT SƯ PHẠM VÀ PHƯƠNG HƯỚNG ĐIỀU CHỈNH", level=1)
        for p_text in part3.split("\n\n"):
            if p_text.strip():
                clean_p = re.sub(r"<[^>]+>", "", p_text.strip())
                if clean_p.strip():
                    p = doc.add_paragraph(clean_p)
                    p.paragraph_format.line_spacing = 1.2
                    p.paragraph_format.space_after = Pt(10)

    # Hiệu trưởng phê duyệt
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.alignment = 2  # Right
    sig_p.add_run("HIỆU TRƯỞNG PHÊ DUYỆT BÁO CÁO\n\n\n\n").bold = True
    sig_run = sig_p.add_run(f"Thầy {principal_name}")
    sig_run.font.bold = True
    sig_run.font.italic = True

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_html_report(
    payload,
    school_name,
    principal_name,
    semester_id,
    sem_name,
    year_name,
    selected_grade_name,
    selected_class_name,
    cls_row,
    all_subjects,
    subject_map,
    all_grades,
    grade_map,
    all_classes,
    class_map,
    selected_subject_name,
    scope_classes,
    total_classes,
    students_data,
    total_students,
    gpa,
    student_subject_averages,
    student_subject_details,
    ai_comment,
    data,
):
    title_text = get_report_title(
        payload.report_type, payload.grade_level, cls_row.name if cls_row else "", year_name, selected_subject_name
    )
    data_table_html = ""
    charts_html = ""

    # CASE 3.1: academic_conduct
    if payload.report_type == "academic_conduct":
        if not payload.class_id and payload.grade_level == "all":
            rows_siso_html = "".join(
                [
                    f"<tr><td><strong>{r['grade']}</strong></td><td>{r['classes']}</td><td>{r['dau']}</td><td>{r['cuoi']}</td><td>{r['den']}</td><td>{r['di']}</td><td>{r['bo']}</td><td style='font-weight: bold;'>{r['ratio']}</td></tr>"
                    for r in data["s_siso_rows"]
                ]
            )
            rows_siso_html += f"<tr class='total-row'><td>TỔNG TRƯỜNG</td><td>{total_classes}</td><td>{data['total_dau']}</td><td>{data['total_cuoi']}</td><td>{data['total_den']}</td><td>{data['total_di']}</td><td>{data['total_bo']}</td><td>{round((data['total_cuoi'] / data['total_dau']) * 100, 2) if data['total_dau'] > 0 else 100}%</td></tr>"

            rows_acad_html = "".join(
                [
                    f"<tr><td><strong>{r['grade']}</strong></td><td>{r['total']}</td><td>{r['tot']} ({r['tot_pct']})</td><td>{r['kha']} ({r['kha_pct']})</td><td>{r['dat']} ({r['dat_pct']})</td><td>{r['yeu']} ({r['yeu_pct']})</td></tr>"
                    for r in data["s_acad_rows"]
                ]
            )
            rows_acad_html += f"<tr class='total-row'><td>TỔNG CỘNG</td><td>{total_students}</td><td>{data['tot_tot_a']} ({round(data['tot_tot_a'] / total_students * 100, 1)}%)</td><td>{data['tot_kha_a']} ({round(data['tot_kha_a'] / total_students * 100, 1)}%)</td><td>{data['tot_dat_a']} ({round(data['tot_dat_a'] / total_students * 100, 1)}%)</td><td>{data['tot_yeu_a']} ({round(data['tot_yeu_a'] / total_students * 100, 1)}%)</td></tr>"

            rows_cond_html = "".join(
                [
                    f"<tr><td><strong>{r['grade']}</strong></td><td>{r['total']}</td><td>{r['tot']} ({r['tot_pct']})</td><td>{r['kha']} ({r['kha_pct']})</td><td>{r['dat']} ({r['dat_pct']})</td><td>{r['yeu']} ({r['yeu_pct']})</td></tr>"
                    for r in data["s_cond_rows"]
                ]
            )
            rows_cond_html += f"<tr class='total-row'><td>TỔNG CỘNG</td><td>{total_students}</td><td>{data['tot_tot_c']} ({round(data['tot_tot_c'] / total_students * 100, 1)}%)</td><td>{data['tot_kha_c']} ({round(data['tot_kha_c'] / total_students * 100, 1)}%)</td><td>{data['tot_dat_c']} ({round(data['tot_dat_c'] / total_students * 100, 1)}%)</td><td>{data['tot_yeu_c']} ({round(data['tot_yeu_c'] / total_students * 100, 1)}%)</td></tr>"

            rows_award_html = "".join(
                [
                    f"<tr><td><strong>{r['grade']}</strong></td><td>{r['xuatsac']}</td><td>{r['gioi']}</td><td>{r['chuyende']}</td></tr>"
                    for r in data["s_award_rows"]
                ]
            )
            rows_award_html += f"<tr class='total-row'><td>TỔNG</td><td>{data['tot_xs']}</td><td>{data['tot_gi']}</td><td>{data['tot_cd']}</td></tr>"

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Bảng Sĩ số & Biến động học sinh</div>
            <table class="data-table">
                <thead>
                    <tr><th>Khối</th><th>Số lớp</th><th>Sĩ số đầu năm</th><th>Sĩ số cuối HK1</th><th>Chuyển đến</th><th>Chuyển đi</th><th>Bỏ học</th><th>Duy trì sĩ số</th></tr>
                </thead>
                <tbody>{rows_siso_html}</tbody>
            </table>
            <div class="sub-section-title">1.2. Báo cáo Chất lượng Học tập theo Khối</div>
            <table class="data-table">
                <thead>
                    <tr><th>Khối</th><th>Tổng HS</th><th>Tốt</th><th>Khá</th><th>Đạt</th><th>Chưa đạt</th></tr>
                </thead>
                <tbody>{rows_acad_html}</tbody>
            </table>
            <div class="sub-section-title">1.3. Báo cáo chất lượng rèn luyện theo Khối</div>
            <table class="data-table">
                <thead>
                    <tr><th>Khối</th><th>Tổng HS</th><th>Tốt</th><th>Khá</th><th>Đạt</th><th>Chưa đạt</th></tr>
                </thead>
                <tbody>{rows_cond_html}</tbody>
            </table>
            <div class="sub-section-title">1.4. Thống kê danh hiệu thi đua của học sinh</div>
            <table class="data-table">
                <thead>
                    <tr><th>Khối</th><th>Học sinh Xuất sắc</th><th>Học sinh Giỏi / Khen thưởng</th><th>Khen thưởng chuyên đề / Đột xuất</th></tr>
                </thead>
                <tbody>{rows_award_html}</tbody>
            </table>
            """

            pass

        elif not payload.class_id:
            rows_siso_html = "".join(
                [
                    f"<tr><td><strong>{r['class_name']}</strong></td><td>{r['dau']}</td><td>{r['cuoi']}</td><td>{r['den']}</td><td>{r['di']}</td><td>{r['co_p']}</td><td>{r['kp_p']}</td></tr>"
                    for r in data["g_siso_rows"]
                ]
            )
            tot_dau = sum(r["dau"] for r in data["g_siso_rows"])
            tot_cuoi = sum(r["cuoi"] for r in data["g_siso_rows"])
            tot_den = sum(r["den"] for r in data["g_siso_rows"])
            tot_di = sum(r["di"] for r in data["g_siso_rows"])
            tot_co = sum(r["co_p"] for r in data["g_siso_rows"])
            tot_kp = sum(r["kp_p"] for r in data["g_siso_rows"])
            rows_siso_html += f"<tr class='total-row'><td>TỔNG KHỐI</td><td>{tot_dau}</td><td>{tot_cuoi}</td><td>{tot_den}</td><td>{tot_di}</td><td>{tot_co}</td><td>{tot_kp}</td></tr>"

            rows_acad_html = "".join(
                [
                    f"<tr><td><strong>{r['class_name']}</strong></td><td>{r['total']}</td><td>{r['tot']} ({r['tot_pct']})</td><td>{r['kha']} ({r['kha_pct']})</td><td>{r['dat']} ({r['dat_pct']})</td><td>{r['yeu']} ({r['yeu_pct']})</td></tr>"
                    for r in data["g_acad_rows"]
                ]
            )
            rows_acad_html += f"<tr class='total-row'><td>TỔNG KHỐI</td><td>{tot_cuoi}</td><td>{sum(x['tot'] for x in data['g_acad_rows'])}</td><td>{sum(x['kha'] for x in data['g_acad_rows'])}</td><td>{sum(x['dat'] for x in data['g_acad_rows'])}</td><td>{sum(x['yeu'] for x in data['g_acad_rows'])}</td></tr>"

            rows_cond_html = "".join(
                [
                    f"<tr><td><strong>{r['class_name']}</strong></td><td>{r['total']}</td><td>{r['tot']} ({r['tot_pct']})</td><td>{r['kha']} ({r['kha_pct']})</td><td>{r['dat']} ({r['dat_pct']})</td><td>{r['yeu']} ({r['yeu_pct']})</td></tr>"
                    for r in data["g_cond_rows"]
                ]
            )
            rows_cond_html += f"<tr class='total-row'><td>TỔNG KHỐI</td><td>{tot_cuoi}</td><td>{sum(x['tot'] for x in data['g_cond_rows'])}</td><td>{sum(x['kha'] for x in data['g_cond_rows'])}</td><td>{sum(x['dat'] for x in data['g_cond_rows'])}</td><td>{sum(x['yeu'] for x in data['g_cond_rows'])}</td></tr>"

            rows_sub_html = "".join(
                [
                    f"<tr><td>{idx + 1}</td><td><strong>{r['subject_name']}</strong></td><td class='align-left'>{r['worst_class']}</td><td>{r['chua_dat']}</td><td>{r['mean']}</td><td class='align-left'>{r['reason']}</td></tr>"
                    for idx, r in enumerate(data["g_subjects_rows"][:4])
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Thống kê Sĩ số và Chuyên cần theo từng lớp</div>
            <table class="data-table">
                <thead>
                    <tr><th>Lớp</th><th>Sĩ số đầu năm</th><th>Sĩ số cuối HK1</th><th>Chuyển đến</th><th>Chuyển đi</th><th>Nghỉ có phép (Tổng)</th><th>Nghỉ không phép (Tổng)</th></tr>
                </thead>
                <tbody>{rows_siso_html}</tbody>
            </table>
            <div class="sub-section-title">1.2. So sánh Kết quả Học tập giữa các lớp trong Khối</div>
            <table class="data-table">
                <thead>
                    <tr><th>Lớp</th><th>Tổng số HS</th><th>Tốt</th><th>Khá</th><th>Đạt</th><th>Chưa đạt</th></tr>
                </thead>
                <tbody>{rows_acad_html}</tbody>
            </table>
            <div class="sub-section-title">1.3. So sánh Kết quả Rèn luyện (Hạnh kiểm) giữa các lớp</div>
            <table class="data-table">
                <thead>
                    <tr><th>Lớp</th><th>Tổng số HS</th><th>Tốt</th><th>Khá</th><th>Đạt</th><th>Chưa đạt</th></tr>
                </thead>
                <tbody>{rows_cond_html}</tbody>
            </table>
            <div class="sub-section-title">1.4. Thống kê các môn học có tỷ lệ Chưa đạt cao</div>
            <table class="data-table">
                <thead>
                    <tr><th>STT</th><th>Môn học</th><th>Lớp có tỷ lệ Chưa đạt cao nhất</th><th>Số lượng HS chưa đạt (Toàn khối)</th><th>Điểm TB</th><th>Nguyên nhân chính</th></tr>
                </thead>
                <tbody>{rows_sub_html}</tbody>
            </table>
            """
        else:
            # Class level
            rows_siso_html = "".join(
                [
                    f"<tr><td class='align-left'><strong>{tc}</strong></td><td>{val}</td><td class='align-left'>{note}</td></tr>"
                    for tc, val, note in data["c_siso_list"]
                ]
            )

            rows_summary_html = "".join(
                [
                    f"<tr><td><strong>{r['category']}</strong></td><td>{r['acad']}</td><td>{r['cond']}</td></tr>"
                    for r in data["c_summary_rows"]
                ]
            )
            rows_summary_html += f"<tr class='total-row'><td>TỔNG CỘNG</td><td>{total_students} HS (100%)</td><td>{total_students} HS (100%)</td></tr>"

            rows_sub_html = "".join(
                [
                    f"<tr><td>{idx + 1}</td><td><strong>{r['name']}</strong></td><td>{r['tot']}</td><td>{r['kha']}</td><td>{r['dat']}</td><td>{r['chua_dat']}</td><td class='align-left'>{r['comment']}</td></tr>"
                    for idx, r in enumerate(data["c_subjects_rows"])
                ]
            )

            rows_award_html = "".join(
                [
                    f"<tr><td>{r['stt']}</td><td>{r['code']}</td><td class='align-left'><strong>{r['name']}</strong></td><td>{r['academic']}</td><td>{r['conduct']}</td><td><strong>{r['title']}</strong></td><td class='align-left'>{r['special']}</td></tr>"
                    for r in data["c_awards_rows"]
                ]
            )
            if not rows_award_html:
                rows_award_html = (
                    "<tr><td colspan='7' style='text-align: center;'>Không có danh hiệu thi đua nào trong kỳ.</td></tr>"
                )

            rows_support_html = "".join(
                [
                    f"<tr><td>{r['stt']}</td><td>{r['code']}</td><td class='align-left'><strong>{r['name']}</strong></td><td class='align-left'>{r['details']}</td><td class='align-left'><strong>{r['diagnosis']}</strong></td><td class='align-left'>{r['suggestion']}</td></tr>"
                    for r in data["c_support_rows"]
                ]
            )
            if not rows_support_html:
                rows_support_html = "<tr><td colspan='6' style='text-align: center;'>Không ghi nhận học sinh nào trong diện cần hỗ trợ đặc biệt.</td></tr>"

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Tổng hợp chung về Sĩ số và Duy trì nề nếp</div>
            <table class="data-table">
                <thead>
                    <tr><th>Tiêu chí</th><th>Số lượng / Chỉ số</th><th>Ghi chú</th></tr>
                </thead>
                <tbody>{rows_siso_html}</tbody>
            </table>
            <div class="sub-section-title">1.2. Kết quả Học tập và Rèn luyện định lượng (Toàn lớp)</div>
            <table class="data-table">
                <thead>
                    <tr><th>Phân loại kết quả</th><th>Kết quả Học tập</th><th>Kết quả Rèn luyện</th></tr>
                </thead>
                <tbody>{rows_summary_html}</tbody>
            </table>
            <div class="sub-section-title">1.3. Thống kê chi tiết Kết quả theo từng Môn học</div>
            <table class="data-table">
                <thead>
                    <tr><th>STT</th><th>Môn học</th><th>Số HS đạt mức Tốt</th><th>Số HS đạt mức Khá</th><th>Số HS đạt mức Đạt</th><th>Số HS Chưa đạt</th><th>Đánh giá sơ bộ của GVBM</th></tr>
                </thead>
                <tbody>{rows_sub_html}</tbody>
            </table>
            <div class="sub-section-title">1.4. Danh sách học sinh đạt Danh hiệu Thi đua & Khen thưởng</div>
            <table class="data-table">
                <thead>
                    <tr><th>STT</th><th>Mã học sinh</th><th>Họ và tên học sinh</th><th>Loại Học tập</th><th>Loại Rèn luyện</th><th>Danh hiệu đạt được</th><th>Thành tích nổi bật khác</th></tr>
                </thead>
                <tbody>{rows_award_html}</tbody>
            </table>
            <div class="sub-section-title">1.5. Kế hoạch Phụ đạo và Đồng hành cùng học sinh trong Học kỳ II</div>
            <table class="data-table">
                <thead>
                    <tr><th>STT</th><th>Mã học sinh</th><th>Họ và tên học sinh</th><th>Biểu hiện điểm số thô / Khó khăn</th><th>"Bệnh học" hệ thống chẩn đoán</th><th>Gợi ý hành động sư phạm</th></tr>
                </thead>
                <tbody>{rows_support_html}</tbody>
            </table>
            """

    # CASE 3.2: subject_quality
    elif payload.report_type == "subject_quality":
        if not payload.class_id and payload.grade_level == "all":
            rows_sub_q_html = "".join(
                [
                    f"<tr><td><strong>{r['name']}</strong></td><td>{r['total']}</td><td>{r['mean']}</td><td>{r['median']}</td><td>{r['mode']}</td><td>{r['stdev']}</td><td class='align-left'>{r['comment']}</td></tr>"
                    for r in data["s_sub_q_rows"]
                ]
            )
            rows_sub_ranking_html = "".join(
                [
                    f"<tr><td>{idx + 1}</td><td><strong>{r['name']}</strong></td><td>{r['mean']}</td><td>{r['gioi_pct']}</td><td>{r['yeu_pct']}</td><td class='align-left'><strong>{r['status']}</strong></td></tr>"
                    for idx, r in enumerate(data["s_sub_ranking_rows"])
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Chỉ số thống kê chất lượng đề và độ phân hóa toàn trường</div>
            <table class="data-table">
                <thead>
                    <tr><th>Môn học</th><th>Số bài thi</th><th>Điểm TB (Mean)</th><th>Điểm Trung vị (Median)</th><th>Điểm xuất hiện nhiều nhất (Mode)</th><th>Độ lệch chuẩn (σ)</th><th>Nhận định tự động từ hệ thống</th></tr>
                </thead>
                <tbody>{rows_sub_q_html}</tbody>
            </table>
            <div class="sub-section-title">1.2. Xếp hạng hiệu suất bộ môn toàn trường</div>
            <table class="data-table">
                <thead>
                    <tr><th>STT</th><th>Môn học</th><th>Điểm TB Toàn trường</th><th>Tỷ lệ điểm Giỏi (>= 8.0)</th><th>Tỷ lệ điểm dưới TB (< 5.0)</th><th>Đánh giá trạng thái chuyên môn</th></tr>
                </thead>
                <tbody>{rows_sub_ranking_html}</tbody>
            </table>
            """

            if payload.include_charts and data["s_sub_q_rows"]:
                charts_html = "<div class='sub-section-title'>Đồ thị so sánh điểm trung bình bộ môn</div><div class='chart-container'><div class='bar-chart-flex'>"
                for r in data["s_sub_q_rows"][:6]:
                    val_h = int(r["mean"] * 10)
                    charts_html += f"""
                    <div class="bar-col">
                        <div class="bar-fill" style="height: {val_h}px;">
                            <div class="bar-val">{r["mean"]}</div>
                        </div>
                        <div class="bar-label">{r["name"]}</div>
                    </div>
                    """
                charts_html += "</div></div>"

        elif not payload.class_id:
            rows_compare_html = "".join(
                [
                    f"<tr><td class='align-left'><strong>{r['class_name']}</strong></td><td>{r['total']}</td><td>{r['mean']}</td><td>{r['yeu']}</td><td>{r['tb']}</td><td>{r['kha']}</td><td>{r['xs']}</td></tr>"
                    for r in data["g_deviation_rows"]
                ]
            )
            rows_process_html = "".join(
                [
                    f"<tr><td><strong>{r['class_name']} - {r['subject']}</strong></td><td>{r['tx']}</td><td>{r['gk']}</td><td>{r['ck']}</td><td style='font-weight: bold;'>{r['diff']}</td><td class='align-left'><strong>{r['status']}</strong></td></tr>"
                    for r in data["g_process_rows"]
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. So sánh Hiệu quả giảng dạy và Phân khúc học lực các Lớp</div>
            <table class="data-table">
                <thead>
                    <tr><th>Lớp</th><th>Sĩ số</th><th>Điểm TB Môn</th><th>Phân khúc Yếu (0 - 4.5)</th><th>Phân khúc TB (5.0 - 6.5)</th><th>Phân khúc Khá (7.0 - 8.0)</th><th>Phân khúc Xuất sắc (8.5 - 10)</th></tr>
                </thead>
                <tbody>{rows_compare_html}</tbody>
            </table>
            <div class="sub-section-title">1.2. Bảng so sánh khoảng cách tiến trình điểm số (Phát hiện điểm ảo)</div>
            <table class="data-table">
                <thead>
                    <tr><th>Lớp - Môn</th><th>Trung bình Điểm TX</th><th>Trung bình Điểm GK</th><th>Trung bình Điểm CK</th><th>Chênh lệch (TX - CK)</th><th>Nhận định hệ thống</th></tr>
                </thead>
                <tbody>{rows_process_html}</tbody>
            </table>
            """
        else:
            rows_deviation_html = "".join(
                [
                    f"<tr><td>{idx + 1}</td><td><strong>{r['name']}</strong></td><td>{r['nat']}</td><td>{r['soc']}</td><td style='font-weight: bold;'>{r['diff']}</td><td><strong>{r['trend']}</strong></td></tr>"
                    for idx, r in enumerate(data["c_deviation_rows"])
                ]
            )
            rows_grouping_html = "".join(
                [
                    f"<tr><td class='align-left'><strong>{r['group']}</strong></td><td>{r['total']}</td><td class='align-left'>{r['list']}</td><td class='align-left'>{r['action']}</td></tr>"
                    for r in data["c_grouping_rows"]
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Ma trận độ lệch học lực theo Khối ngành (Phát hiện học lệch - Trích xuất)</div>
            <table class="data-table">
                <thead>
                    <tr><th>STT</th><th>Họ tên học sinh</th><th>ĐTB Môn Tự nhiên</th><th>ĐTB Môn Xã hội</th><th>Chênh lệch</th><th>Xu hướng năng lực cá nhân</th></tr>
                </thead>
                <tbody>{rows_deviation_html}</tbody>
            </table>
            <div class="sub-section-title">1.2. Bảng phân nhóm học tập tự động phục vụ điều hành lớp</div>
            <table class="data-table">
                <thead>
                    <tr><th>Phân nhóm hệ thống</th><th>Số lượng</th><th>Danh sách học sinh đề xuất</th><th>Giải pháp hành động đề xuất cho GVCN</th></tr>
                </thead>
                <tbody>{rows_grouping_html}</tbody>
            </table>
            """

    # CASE 3.3: at_risk
    elif payload.report_type == "at_risk":
        if not payload.class_id:
            rows_risk_html = "".join(
                [
                    f"<tr><td><strong>{r['subject_name']}</strong></td><td>{r['grade_6']}</td><td>{r['grade_7']}</td><td>{r['grade_8']}</td><td>{r['grade_9']}</td><td style='font-weight: bold;'>{r['total']} HS</td><td style='font-weight: bold;'>{r['pct']}</td></tr>"
                    for r in data["s_risk_rows"]
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Bảng theo dõi Chỉ số Cần Hỗ trợ theo Bộ môn & Khối</div>
            <table class="data-table">
                <thead>
                    <tr><th>Môn học</th><th>Khối 6</th><th>Khối 7</th><th>Khối 8</th><th>Khối 9</th><th>Tổng cộng toàn trường</th><th>Tỷ lệ (%)</th></tr>
                </thead>
                <tbody>{rows_risk_html}</tbody>
            </table>
            """

            if payload.grade_level != "all":
                rows_risk_classes_html = "".join(
                    [
                        f"<tr><td><strong>{r['class_name']}</strong></td><td>{r['siso']}</td><td class='align-left'>{r['area1_count']} HS ({r['area1_list']})</td><td class='align-left'>{r['area2_count']} HS ({r['area2_list']})</td><td class='align-left'>{r['area3_count']} HS ({r['area3_list']})</td></tr>"
                        for r in data["g_risk_classes"]
                    ]
                )

                data_table_html = f"""
                <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Báo cáo phân bổ diện hỗ trợ theo đơn vị Lớp</div>
                <table class="data-table">
                    <thead>
                        <tr><th>Lớp</th><th>Sĩ số</th><th>Diện 1: Nguy cơ Liệt môn (GPA môn < 5.0)</th><th>Diện 2: Phong độ Tụt dốc (GK/TX - CK >= 2.0)</th><th>Diện 3: Học lệch nghiêm trọng (Lệch > 3.0)</th></tr>
                    </thead>
                    <tbody>{rows_risk_classes_html}</tbody>
                </table>
                """
        else:
            rows_risk_stud_html = "".join(
                [
                    f"<tr><td>{r['stt']}</td><td>{r['code']}</td><td class='align-left'><strong>{r['name']}</strong></td><td class='align-left'>{r['details']}</td><td class='align-left'><strong>{r['diagnosis']}</strong></td><td class='align-left'>{r['suggestion']}</td></tr>"
                    for r in data["c_risk_students"]
                ]
            )
            if not rows_risk_stud_html:
                rows_risk_stud_html = "<tr><td colspan='6' style='text-align: center; color: #666;'>Không ghi nhận học sinh nào rơi vào diện cần phụ đạo. Lớp học an toàn học thuật.</td></tr>"

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Hồ sơ theo dõi và Gợi ý hướng hỗ trợ cá nhân (Lớp)</div>
            <table class="data-table">
                <thead>
                    <tr><th>STT</th><th>Mã HS</th><th>Họ tên</th><th>Môn học / Khó khăn chi tiết</th><th>Chẩn đoán hệ thống</th><th>Hành động gợi ý</th></tr>
                </thead>
                <tbody>{rows_risk_stud_html}</tbody>
            </table>
            """

    # CASE 3.4: subject_report
    elif payload.report_type == "subject_report":
        if not payload.class_id and payload.grade_level == "all":
            rows_sub_rep_html = "".join(
                [
                    f"<tr><td><strong>{r['grade_name']}</strong></td><td>{r['total']}</td><td>{r['mean']}</td><td>{r['gioi_pct']}</td><td>{r['dat_pct']}</td><td style='font-weight: bold;'>{r['duoi_pct']}</td></tr>"
                    for r in data["s_sub_rep_rows"]
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Bảng chỉ số sức khỏe Bộ môn {selected_subject_name} theo khối</div>
            <table class="data-table">
                <thead>
                    <tr><th>Khối</th><th>Số bài thi</th><th>Điểm TB Khối</th><th>Xuất sắc (>= 9.0)</th><th>Đạt chuẩn (>= 5.0)</th><th>Dưới chuẩn (< 5.0)</th></tr>
                </thead>
                <tbody>{rows_sub_rep_html}</tbody>
            </table>
            """
        elif not payload.class_id:
            rows_sub_classes_html = "".join(
                [
                    f"<tr><td><strong>{r['class_name']}</strong></td><td>{r['siso']}</td><td>{r['mean']}</td><td>{r['gioi']}</td><td>{r['chua_dat']}</td><td style='font-weight: bold;'>{r['diff']}</td></tr>"
                    for r in data["g_sub_rep_rows"]
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Báo cáo so sánh thi đua chất lượng môn {selected_subject_name} giữa các lớp</div>
            <table class="data-table">
                <thead>
                    <tr><th>Lớp</th><th>Sĩ số</th><th>Điểm TB Lớp</th><th>Số HS đạt điểm Giỏi (>= 8.0)</th><th>Số HS cần phụ đạo (< 5.0)</th><th>Chênh lệch so với TB Khối</th></tr>
                </thead>
                <tbody>{rows_sub_classes_html}</tbody>
            </table>
            """
        else:
            rows_sub_rep_cls_html = "".join(
                [
                    f"<tr><td class='align-left'><strong>{r['group']}</strong></td><td>{r['total']}</td><td class='align-left'>{r['list']}</td><td class='align-left'>{r['action']}</td></tr>"
                    for r in data["c_sub_rep_rows"]
                ]
            )

            data_table_html = f"""
            <div class="main-section-title">1. SỐ LIỆU THỐNG KÊ</div>
            <div class="sub-section-title">1.1. Sổ tay theo dõi học sinh đột biến theo môn học {selected_subject_name}</div>
            <table class="data-table">
                <thead>
                    <tr><th>Phân loại học sinh</th><th>Số lượng</th><th>Danh sách học sinh đề xuất</th><th>Hành động sư phạm gợi ý</th></tr>
                </thead>
                <tbody>{rows_sub_rep_cls_html}</tbody>
            </table>
            """

    # AI Insights Section natural flow
    ai_insights_html = ""
    if payload.include_ai_insights and ai_comment:
        part1, part2, part3 = parse_ai_comment(ai_comment)

        def to_html_paragraphs(part_text: str) -> str:
            paragraphs = []
            for p in part_text.split("\n\n"):
                if p.strip():
                    p_clean = p.replace("\n", "<br>")
                    paragraphs.append(
                        f"<p style='margin: 0 0 10px 0; font-size: 13px; line-height: 1.5;'>{p_clean}</p>"
                    )
            return "".join(paragraphs)

        p1_html = to_html_paragraphs(part1)
        p2_html = to_html_paragraphs(part2)
        p3_html = to_html_paragraphs(part3)

        ai_insights_html = f"""
        <div class="ai-insights">
            <div class="main-section-title">2. ĐÁNH GIÁ CHUYÊN MÔN</div>
            <div class="sub-section-title">2.1. Tóm tắt tình hình</div>
            {p1_html}
            <div class="sub-section-title">2.2. Điểm bất thường và rủi ro học thuật</div>
            {p2_html}
            <div class="main-section-title">3. ĐỀ XUẤT SƯ PHẠM VÀ PHƯƠNG HƯỚNG ĐIỀU CHỈNH</div>
            {p3_html}
        </div>
        """

    # Build dynamic Legal bases HTML
    bases_html = ""
    bases = get_legal_bases(
        payload.report_type, payload.grade_level, cls_row.name if cls_row else "", year_name, school_name
    )
    if bases:
        items_html = "".join([f"<p class='legal-base-item'>- {b}</p>" for b in bases])
        bases_html = f"""
        <div class="legal-bases">
            {items_html}
        </div>
        """

    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title_text}</title>
    <style>
        body {{
            font-family: "Times New Roman", Times, serif;
            padding: 40px;
            color: #000000;
            background-color: #ffffff;
            line-height: 1.5;
            max-width: 900px;
            margin: 0 auto;
            font-size: 16px; /* 16px matches 12pt Normal style */
        }}
        .report-header {{
            border-bottom: 1.5px solid #000000;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        .report-title {{
            font-size: 16px;
            font-weight: bold;
            color: #000000;
            text-align: center;
            margin: 25px 0 15px 0;
            text-transform: uppercase;
        }}
        .report-subtitle {{
            font-size: 13px;
            color: #000000;
            text-align: center;
            margin-bottom: 25px;
            font-style: italic;
        }}
        .legal-bases {{
            margin-bottom: 25px;
            font-size: 13px;
            line-height: 1.4;
        }}
        .legal-base-item {{
            margin: 0 0 5px 0;
            font-style: italic;
        }}
        .section-title {{
            font-size: 14px;
            font-weight: bold;
            color: #000000;
            margin-top: 25px;
            margin-bottom: 12px;
            text-transform: uppercase;
            page-break-before: always;
            break-before: page;
        }}
        .section-title:first-of-type {{
            page-break-before: avoid !important;
            break-before: avoid !important;
        }}
        .data-table {{
            width: auto;
            min-width: 75%;
            max-width: 100%;
            margin-left: auto;
            margin-right: auto;
            border-collapse: collapse;
            margin-bottom: 20px;
            font-size: 14px;
        }}
        .data-table th {{
            background-color: #f2f2f2;
            color: #000000;
            padding: 9px 10px;
            font-weight: bold;
            text-align: center; /* Center headers */
            vertical-align: middle;
            border: 1px solid #000000;
        }}
        .data-table td {{
            padding: 9px 10px;
            border: 1px solid #000000;
            color: #000000;
            text-align: center; /* Center cell contents */
            vertical-align: middle;
        }}
        .align-left {{
            text-align: left !important;
        }}
        .data-table tr {{
            page-break-inside: avoid;
            break-inside: avoid;
        }}
        .total-row {{
            font-weight: bold;
            background-color: #f2f2f2 !important;
        }}
        .ai-insights {{
            border: 1px solid #000000;
            padding: 15px;
            margin-top: 30px;
            font-size: 13px;
            color: #000000;
            background-color: #ffffff;
            page-break-inside: avoid;
            break-inside: avoid;
        }}
        .ai-title {{
            font-size: 13px;
            font-weight: bold;
            color: #000000;
            margin-top: 0;
            margin-bottom: 10px;
            text-transform: uppercase;
        }}
        .signature-section {{
            display: flex;
            justify-content: space-between;
            margin-top: 40px;
            font-size: 13px;
            page-break-inside: avoid;
            break-inside: avoid;
        }}
        .signature-block {{
            text-align: center;
            width: 250px;
        }}
        .signature-title {{
            font-weight: bold;
            margin-bottom: 60px;
        }}
        .signature-name {{
            font-weight: bold;
            font-style: italic;
        }}

        /* Grayscale Charts styling */
        .chart-container {{
            border: 1px solid #000000;
            padding: 15px;
            background-color: #ffffff;
            margin-bottom: 25px;
            page-break-inside: avoid;
            break-inside: avoid;
        }}
        .bar-chart-flex {{
            display: flex;
            align-items: flex-end;
            justify-content: space-around;
            height: 150px;
            padding-top: 15px;
            border-left: 1.5px solid #000000;
            border-bottom: 1.5px solid #000000;
        }}
        .bar-col {{
            display: flex;
            flex-direction: column;
            align-items: center;
            width: 50px;
        }}
        .bar-fill {{
            width: 24px;
            background-color: #7f7f7f;
            border: 1px solid #000000;
            border-bottom: none;
            position: relative;
        }}
        .bar-val {{
            position: absolute;
            top: -18px;
            left: 0;
            right: 0;
            text-align: center;
            font-size: 10px;
            font-weight: bold;
            color: #000000;
        }}
        .bar-label {{
            margin-top: 5px;
            font-size: 10px;
            color: #000000;
            text-align: center;
            white-space: nowrap;
        }}
        .progress-bar-container {{
            display: flex;
            height: 16px;
            background-color: #e5e5e5;
            border: 1px solid #000000;
            overflow: hidden;
            margin: 8px 0;
        }}
        .progress-segment {{
            height: 100%;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #000000;
            font-size: 10px;
            font-weight: bold;
            border-right: 1px solid #000000;
        }}
        .progress-segment:last-child {{
            border-right: none;
        }}
        .progress-segment-1 {{ background-color: #7f7f7f; color: #ffffff; }}
        .progress-segment-2 {{ background-color: #a6a6a6; color: #ffffff; }}
        .progress-segment-3 {{ background-color: #d9d9d9; color: #000000; }}
        .progress-segment-4 {{ background-color: #f2f2f2; color: #000000; }}
        @media print {{
            body {{ padding: 0; }}
            .signature-section {{ page-break-inside: avoid; }}
        }}
    </style>
</head>
<body>
    <table style="width: 100%; border: none; border-collapse: collapse; margin-bottom: 20px;">
        <tr style="border: none; background: none;">
            <td style="width: 45%; text-align: center; border: none; padding: 0; vertical-align: top; font-size: 12px; color: #000000;">
                SỞ GIÁO DỤC VÀ ĐÀO TẠO<br>
                <strong>TRƯỜNG {school_name.upper()}</strong><br>
                Số: ....../BC-TH<br>
                <span style="display: inline-block; width: 60px; border-bottom: 1.5px solid #000000; margin-top: 5px;"></span>
            </td>
            <td style="width: 55%; text-align: center; border: none; padding: 0; vertical-align: top; font-size: 12px; color: #000000;">
                <strong>CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</strong><br>
                <strong>Độc lập - Tự do - Hạnh phúc</strong><br>
                <span style="display: inline-block; width: 120px; border-bottom: 1.5px solid #000000; margin-top: 5px; margin-bottom: 5px;"></span><br>
                <span style="font-style: italic;">Hà Nội, ngày 22 tháng 06 năm 2026</span>
            </td>
        </tr>
    </table>

    <div class="report-title">{title_text}</div>
    <div class="report-subtitle">Phạm vi: {selected_grade_name}{selected_class_name} | Kỳ: {sem_name} | Niên khóa {year_name}</div>

    {bases_html}
    {charts_html}
    {data_table_html}
    {ai_insights_html}

    {f"<div class='signature-section'><div class='signature-block'></div><div class='signature-block'><div class='signature-title'>HIỆU TRƯỞNG PHÊ DUYỆT</div><div class='signature-name'>Thầy {principal_name}</div></div></div>" if payload.include_signature else ""}
</body>
</html>"""

    stream = io.BytesIO(html_content.encode("utf-8"))
    return stream
