"""Nạp sách giáo khoa (PDF/DOCX/TXT/MD) → tự tách mục lục → node chương/bài — KHÔNG RAG.

M5 mở rộng: thay vì người dùng tự tổng hợp file JSON/markdown, upload chính cuốn SGK.

- PDF: 2 lượt quét bằng Qwen3-VL-Flash — (A) tìm trang MỤC LỤC (~15 trang đầu) → cây chương→bài
  + danh sách NEO; (B) phân loại trang nội dung CÓ NEO (VLM chỉ chọn ID từ danh sách MỤC LỤC,
  không tự đặt tên → không bịa đơn vị mới từ mục con). Rồi làm giàu từng bài (tóm tắt + từ khóa
  + mục con) bằng TOÀN BỘ trang của bài + tên bài làm neo. Khoảng trang bài xác định theo TÊN
  (không dùng số trang in) → chịu được file cắt ngắn và MỤC LỤC không có số trang.
  Dữ liệu lấy TỪ VLM, không dùng regex/hardcode để bóc nội dung — code chỉ nối/ghép JSON VLM trả về.
- DOCX: heading styles (Heading 1 = chương, Heading 2 = bài).
- TXT/MD: regex dòng "Chương X:" / "Bài n:" (mục lục tay — user kiểm soát nội dung).

Mọi nguồn đều qua chuẩn hóa (_normalize_title), lọc placeholder (_is_placeholder),
gắn cờ phụ (_is_phu_title) và sanity-check (_sanity_check) trước khi preview/lưu.
Hỗ trợ dry_run (xem trước cây dự kiến trước khi ghi DB).
"""

from __future__ import annotations

import json
import re
import uuid
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.orm import Session

from src.config import Settings, get_settings
from src.models.tables import CurriculumUnit
from src.services import layout_detector as ld
from src.services import vlm
from src.services.curriculum_catalog import deactivate_placeholder_units, resolve_subject_ids

_TMP_DIR = Path(__file__).resolve().parents[2] / "temp"

_CHAPTER_RE = re.compile(r"^\s*Chương\s+([IVXLCDM]+|\d+)\s*[.:]?\s*(.+)$", re.IGNORECASE)
_LESSON_RE = re.compile(r"^\s*Bài\s+(\d+)\s*[.:]?\s*(.+)$", re.IGNORECASE)
_SEMESTER_RE = re.compile(r"(?:tập|tap|hk)\s*([12])", re.IGNORECASE)
_TRAILING_PAGE_RE = re.compile(r"[\s.…]+\d{1,3}[\s.…]*$")

# Placeholder xuất hiện trong bản mẫu/template sách — loại hẳn.
_PLACEHOLDER_RE = re.compile(
    r"^(tên chương|tên bài|tên mục|tên hoạt động|tên thực hành|tên luyện tập|tên đề mục|tên phần)$",
    re.IGNORECASE,
)
# Node phụ (giữ trong cây nhưng loại khỏi shortlist map đề thi).
_PHU_TITLE_RE = re.compile(
    r"^(ôn tập|kiểm tra|hoạt động thực hành|luyện tập chung|bài tập cuối|tổng kết chương|câu hỏi ôn tập)",
    re.IGNORECASE,
)

_MAX_CHAPTERS_PER_SEMESTER = 6
_MAX_LESSONS_PER_CHAPTER = 30
# Số trang tối đa đưa vào 1 lần gọi VLM làm giàu 1 bài — gửi TOÀN BỘ trang bài; bài dài hơn
# thì lấy mẫu đều nhưng LUÔN giữ trang đầu + trang cuối (giữ ngữ cảnh đầu/cuối bài).
_MAX_REFINE_PAGES = 10
# Số trang đầu quét để TÌM MỤC LỤC (bước A) trước khi phân loại nội dung có neo (bước B).
_TOC_MAX_PAGES = 15
_MAX_KEYWORDS = 8
# Số mục con tối đa 1 bài (trần an toàn, thực tế thường ≤ 10).
_MAX_SECTIONS = 30
_MAX_SUMMARY_CHARS = 600

TocEntry = tuple[int, str, int]  # (level: 1=chương, 2=bài, page)


def detect_semester_from_filename(filename: str) -> int | None:
    """Đoán học kỳ từ tên file ("tap 1"/"tập 2"/"HK1") → 1 hoặc 2; None nếu không rõ."""
    match = _SEMESTER_RE.search(filename)
    if not match:
        return None
    return int(match.group(1))


def _normalize_title(title: str) -> str:
    """Bỏ số trang cuối dòng + dấu chấm chấm; trim. Giữ nguyên cách viết hoa của sách."""
    return _TRAILING_PAGE_RE.sub("", title).strip(" .…:").strip()


def _is_placeholder(title: str) -> bool:
    """True nếu là placeholder bản mẫu (Tên chương/Tên bài...) → loại hẳn."""
    return bool(_PLACEHOLDER_RE.match(title.strip()))


def _is_phu_title(title: str) -> bool:
    """True nếu là node phụ (Ôn tập chương/Kiểm tra/Hoạt động thực hành...) → gắn cờ is_phu."""
    return bool(_PHU_TITLE_RE.match(title.strip()))


def _parse_json_object(text: str) -> dict[str, Any] | None:
    """Parse JSON object từ text VLM (bóc code fence, lấy object đầu tiên, xử lý LaTeX escapes). None nếu hỏng."""
    cleaned = re.sub(r"^`{3}(?:json)?|`{3}$", "", text.strip(), flags=re.MULTILINE).strip()
    start = cleaned.find("{")
    if start < 0:
        return None
    sub_text = cleaned[start:]
    try:
        obj, _ = json.JSONDecoder().raw_decode(sub_text)
        return obj if isinstance(obj, dict) else None
    except json.JSONDecodeError:
        pass

    # Fallback: Sửa các ký tự escape LaTeX không hợp lệ (ví dụ \in, \cdot, \ge, \{, \})
    try:
        fixed_text = re.sub(r'\\(?![/\\\"bfnrt]|u[0-9a-fA-F]{4})', r'\\\\', sub_text)
        obj, _ = json.JSONDecoder().raw_decode(fixed_text)
        return obj if isinstance(obj, dict) else None
    except json.JSONDecodeError:
        return None


def _as_page_int(value: Any) -> int | None:
    """Chuyển số trang VLM trả về (int/str) → int; None nếu không hợp lệ."""
    if isinstance(value, bool):
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _parse_scan_batch(raw: str | None, expected_count: int) -> list[dict[str, Any]] | None:
    """Parse response 1 lô quét toàn cuốn → list dict theo từng trang; None nếu JSON hỏng hoàn toàn."""
    if not raw:
        return None
    obj = _parse_json_object(raw)
    if obj is None or not isinstance(obj.get("pages"), list):
        return None
    pages = [p if isinstance(p, dict) else {} for p in obj["pages"]]
    if not pages:
        return None
    # Nếu AI trả thiếu/thừa một vài phần tử, tự điều chỉnh vừa khít expected_count thay vì vứt bỏ cả lô
    if len(pages) < expected_count:
        pages.extend([{}] * (expected_count - len(pages)))
    elif len(pages) > expected_count:
        pages = pages[:expected_count]
    return pages


def _merge_toc_chapters(parsed_pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Gom kết quả từng trang VLM → danh sách chương (gộp trang TOC trải nhiều trang).

    Giữ 'page' (số trang in) của chương/bài nếu VLM trả — chỉ là metadata; khoảng trang bài
    làm giàu KHÔNG dùng số trang này (dùng nhãn neo, xem _ranges_from_anchors).
    """
    chapters: list[dict[str, Any]] = []
    for page in parsed_pages:
        if not page.get("toc_page"):
            continue
        for ch in page.get("chapters", []) or []:
            name = _normalize_title(str(ch.get("name", "")))
            if not name or _is_placeholder(name):
                continue
            # Tìm xem chương này đã tồn tại trong toàn bộ danh sách chưa (tránh lặp chương khi TOC trải nhiều trang)
            target = next((c for c in chapters if c["name"].lower() == name.lower()), None)
            if target is None:
                target = {
                    "name": name,
                    "is_phu": _is_phu_title(name),
                    "page": _as_page_int(ch.get("page")),
                    "lessons": [],
                }
                chapters.append(target)

            for lesson in ch.get("lessons", []) or []:
                lesson_name = _normalize_title(str(lesson.get("name", "")))
                if not lesson_name or _is_placeholder(lesson_name):
                    continue
                # Tránh lặp bài con trong cùng một chương
                existing_lesson = next((ls for ls in target["lessons"] if ls["name"].lower() == lesson_name.lower()), None)
                if existing_lesson is not None:
                    continue

                kind = lesson.get("kind")
                target["lessons"].append(
                    {
                        "name": lesson_name,
                        "is_phu": kind == "phu" or _is_phu_title(lesson_name),
                        "page": _as_page_int(lesson.get("page")),
                    }
                )
    return chapters


def _build_tree_from_labels(page_items: list[dict[str, Any] | None]) -> list[dict[str, Any]]:
    """Dựng cây chương/bài từ nhãn VLM trên trang nội dung (fallback khi sách không có MỤC LỤC)."""
    chapters: list[dict[str, Any]] = []
    for page in page_items:
        if page is None or page.get("kind") != "content":
            continue
        ch = (page.get("chapter") or "").strip()
        ls = (page.get("lesson") or "").strip()
        if not ch and not ls:
            continue
        if not ch:
            ch = chapters[-1]["name"] if chapters else ""
            if not ch:
                continue
        if not chapters or chapters[-1]["name"] != ch:
            chapters.append({"name": ch, "is_phu": _is_phu_title(ch), "lessons": []})
        cur = chapters[-1]
        if ls and (not cur["lessons"] or cur["lessons"][-1]["name"] != ls):
            cur["lessons"].append({"name": ls, "is_phu": _is_phu_title(ls)})
    return chapters


def _clean_group_name(name: str) -> str:
    """Nhãn VLM → khóa khớp tên node: bỏ tiền tố 'Bài 1'/'1'/'Chương I' (lấy đoạn sau ':'/'.'),
    bỏ hậu tố '(tiếp theo)'. Chỉ thao tác chuỗi — không bóc dữ liệu, không regex."""
    name = name.strip()
    for sep in (":", "."):
        if sep in name:
            name = name.rsplit(sep, 1)[-1]
    if "(" in name:
        name = name.split("(", 1)[0]
    return name.strip().lower()


def _anchor_id(value: Any) -> int | None:
    """Lấy ID (số) từ giá trị `lesson` VLM trả: số nguyên, hoặc chuỗi bắt đầu bằng số (vd '3. [...]').

    Không dùng regex — duyệt từng ký tự chữ số ở đầu chuỗi.
    """
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return int(value)
    digits = ""
    for ch in str(value or "").strip():
        if ch.isdigit():
            digits += ch
        else:
            break
    return int(digits) if digits else None


def _build_anchor_list(chapters: list[dict[str, Any]]) -> str:
    """Danh sách neo MỤC LỤC: mỗi dòng 'ID. [Tên chương] Tên bài' — ID để VLM chọn.

    Kèm tiền tố chương để phân biệt các bài trùng tên ở các chương khác nhau. Thứ tự ID
    khớp thứ tự (ci, li) khi enumerate lại chapters → ánh xạ ID → bài là ánh xạ 1-1.
    """
    lines: list[str] = []
    for ch in chapters:
        lessons = ch.get("lessons") or []
        if lessons:
            for lesson in lessons:
                lines.append(f"{len(lines) + 1}. [{ch['name']}] {lesson['name']}")
        else:
            lines.append(f"{len(lines) + 1}. {ch['name']}")
    return "\n".join(lines)


def _ranges_from_anchors(
    page_items: list[dict[str, Any] | None], chapters: list[dict[str, Any]]
) -> dict[tuple[int, int | None], tuple[int, int]] | None:
    """Bản đồ (chapter_idx, lesson_idx) → khoảng trang [start, end) từ NHÃN NEO VLM.

    Mỗi trang nội dung được VLM gán `lesson` = ID trong danh sách MỤC LỤC (bước B) → bài X
    = [trang gán X đầu tiên, trang cuối cùng + 1). KHÔNG dùng số trang in → chịu được file
    cắt ngắn và MỤC LỤC không có số trang. Trang không gán được → bỏ qua (không thuộc bài nào).
    """
    id_to_key: dict[int, tuple[int, int | None]] = {}
    name_to_key: dict[str, tuple[int, int | None]] = {}
    idx = 1
    for ci, ch in enumerate(chapters):
        lessons = ch.get("lessons") or []
        if lessons:
            for li, lesson in enumerate(lessons):
                id_to_key[idx] = (ci, li)
                name_to_key.setdefault(_clean_group_name(lesson["name"]), (ci, li))
                idx += 1
        else:
            id_to_key[idx] = (ci, None)
            name_to_key.setdefault(_clean_group_name(ch["name"]), (ci, None))
            idx += 1
    if not id_to_key:
        return None

    spans: dict[tuple[int, int | None], list[int]] = {}
    for i, page in enumerate(page_items):
        if page is None or page.get("kind") != "content":
            continue
        key: tuple[int, int | None] | None = None
        aid = _anchor_id(page.get("lesson"))
        if aid is not None:
            key = id_to_key.get(aid)
        if key is None:
            # Fallback: VLM trả tên bài (không phải ID) → khớp tên đã làm sạch với cây.
            name = _clean_group_name(str(page.get("lesson") or ""))
            if name:
                key = name_to_key.get(name)
        if key is not None:
            spans.setdefault(key, []).append(i)
    if not spans:
        return None
    return {key: (min(pages), max(pages) + 1) for key, pages in spans.items()}


def _sample_indices(start: int, end: int, cap: int) -> list[int]:
    """Chỉ số trang trong [start, end] (inclusive) — lấy mẫu đều nếu dài hơn cap, LUÔN giữ trang đầu + cuối."""
    if end < start:
        return []
    total = end - start + 1
    if total <= cap:
        return list(range(start, end + 1))
    if cap <= 1:
        return [start]
    step = (total - 1) / (cap - 1)
    return [start + round(i * step) for i in range(cap)]


def _clean_keywords(value: Any) -> list[str]:
    """Làm sạch keywords do VLM trả (danh sách chuỗi) — dedup (không phân biệt hoa thường) + giới hạn."""
    if not isinstance(value, list):
        return []
    seen: list[str] = []
    lowered: set[str] = set()
    for item in value:
        if not isinstance(item, str):
            continue
        kw = item.strip().strip(".,;:")
        key = kw.lower()
        if not kw or key in lowered:
            continue
        seen.append(kw)
        lowered.add(key)
        if len(seen) >= _MAX_KEYWORDS:
            break
    return seen


def _clean_sections(value: Any) -> list[dict[str, str]]:
    """Làm sạch sections do VLM trả — danh sách tên mục con theo thứ tự [{name}], dedup + giới hạn.

    Không có kind taxonomy — nội dung do VLM quyết định, code chỉ kiểm tra {name} hợp lệ.
    """
    if not isinstance(value, list):
        return []
    out: list[dict[str, str]] = []
    for item in value:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "").strip().strip(".,;:")
        if not name or any(name == s["name"] for s in out):
            continue
        out.append({"name": name})
        if len(out) >= _MAX_SECTIONS:
            break
    return out


def _join_summaries(summaries: list[str]) -> str:
    """Nối tóm tắt các bài con thành tóm tắt chương (giới hạn độ dài)."""
    joined = " ".join(s.strip() for s in summaries if s and s.strip())
    if len(joined) > _MAX_SUMMARY_CHARS:
        joined = joined[: _MAX_SUMMARY_CHARS].rstrip() + "…"
    return joined


def _get_runtime_settings(settings: Settings | None = None, vlm_model: str | None = None) -> Settings:
    """Tạo runtime settings với model và provider được suy luận linh hoạt từ tên model."""
    base = settings or get_settings()
    if not vlm_model or not vlm_model.strip():
        return base
    model = vlm_model.strip()
    update_dict: dict[str, Any] = {}
    if model == "qwen3-vl-flash":
        update_dict["vlm_provider"] = "qwen"
        update_dict["qwen_vlm_model"] = model
    elif "/" in model or any(k in model.lower() for k in ("gemini", "mimo", "gpt", "claude")):
        update_dict["vlm_provider"] = "openrouter"
        update_dict["openrouter_vlm_model"] = model
    else:
        update_dict["vlm_model"] = model
    return base.model_copy(update=update_dict)


def _spot_check_offset_with_vlm(
    pdf_path: Path, est_pdf_idx: int, total_pages: int, settings: Settings | None = None
) -> int | None:
    """Spot check 1 trang bằng VLM (1 request duy nhất, ~0.8s) để xác định Ground-Truth Offset khi page_map rỗng."""
    if not (0 <= est_pdf_idx < total_pages):
        return None
    try:
        import fitz
        with fitz.open(pdf_path) as doc:
            page = doc.load_page(est_pdf_idx)
            img_b64 = base64.b64encode(page.get_pixmap(dpi=120).tobytes("png")).decode("ascii")
        s = settings or get_settings()
        prompt = (
            "Đọc số trang in (số trang sách) ở góc chân trang (footer) hoặc đầu trang (header) của ảnh. "
            'Trả về đúng 1 JSON: {"printed_page": <số nguyên hoặc null>}'
        )
        ans = vlm._chat_completions(img_b64, s, prompt)
        data = json.loads(vlm._clean_json(ans))
        printed = data.get("printed_page")
        if printed and isinstance(printed, int) and printed > 0:
            return est_pdf_idx - printed
    except Exception as exc:
        logger.warning("VLM spot check offset failed: %s", exc)
    return None


def _align_unit_page_ranges(
    chapters: list[dict[str, Any]],
    page_map: dict[int, int],
    total_pages: int,
    toc_pages: list[int] | None = None,
    pdf_path: Path | None = None,
    settings: Settings | None = None,
) -> None:
    """Khớp số trang in từ TOC với bảng ánh xạ page_map hoặc Dynamic Offset để chốt dải trang [start_page, end_page] cho từng bài."""
    flat_units: list[dict[str, Any]] = []
    for ch in chapters:
        lessons = ch.get("lessons") or []
        if lessons:
            for ls in lessons:
                flat_units.append(ls)
        else:
            flat_units.append(ch)

    # 1. Tính Dynamic Offset nếu page_map rỗng (PDF Scan không có text layer)
    dom_offset = 0
    if page_map:
        offsets = [pdf_idx - pr_pg for pr_pg, pdf_idx in page_map.items()]
        dom_offset = max(set(offsets), key=offsets.count)
    else:
        all_printed: list[int] = []
        for ch in chapters:
            p = ch.get("page")
            if p is not None and isinstance(p, int) and p > 0:
                all_printed.append(p)
            for ls in ch.get("lessons", []):
                lp = ls.get("page")
                if lp is not None and isinstance(lp, int) and lp > 0:
                    all_printed.append(lp)
        first_printed = min(all_printed) if all_printed else None

        # Spot check calibration bằng VLM trên bài học đầu tiên (nếu có file PDF)
        spot_offset = None
        s = settings or get_settings()
        if pdf_path and first_printed is not None and vlm.is_configured(s):
            raw_est = (max(toc_pages) + 1) if toc_pages else 0
            # Thử spot check tại trang first_printed hoặc raw_est
            spot_offset = _spot_check_offset_with_vlm(pdf_path, first_printed, total_pages, settings=s)
            if spot_offset is None and raw_est != first_printed:
                spot_offset = _spot_check_offset_with_vlm(pdf_path, raw_est, total_pages, settings=s)

        if spot_offset is not None and -5 <= spot_offset <= 10:
            dom_offset = spot_offset
        elif toc_pages and first_printed is not None:
            raw_offset = (max(toc_pages) + 1) - first_printed
            # Sanity bound [-5 .. 10]: Offset của SGK thông thường chỉ nằm trong dải này
            if -5 <= raw_offset <= 10:
                dom_offset = raw_offset
            else:
                dom_offset = 0
        else:
            dom_offset = 0

    # 2. Gán start_page từ page_map hoặc Dynamic Offset
    for unit in flat_units:
        printed = unit.get("page")
        if printed and isinstance(printed, int) and printed in page_map:
            pdf_idx = page_map[printed]
            unit["start_page"] = pdf_idx if 0 <= pdf_idx < total_pages else None
        elif printed and isinstance(printed, int):
            est_idx = printed + dom_offset
            # CHỈ gán start_page nếu trang ước tính NẰM TRONG file PDF thực tế
            if 0 <= est_idx < total_pages:
                unit["start_page"] = est_idx
            else:
                # Trang vượt ngoài file PDF (file bị cắt ngắn/chỉ có 1 phần) -> gán None, KHÔNG clamp!
                unit["start_page"] = None
        else:
            unit["start_page"] = None

    # 3. Lọc các unit có start_page hợp lệ trong file PDF và chốt end_page
    valid_units = [u for u in flat_units if u.get("start_page") is not None]
    valid_units.sort(key=lambda u: u["start_page"])

    for i, unit in enumerate(valid_units):
        start_p = unit["start_page"]
        if i + 1 < len(valid_units):
            next_start = valid_units[i + 1]["start_page"]
            end_p = max(start_p, next_start - 1)
        else:
            end_p = total_pages - 1
        unit["end_page"] = end_p

    # Các unit không có start_page trong file thì end_page cũng là None
    for unit in flat_units:
        if unit.get("start_page") is None:
            unit["end_page"] = None


def extract_book_structure(
    content: bytes,
    progress_cb: ProgressCb | None = None,
    vlm_model: str | None = None,
    settings: Settings | None = None,
) -> tuple[list[dict[str, Any]], dict[int, int], list[str], Path]:
    """Quét PDF theo Windowed Batch TOC Pipeline (tối ưu hóa độ chính xác và tốc độ):
    1. Đọc số trang in ở footer/header (PyMuPDF) + nội suy tuyến tính -> page_map (0.1s).
    2. Định vị toàn bộ dải trang MỤC LỤC bằng find_toc_pages (hỗ trợ cả TOC 1 trang và nhiều trang).
    3. Gửi toàn bộ dải ảnh Mục Lục vào VLM trong 1 request duy nhất -> trích xuất cây Chương/Bài toàn diện.
    4. Khớp số trang in sang PDF index thực tế, chốt ranh giới [start_page, end_page] cho từng bài.
    """
    s = _get_runtime_settings(settings, vlm_model)
    if not vlm.is_configured(s):
        raise ValueError("Cần cấu hình VLM_API_KEY để nạp sách PDF — VLM là bắt buộc cho luồng này.")
    _TMP_DIR.mkdir(parents=True, exist_ok=True)
    tmp = _TMP_DIR / f"sweep_{uuid.uuid4().hex}.pdf"
    tmp.write_bytes(content)
    try:
        import fitz

        with fitz.open(tmp) as doc:
            total_pages = doc.page_count

        warnings: list[str] = []
        if progress_cb:
            progress_cb(1, 4, "scan")

        # Giai đoạn 1: Bảng ánh xạ số trang in -> PDF page index
        page_map = ld.build_page_number_map(tmp)

        # Giai đoạn 2: Định vị các trang Mục Lục (TOC) (hỗ trợ cả TOC 1 trang và TOC nhiều trang liên tiếp)
        toc_pages = ld.find_toc_pages(tmp, settings=s)
        if progress_cb:
            progress_cb(2, 4, "scan")

        # Giai đoạn 3: Gọi VLM trích xuất Cây Mục Lục Động từ toàn bộ dải trang TOC trong 1 request
        if not toc_pages:
            warnings.append("Không tìm thấy trang Mục Lục trong sách.")
            chapters = []
        else:
            raw_toc = vlm.read_toc_pages(tmp, toc_pages, settings=s)
            chapters = vlm.parse_dynamic_toc_json(raw_toc)
            if not chapters:
                warnings.append("Không bóc tách được cây Mục Lục từ các trang TOC.")

        # Giai đoạn 4: Ánh xạ ranh giới [start_page, end_page] cho từng bài học
        _align_unit_page_ranges(chapters, page_map, total_pages, toc_pages=toc_pages, pdf_path=tmp, settings=s)
        if progress_cb:
            progress_cb(4, 4, "scan")

        return chapters, page_map, warnings, tmp
    except vlm.VlmUnavailableError as exc:
        tmp.unlink(missing_ok=True)
        raise ValueError(str(exc)) from exc
    except Exception as exc:
        tmp.unlink(missing_ok=True)
        raise exc


def _enrich_units_chunked(
    chapters: list[dict[str, Any]],
    page_map: dict[int, int],
    pdf_path: Path,
    progress_cb: ProgressCb | None = None,
    vlm_model: str | None = None,
    settings: Settings | None = None,
) -> list[str]:
    """Làm giàu từng bài bằng VLM song song (ThreadPoolExecutor): summary + keywords + sections.
    Cắt lát chính xác các trang theo dải [start_page, end_page] của bài học.
    """
    s = _get_runtime_settings(settings, vlm_model)
    warnings: list[str] = []
    targets: list[tuple[int, int | None, dict[str, Any]]] = []
    for ci, ch in enumerate(chapters):
        lessons = ch.get("lessons") or []
        if lessons:
            targets.extend((ci, li, lesson) for li, lesson in enumerate(lessons))
        else:
            targets.append((ci, None, ch))

    total = len(targets)
    done_count = 0

    valid_tasks: list[tuple[int, int | None, dict[str, Any], int, int]] = []
    for ci, li, node in targets:
        start_p = node.get("start_page")
        end_p = node.get("end_page")
        if start_p is None or end_p is None:
            warnings.append(
                f"Bài '{node['name']}': không có trang nội dung trong file này (file có thể bị cắt ngắn / chỉ tải 1 phần sách) — bỏ làm giàu."
            )
            node["summary"] = None
            node["keywords"] = []
            node["sections"] = []
            done_count += 1
            if progress_cb:
                progress_cb(done_count, total, "enrich")
        else:
            valid_tasks.append((ci, li, node, start_p, end_p))

    if valid_tasks:
        max_workers = max(1, min(s.vlm_max_concurrency, len(valid_tasks)))

        def process_one(task: tuple[int, int | None, dict[str, Any], int, int]) -> tuple[dict[str, Any], str | None]:
            ci, _li, node, start, end = task
            indices = _sample_indices(start, end, _MAX_REFINE_PAGES)
            chapter_name = chapters[ci].get("name") if ci < len(chapters) else None
            try:
                raw = vlm.read_lesson_pages(
                    pdf_path, indices, lesson_name=node["name"], chapter_name=chapter_name, settings=s
                )
            except vlm.VlmUnavailableError as exc:
                return node, f"Bài '{node['name']}': {exc}"

            data = _parse_json_object(raw)
            if data is None or not isinstance(data.get("summary"), str):
                return node, f"Bài '{node['name']}': VLM trả JSON làm giàu không hợp lệ — bỏ làm giàu."

            node["summary"] = data.get("summary").strip() or None
            node["keywords"] = _clean_keywords(data.get("keywords"))
            node["sections"] = _clean_sections(data.get("sections"))
            return node, None

        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = [pool.submit(process_one, t) for t in valid_tasks]
            for fut in as_completed(futures):
                try:
                    _, warn = fut.result()
                    if warn:
                        warnings.append(warn)
                except Exception as exc:
                    warnings.append(f"Lỗi khi làm giàu bài: {exc}")
                finally:
                    done_count += 1
                    if progress_cb:
                        progress_cb(done_count, total, "enrich")

    # Tổng hợp tóm tắt & từ khóa cấp chương
    for ch in chapters:
        lessons = ch.get("lessons") or []
        summaries = [ls.get("summary") for ls in lessons if ls.get("summary")]
        if summaries:
            ch["summary"] = _join_summaries(summaries)
        keywords: list[str] = []
        for ls in lessons:
            for kw in ls.get("keywords") or []:
                if kw not in keywords:
                    keywords.append(kw)
        if keywords:
            ch["keywords"] = keywords[:_MAX_KEYWORDS]
    return warnings


def _enrich_chapters(
    chapters: list[dict[str, Any]],
    page_items_or_map: Any,
    pdf_path: Path,
    progress_cb: ProgressCb | None = None,
) -> list[str]:
    """Alias tương thích cho _enrich_units_chunked."""
    page_map = page_items_or_map if isinstance(page_items_or_map, dict) else {}
    return _enrich_units_chunked(chapters, page_map, pdf_path, progress_cb=progress_cb)



def extract_toc_from_text(text: str) -> list[TocEntry]:
    """Dò dòng "Chương X: Tên" / "Bài n: Tên" trong text → TOC entries (dùng cho TXT/MD tay)."""
    entries: list[TocEntry] = []
    for idx, line in enumerate(text.splitlines()):
        line = line.strip()
        chapter = _CHAPTER_RE.match(line)
        if chapter:
            entries.append((1, chapter.group(2).strip(), idx))
            continue
        lesson = _LESSON_RE.match(line)
        if lesson:
            entries.append((2, lesson.group(2).strip(), idx))
    return entries


def extract_toc_from_docx(content: bytes) -> list[TocEntry]:
    """Trích TOC từ DOCX bằng heading styles (Heading 1 = chương, Heading 2 = bài)."""
    import docx

    document = docx.Document(__import__("io").BytesIO(content))
    entries: list[TocEntry] = []
    for idx, para in enumerate(document.paragraphs):
        style = (para.style.name if para.style else "") or ""
        if style.lower().startswith("heading"):
            level = 1 if "1" in style else 2
            if para.text.strip():
                entries.append((level, para.text.strip(), idx))
    return entries


def _entries_to_chapters(entries: list[TocEntry]) -> list[dict[str, Any]]:
    """Chuyển TocEntry (level/title) → chapters dict; chuẩn hóa + tag phụ + loại placeholder."""
    chapters: list[dict[str, Any]] = []
    for level, title, _page in entries:
        name = _normalize_title(title)
        if not name or _is_placeholder(name):
            continue
        if level == 1:
            chapters.append({"name": name, "is_phu": _is_phu_title(name), "lessons": []})
        elif chapters:
            chapters[-1]["lessons"].append({"name": name, "is_phu": _is_phu_title(name)})
    return chapters


def _sanity_check(chapters: list[dict[str, Any]]) -> list[str]:
    """Cảnh báo bất thường (số chương/bài vượt ngưỡng, trùng tên) — không chặn, chỉ cảnh báo."""
    warnings: list[str] = []
    if len(chapters) > _MAX_CHAPTERS_PER_SEMESTER:
        warnings.append(
            f"Phát hiện {len(chapters)} chương — nhiều hơn mức thường gặp cho 1 học kỳ "
            f"({_MAX_CHAPTERS_PER_SEMESTER}); kiểm tra xem có lẫn nội dung ruột sách không."
        )
    for ch in chapters:
        if len(ch["lessons"]) > _MAX_LESSONS_PER_CHAPTER:
            warnings.append(
                f"Chương '{ch['name']}' có {len(ch['lessons'])} bài — nhiều hơn mức bình thường; kiểm tra lại."
            )
    names = [ch["name"] for ch in chapters]
    if len(set(names)) != len(names):
        warnings.append("Có chương trùng tên — kiểm tra lại trước khi lưu.")
    return warnings


def build_unit_specs_from_chapters(
    chapters: list[dict[str, Any]],
    subject_code: str,
    grade: int,
    semester: int | None,
    include_lessons: bool,
    volume: str | None = None,
) -> list[dict[str, Any]]:
    """Chuyển chapters → spec curriculum_units: chương C1.., bài con {chương}_B{n} (parent_code).

    Hỗ trợ phân lập mã theo volume (Tập 1 -> _T1_, Tập 2 -> _T2_).
    Truyền cả nội dung làm giàu (summary/keywords/sections) vào spec để lưu cùng node.
    """
    specs: list[dict[str, Any]] = []
    # subject_code có thể đã gắn khối (TOAN_6, TOAN_7...) từ dropdown 24 môn — nếu hậu tố
    # khớp grade thì bỏ để code node gọn (TOAN6_C1 thay vì TOAN_66_C1).
    base = subject_code.upper().strip()
    if base.endswith(f"_{grade}"):
        base = base[: -len(f"_{grade}")]

    # Xác định tag phân lập theo tập sách (volume)
    vol_tag = ""
    if volume:
        v_clean = volume.strip().lower()
        if "2" in v_clean or "t2" in v_clean:
            vol_tag = "_T2"
        elif "1" in v_clean or "t1" in v_clean:
            vol_tag = "_T1"
    elif semester == 2:
        vol_tag = "_T2"
    elif semester == 1:
        vol_tag = "_T1"

    prefix = f"{base}{grade}{vol_tag}"
    for idx, ch in enumerate(chapters, start=1):
        code = f"{prefix}_C{idx}"
        specs.append(
            {
                "code": code,
                "name": ch["name"],
                "semester_number": semester,
                "parent_code": None,
                "is_phu": ch.get("is_phu", False),
                "summary": ch.get("summary"),
                "keywords": ch.get("keywords") or None,
                "sections": ch.get("sections") or None,
                "start_page": ch.get("start_page"),
                "end_page": ch.get("end_page"),
            }
        )
        if not include_lessons:
            continue
        for jdx, lesson in enumerate(ch.get("lessons", []), start=1):
            specs.append(
                {
                    "code": f"{code}_B{jdx}",
                    "name": lesson["name"],
                    "semester_number": semester,
                    "parent_code": code,
                    "is_phu": lesson.get("is_phu", False),
                    "summary": lesson.get("summary"),
                    "keywords": lesson.get("keywords") or None,
                    "sections": lesson.get("sections") or None,
                    "start_page": lesson.get("start_page"),
                    "end_page": lesson.get("end_page"),
                }
            )
    return specs


def upsert_unit_tree(
    db: Session,
    specs: list[dict[str, Any]],
    subject_id: int,
    grade: int,
    book_id: int | None = None,
    overwrite_enrichment: bool = True,
) -> tuple[int, int]:
    """Upsert chương trước, rồi bài con gắn parent_id theo parent_code. Trả (inserted, updated).

    book_id (nếu có) sẽ gắn vào từng node để biết cuốn SGK nguồn.
    overwrite_enrichment=True (mặc định): cập nhật trực tiếp dữ liệu làm giàu theo spec mới nhất
    (xóa sạch tóm tắt/từ khóa cũ nếu spec mới là None).
    """
    inserted = updated = 0
    code_to_id: dict[str, int] = {}
    for spec in specs:
        unit = db.execute(
            select(CurriculumUnit).where(
                CurriculumUnit.subject_id == subject_id,
                CurriculumUnit.grade_number == grade,
                CurriculumUnit.code == spec["code"],
            )
        ).scalars().first()
        parent_id = code_to_id.get(spec["parent_code"]) if spec["parent_code"] else None
        if unit is None:
            unit = CurriculumUnit(
                subject_id=subject_id,
                grade_number=grade,
                code=spec["code"],
                name=spec["name"],
                semester_number=spec["semester_number"],
                parent_id=parent_id,
                is_phu=spec.get("is_phu", False),
                book_id=book_id,
                summary=spec.get("summary"),
                keywords=spec.get("keywords"),
                sections=spec.get("sections"),
                start_page=spec.get("start_page"),
                end_page=spec.get("end_page"),
            )
            db.add(unit)
            inserted += 1
        else:
            unit.name = spec["name"]
            unit.semester_number = spec["semester_number"]
            unit.parent_id = parent_id
            unit.is_phu = spec.get("is_phu", False)
            unit.is_active = True
            if spec.get("start_page") is not None:
                unit.start_page = spec["start_page"]
            if spec.get("end_page") is not None:
                unit.end_page = spec["end_page"]
            if overwrite_enrichment:
                unit.summary = spec.get("summary")
                unit.keywords = spec.get("keywords")
                unit.sections = spec.get("sections")
            else:
                if spec.get("summary") is not None:
                    unit.summary = spec["summary"]
                if spec.get("keywords") is not None:
                    unit.keywords = spec["keywords"]
                if spec.get("sections") is not None:
                    unit.sections = spec["sections"]
            if book_id is not None:
                unit.book_id = book_id
            updated += 1
        db.flush()
        code_to_id[spec["code"]] = unit.id
    db.commit()
    return inserted, updated


def save_catalog_from_preview(
    db: Session,
    chapters: list[dict[str, Any]],
    subject_code: str,
    grade: int,
    semester: int | None = None,
    book_id: int | None = None,
    overwrite_enrichment: bool = True,
) -> dict[str, Any]:
    """Lưu cây chương/bài (đã trích xuất ở bước dry_run) thẳng vào curriculum_units.

    KHÔNG trích lại file, KHÔNG gọi VLM — upsert theo code đã xem trước (idempotent).
    book_id (nếu có) gắn vào node để biết cuốn SGK nguồn.
    """
    specs: list[dict[str, Any]] = []
    for chapter in chapters:
        code = chapter["code"]
        specs.append(
            {
                "code": code,
                "name": chapter["name"],
                "semester_number": chapter.get("semester_number") or semester,
                "parent_code": None,
                "is_phu": chapter.get("is_phu", False),
                "summary": chapter.get("summary"),
                "keywords": chapter.get("keywords") or None,
                "sections": chapter.get("sections") or None,
            }
        )
        for lesson in chapter.get("lessons", []):
            specs.append(
                {
                    "code": lesson["code"],
                    "name": lesson["name"],
                    "semester_number": chapter.get("semester_number") or semester,
                    "parent_code": code,
                    "is_phu": lesson.get("is_phu", False),
                    "summary": lesson.get("summary"),
                    "keywords": lesson.get("keywords") or None,
                    "sections": lesson.get("sections") or None,
                }
            )
    if not specs:
        raise ValueError("Không có chương nào để lưu.")
    subject_ids = resolve_subject_ids(db, subject_code, [grade])
    subject_id = subject_ids.get(grade)
    if subject_id is None:
        raise ValueError(f"Không có s360.dim_subject cho {subject_code.upper()}_{grade} — nạp môn trước.")
    inserted, updated = upsert_unit_tree(db, specs, subject_id, grade, book_id=book_id, overwrite_enrichment=overwrite_enrichment)
    hidden = deactivate_placeholder_units(db)
    return {
        "subject_code": subject_code.upper(),
        "grade": grade,
        "semester": semester,
        "source": "preview",
        "chapters": chapters,
        "inserted": inserted,
        "updated": updated,
        "hidden_placeholders": hidden,
        "warnings": [],
        "dry_run": False,
        "book_id": book_id,
    }


def _extract_non_pdf(content: bytes, filename: str) -> tuple[list[dict[str, Any]], str]:
    """DOCX/TXT/MD → (chapters, source). Không có trang nội dung nên KHÔNG làm giàu."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        return _entries_to_chapters(extract_toc_from_docx(content)), "docx"
    if ext in (".txt", ".md"):
        text = content.decode("utf-8", errors="replace")
        return _entries_to_chapters(extract_toc_from_text(text)), "text"
    raise ValueError(f"Định dạng không hỗ trợ: {ext} (chỉ PDF/DOCX/TXT/MD)")


def ingest_book(
    db: Session,
    filename: str,
    content: bytes,
    subject_code: str,
    grade: int,
    semester: int | None = None,
    include_lessons: bool = False,
    dry_run: bool = False,
    book_id: int | None = None,
    enrich: bool = True,
    progress_cb: ProgressCb | None = None,
    vlm_model: str | None = None,
    settings: Settings | None = None,
    volume: str | None = None,
) -> dict[str, Any]:
    """Nạp sách → (PDF: quét toàn cuốn 1 lần + làm giàu | TXT/DOCX: như cũ) → preview/lưu. KHÔNG RAG."""
    s = _get_runtime_settings(settings, vlm_model)
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        pdf_path: Path | None = None
        try:
            chapters, page_map, warnings, pdf_path = extract_book_structure(
                content, progress_cb=progress_cb, vlm_model=vlm_model, settings=s
            )
            if not chapters:
                raise ValueError(
                    "Không trích được mục lục: VLM không tìm thấy trang MỤC LỤC và không nhận diện "
                    "được tiêu đề chương/bài trên trang nội dung. Hãy thử lại, hoặc dùng file mục lục JSON/markdown."
                )
            if enrich:
                warnings += _enrich_units_chunked(
                    chapters, page_map, pdf_path, progress_cb=progress_cb, vlm_model=vlm_model, settings=s
                )
        finally:
            if pdf_path is not None:
                pdf_path.unlink(missing_ok=True)
            source = "pdf-vlm"
    else:
        chapters, source = _extract_non_pdf(content, filename)
        warnings = []
        if not chapters:
            raise ValueError(
                "Không trích được mục lục: không tìm thấy trang MỤC LỤC trong sách "
                "(PDF cần VLM hoạt động; nếu là DOCX/TXT hãy kiểm tra cấu trúc). "
                "Hãy thử lại, hoặc dùng file mục lục JSON/markdown."
            )
    if semester is None:
        semester = detect_semester_from_filename(filename)
    warnings += _sanity_check(chapters)
    specs = build_unit_specs_from_chapters(chapters, subject_code, grade, semester, include_lessons, volume=volume)
    preview_chapters = [
        {
            "code": spec["code"],
            "name": spec["name"],
            "semester_number": spec["semester_number"],
            "is_phu": spec["is_phu"],
            "start_page": spec.get("start_page"),
            "end_page": spec.get("end_page"),
            "summary": spec.get("summary"),
            "keywords": spec.get("keywords") or None,
            "sections": spec.get("sections") or None,
            "lessons": [
                {
                    "code": child["code"],
                    "name": child["name"],
                    "is_phu": child["is_phu"],
                    "start_page": child.get("start_page"),
                    "end_page": child.get("end_page"),
                    "summary": child.get("summary"),
                    "keywords": child.get("keywords") or None,
                    "sections": child.get("sections") or None,
                }
                for child in specs
                if child["parent_code"] == spec["code"]
            ],
        }
        for spec in specs
        if spec["parent_code"] is None
    ]
    if dry_run:
        return {
            "subject_code": subject_code.upper(),
            "grade": grade,
            "semester": semester,
            "source": source,
            "chapters": preview_chapters,
            "inserted": 0,
            "updated": 0,
            "hidden_placeholders": 0,
            "warnings": warnings,
            "dry_run": True,
        }

    subject_ids = resolve_subject_ids(db, subject_code, [grade])
    subject_id = subject_ids.get(grade)
    if subject_id is None:
        raise ValueError(f"Không có s360.dim_subject cho {subject_code.upper()}_{grade} — nạp môn trước.")
    inserted, updated = upsert_unit_tree(db, specs, subject_id, grade, book_id=book_id)
    hidden = deactivate_placeholder_units(db)
    return {
        "subject_code": subject_code.upper(),
        "grade": grade,
        "semester": semester,
        "source": source,
        "chapters": preview_chapters,
        "inserted": inserted,
        "updated": updated,
        "hidden_placeholders": hidden,
        "warnings": warnings,
        "dry_run": False,
    }
