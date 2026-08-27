import os
import re
import uuid
from typing import Literal

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt, RGBColor
from langchain_core.tools import tool
from sqlalchemy import and_, func, select

from src.agents.context import current_user_school_id
from src.agents.report_agent.chart_generator import generate_chart_for_table
from src.agents.report_agent.queries import (
    compute_report_data,
    is_valid_int,
    resolve_parameters,
)
from src.agents.report_agent.visual_contracts import (
    ColumnAlignment,
    detect_cell_alignment,
    sanitize_delta_value,
)
from src.db.session import SessionLocal
from src.models.tables import User as DBUser
from src.schemas.analytics import ReportExportRequestS360

@tool
def get_report_data_summary(
    report_type: Literal["academic_conduct", "subject_quality", "at_risk", "subject_report"],
    grade_level: str = "all",
    class_id: str = None,
    semester_id: str = None,
    subject_id: str = None,
    school_year_id: str = None,
) -> str:
    """Tra cứu và tổng hợp số liệu báo cáo thống kê phục vụ cho việc hiển thị bảng số liệu trực tiếp.

    Args:
        report_type: Loại báo cáo ('academic_conduct', 'subject_quality', 'at_risk', 'subject_report').
        grade_level: Khối lớp học ('all' hoặc số khối ví dụ: '7', '8', '10', '11', '12').
        class_id: ID hoặc Tên lớp học cụ thể (tùy chọn, ví dụ: '10A1', '8B').
        semester_id: ID hoặc Tên học kỳ/Số học kỳ (tùy chọn, ví dụ: 'Học kỳ 1', 'HK1', '1', 'Học kỳ 2', 'HK2', '2').
        subject_id: ID hoặc Tên/Mã môn học cụ thể (tùy chọn, ví dụ: 'Toán', 'Ngữ văn', 'Tiếng Anh').
        school_year_id: ID hoặc Tên niên khóa (tùy chọn, ví dụ: '2025-2026').
    """
    school_id = current_user_school_id.get()
    if not school_id:
        return "Lỗi: Không xác định được trường học. Vui lòng đăng nhập."

    with SessionLocal() as db:
        data = compute_report_data(db, school_id, report_type, grade_level, class_id, semester_id, subject_id, school_year_id)

        report_titles = {
            "academic_conduct": "BÁO CÁO TỔNG KẾT KẾT QUẢ HỌC TẬP VÀ RÈN LUYỆN",
            "subject_quality": "BÁO CÁO PHÂN TÍCH PHỔ ĐIỂM VÀ CHẤT LƯỢNG BỘ MÔN",
            "at_risk": "BÁO CÁO SÀNG LỌC VÀ THEO DÕI NHÓM HỌC SINH CẦN HỖ TRỢ SƯ PHẠM",
            "subject_report": "BÁO CÁO CHUYÊN SÂU MÔN HỌC",
        }
        title_text = report_titles.get(report_type, f"BÁO CÁO THỐNG KÊ {report_type.upper()}")

        summary = f"### {title_text}\n"
        summary += f"- **Phạm vi**: {data['selected_grade_name']}{data['selected_class_name']}\n"
        summary += f"- **Niên khóa**: {data['year_name']} / **Học kỳ**: {data['sem_name']}\n\n"

        summary += "| Chỉ số thống kê | Giá trị thực tế |\n"
        summary += "| --- | --- |\n"
        summary += f"| Sĩ số học sinh active | {data['total_students']} học sinh |\n"
        if not class_id:
            summary += f"| Tổng số lớp học hoạt động | {data['total_classes']} lớp học |\n"
        summary += f"| GPA trung bình | {data['gpa'] or 0.0} / 10 |\n"
        if not class_id:
            summary += f"| Số lớp cần can thiệp học thuật | {data['at_risk']} lớp |\n"

        if report_type == "subject_quality" and data["subject_averages"]:
            summary += "\n#### Điểm trung bình các môn học:\n"
            summary += "| Môn học | Điểm trung bình |\n"
            summary += "| --- | --- |\n"
            for item in data["subject_averages"]:
                summary += f"| {item['Môn học']} | {item['ĐTB']} |\n"

        elif report_type == "academic_conduct":
            summary += "\n#### Phân loại hạnh kiểm:\n"
            summary += "| Loại hạnh kiểm | Số học sinh |\n"
            summary += "| --- | --- |\n"
            summary += f"| Tốt | {data['conduct_stats']['TOT']} học sinh |\n"
            summary += f"| Khá | {data['conduct_stats']['KHA']} học sinh |\n"
            summary += f"| Đạt | {data['conduct_stats']['TRUNG_BINH']} học sinh |\n"
            summary += f"| Chưa đạt | {data['conduct_stats']['YEU']} học sinh |\n"

        elif report_type == "at_risk":
            summary += "\n#### Thống kê nhóm học sinh cần hỗ trợ sư phạm:\n"
            summary += f"- **Số lượng lớp cảnh báo có ĐTB < 5.0**: {data['at_risk']} lớp\n"

        elif report_type == "subject_report":
            summary += "\n#### Báo cáo chuyên sâu môn học:\n"
            summary += f"- Học sinh tham gia thi: {data['total_students']} học sinh\n"
            summary += f"- Điểm trung bình môn: {data['gpa'] or 0.0} / 10\n"

        return summary


@tool
async def generate_report_download_link(
    report_type: Literal["academic_conduct", "subject_quality", "at_risk", "subject_report"],
    format: Literal["docx", "pdf", "html"],
    grade_level: str = "all",
    class_id: str = None,
    semester_id: str = None,
    subject_id: str = None,
    school_year_id: str = None,
    include_ai_insights: bool = True,
    include_tables: bool = True,
    include_signature: bool = True,
) -> str:
    """Tạo tệp báo cáo thống kê thực tế ở server và trả về link tải trực tiếp trong khung chat.

    IMPORTANT WARNING: Công cụ này sẽ tự động tạo đồng thời cả 3 định dạng file (.docx, .pdf, .html) và trả về đường link của cả 3 định dạng này trong cùng 1 lần gọi.
    BẠN CHỈ ĐƯỢC PHÉP GỌI CÔNG CỤ NÀY ĐÚNG 1 LẦN DUY NHẤT CHO MỖI YÊU CẦU BÁO CÁO. Tuyệt đối KHÔNG gọi công cụ này nhiều lần trong vòng lặp hoặc gọi riêng rẽ cho từng định dạng.

    Args:
        report_type: Loại báo cáo ('academic_conduct', 'subject_quality', 'at_risk', 'subject_report').
        format: Định dạng tệp ('docx', 'pdf', 'html').
        grade_level: Khối lớp học ('all' hoặc số khối ví dụ: '7', '8', '10', '11', '12').
        class_id: ID hoặc Tên lớp học cụ thể (tùy chọn, ví dụ: '10A1', '8B').
        semester_id: ID hoặc Tên học kỳ/Số học kỳ (tùy chọn, ví dụ: 'Học kỳ 1', 'HK1', '1', 'Học kỳ 2', 'HK2', '2').
        subject_id: ID hoặc tên môn học cụ thể (tùy chọn).
        school_year_id: ID hoặc Tên niên khóa (tùy chọn, ví dụ: '2025-2026').
        include_ai_insights: Bao gồm nhận xét phân tích từ AI.
        include_tables: Bao gồm bảng dữ liệu chi tiết.
        include_signature: Bao gồm khung chữ ký phê duyệt.
    """
    school_id = current_user_school_id.get()
    if not school_id:
        return "Lỗi: Không xác định được trường học. Vui lòng đăng nhập."

    file_uuid = str(uuid.uuid4())[:8]
    temp_dir = "temp"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    try:
        from src.api.v1.reports import export_analytics_report_s360

        with SessionLocal() as db:
            resolved_class_id, resolved_school_year_id, resolved_semester_index, resolved_subject_id = resolve_parameters(
                db, school_id, class_id, semester_id, subject_id, school_year_id
            )
            class_id = resolved_class_id or class_id
            school_year_id = resolved_school_year_id or school_year_id
            semester_index = resolved_semester_index or 1
            subject_id = resolved_subject_id or subject_id

        formats_to_generate = ["docx", "html", "pdf"]
        generated_formats = []

        for fmt in formats_to_generate:
            try:
                with SessionLocal() as db:
                    user_row = (
                        db.execute(select(DBUser).where(DBUser.so_school_id == school_id, DBUser.is_active.is_(True)))
                        .scalars()
                        .first()
                    )
                    if not user_row:
                        continue

                    payload = ReportExportRequestS360(
                        report_type=report_type,
                        format=fmt,
                        grade_level=grade_level,
                        class_id=int(str(class_id)) if class_id and is_valid_int(class_id) else None,
                        semester_index=semester_index,
                        subject_id=int(str(subject_id)) if subject_id and is_valid_int(subject_id) else None,
                        school_year_id=int(str(school_year_id)) if school_year_id and is_valid_int(school_year_id) else None,
                        include_charts=True,
                        include_tables=include_tables,
                        include_ai_insights=include_ai_insights,
                        include_signature=include_signature,
                    )

                    response = export_analytics_report_s360(payload=payload, user=user_row, db=db)

                fmt_filepath = os.path.join(temp_dir, f"bao_cao_{report_type}_{file_uuid}.{fmt}")
                if hasattr(response, "body"):
                    with open(fmt_filepath, "wb") as f:
                        f.write(response.body)
                elif hasattr(response, "body_iterator"):
                    with open(fmt_filepath, "wb") as f:
                        async for chunk in response.body_iterator:
                            f.write(chunk)
                generated_formats.append(fmt)
            except Exception as e:
                print(f"Error generating format {fmt}: {str(e)}")

        if not generated_formats:
            return "Lỗi: Không thể tạo bất kỳ định dạng báo cáo nào."

        from src.config import get_settings

        settings = get_settings()
        base_url = settings.backend_url.rstrip('/')

        links = []
        if "html" in generated_formats:
            download_url_html = f"{base_url}/api/v1/reports/download/bao_cao_{report_type}_{file_uuid}.html"
            links.append(f"[Xem Bản Xem Trước Báo Cáo]({download_url_html})")
        if "docx" in generated_formats:
            download_url_docx = f"{base_url}/api/v1/reports/download/bao_cao_{report_type}_{file_uuid}.docx"
            links.append(f"[Tải Báo Cáo Word (.docx)]({download_url_docx})")
        if "pdf" in generated_formats:
            download_url_pdf = f"{base_url}/api/v1/reports/download/bao_cao_{report_type}_{file_uuid}.pdf"
            links.append(f"[Tải Báo Cáo PDF (.pdf)]({download_url_pdf})")

        links_str = " | ".join(links)
        return f"Tệp báo cáo đã được tạo thành công!\n\n👉 {links_str}"

    except Exception as e:
        return f"Lỗi trong quá trình tạo tệp báo cáo: {str(e)}"

def render_markdown_to_docx(title: str, content_markdown: str) -> docx.Document:
    doc = docx.Document()

    # 1. Page Setup
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # 2. Style Setup
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # 3. Parse Markdown Body
    lines = content_markdown.split("\n")

    # 4. Parse Markdown Body helper

    def add_formatted_text(paragraph, text, is_bold_default=False):
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = paragraph.add_run(part[2:-2])
                r.font.bold = True
                r.font.name = "Times New Roman"
            elif part.startswith("*") and part.endswith("*"):
                r = paragraph.add_run(part[1:-1])
                r.font.italic = True
                r.font.name = "Times New Roman"
            elif part.startswith("`") and part.endswith("`"):
                r = paragraph.add_run(part[1:-1])
                r.font.name = "Courier New"
                r.font.size = Pt(10.5)
            else:
                if part:
                    r = paragraph.add_run(part)
                    if is_bold_default:
                        r.font.bold = True
                    r.font.name = "Times New Roman"

    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        parsed_rows = []
        for r in table_rows:
            cols = [c.strip() for c in r.split("|")]
            if len(cols) >= 2:
                if cols[0] == "":
                    cols = cols[1:]
                if cols and cols[-1] == "":
                    cols = cols[:-1]
                if all(re.match(r"^\s*:-?-*:?\s*$", c) or re.match(r"^\s*-+\s*$", c) for c in cols):
                    continue
                parsed_rows.append(cols)

        if parsed_rows:
            num_cols = max(len(row) for row in parsed_rows)
            table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
            table.style = "Table Grid"
            table.alignment = 1  # Center

            headers = parsed_rows[0] if parsed_rows else []
            for r_idx, row_data in enumerate(parsed_rows):
                row = table.rows[r_idx]
                is_header = r_idx == 0
                for c_idx, val in enumerate(row_data):
                    if c_idx < len(row.cells):
                        cell = row.cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        if is_header:
                            p.alignment = 1  # Center
                            add_formatted_text(p, val, is_bold_default=True)
                            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F4F7"/>')
                            cell._tc.get_or_add_tcPr().append(shading_elm)
                        else:
                            header_name = headers[c_idx] if c_idx < len(headers) else ""
                            align = detect_cell_alignment(header_name, val)
                            if align == ColumnAlignment.RIGHT:
                                p.alignment = 2  # Right
                            elif align == ColumnAlignment.CENTER:
                                p.alignment = 1  # Center
                            else:
                                p.alignment = 0  # Left

                            # Sanitize delta if in delta column
                            if "chênh lệch" in header_name.lower() or "(δ)" in header_name.lower() or "(delta)" in header_name.lower():
                                val = sanitize_delta_value(val)

                            add_formatted_text(p, val)

            doc.add_paragraph().paragraph_format.space_after = Pt(4)

            # Tự động sinh và nhúng biểu đồ trực quan dưới bảng (nếu bảng có số liệu)
            if len(parsed_rows) >= 2:
                data_rows = parsed_rows[1:]
                chart_res = generate_chart_for_table(headers, data_rows, report_title=title)
                if chart_res:
                    chart_path, _ = chart_res
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = 1  # Center
                        p_img.paragraph_format.space_before = Pt(8)
                        p_img.paragraph_format.space_after = Pt(2)
                        p_img.paragraph_format.keep_with_next = True
                        doc.add_picture(chart_path, width=Inches(6.0))

                        p_cap = doc.add_paragraph()
                        p_cap.alignment = 1  # Center
                        p_cap.paragraph_format.space_after = Pt(12)
                        r_cap = p_cap.add_run("Hình: Biểu đồ trực quan hóa dữ liệu thống kê")
                        r_cap.font.italic = True
                        r_cap.font.size = Pt(10)
                        r_cap.font.name = "Times New Roman"
                        r_cap.font.color.rgb = RGBColor(100, 116, 139)
                    except Exception as img_err:
                        print(f"[DOCX Renderer] Lỗi khi nhúng ảnh biểu đồ: {img_err}")

        table_rows = []
        in_table = False

    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()

        if line.startswith("|"):
            in_table = True
            table_rows.append(line)
            idx += 1
            continue
        elif in_table:
            flush_table()

        if not line:
            idx += 1
            continue

        # Check for center alignment tags (case-insensitive)
        is_center = False
        line_lower = line.lower()
        if (line_lower.startswith("<center>") and line_lower.endswith("</center>")) or (line_lower.startswith('<p align="center">') and line_lower.endswith('</p>')):
            is_center = True
            line = re.sub(r"^<center>", "", line, flags=re.IGNORECASE)
            line = re.sub(r"</center>$", "", line, flags=re.IGNORECASE)
            line = re.sub(r'^<p align="center">', "", line, flags=re.IGNORECASE)
            line = re.sub(r'</p>$', "", line, flags=re.IGNORECASE)
            line = line.strip()

        if line.startswith("# "):
            heading_text = line[2:].strip()
            p = doc.add_paragraph()
            if is_center:
                p.alignment = 1
            r = p.add_run(heading_text.upper())
            r.font.bold = True
            r.font.size = Pt(13)
            r.font.name = "Times New Roman"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True

        elif line.startswith("## "):
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            if is_center:
                p.alignment = 1
            r = p.add_run(heading_text)
            r.font.bold = True
            r.font.size = Pt(13)
            r.font.name = "Times New Roman"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True

        elif line.startswith("### "):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            if is_center:
                p.alignment = 1
            r = p.add_run(heading_text)
            r.font.bold = True
            r.font.italic = True
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True

        elif re.match(r"^[-*+]\s+", line):
            bullet_text = re.sub(r"^[-*+]\s+", "", line).strip()
            p = doc.add_paragraph(style="List Bullet")
            if is_center:
                p.alignment = 1
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.75)
            add_formatted_text(p, bullet_text)

        elif re.match(r"^\d+\.\s+", line):
            match = re.match(r"^(\d+\.\s+)", line)
            num_prefix = match.group(1) if match else ""
            num_text = re.sub(r"^\d+\.\s+", "", line).strip()
            p = doc.add_paragraph()
            if is_center:
                p.alignment = 1
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.75)
            if num_prefix:
                run_num = p.add_run(num_prefix)
                run_num.font.name = "Times New Roman"
                run_num.font.size = Pt(12)
            add_formatted_text(p, num_text)

        else:
            p = doc.add_paragraph()
            if is_center:
                p.alignment = 1
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, line)

        idx += 1

    if in_table:
        flush_table()

    return doc


def render_markdown_to_html(title: str, content_markdown: str) -> str:
    import re

    lines = content_markdown.split("\n")

    html_body = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        parsed_rows = []
        for r in table_rows:
            cols = [c.strip() for c in r.split("|")]
            if len(cols) >= 2:
                if cols[0] == "":
                    cols = cols[1:]
                if cols and cols[-1] == "":
                    cols = cols[:-1]
                if all(re.match(r"^\s*:-?-*:?\s*$", c) or re.match(r"^\s*-+\s*$", c) for c in cols):
                    continue
                parsed_rows.append(cols)

        if parsed_rows:
            headers = parsed_rows[0] if parsed_rows else []
            html_body.append('<table class="report-table">')
            for r_idx, row_data in enumerate(parsed_rows):
                is_header = r_idx == 0
                html_body.append("<tr>")
                for c_idx, val in enumerate(row_data):
                    if is_header:
                        val_html = parse_inline_markdown(val)
                        html_body.append(f'<th style="text-align: center;">{val_html}</th>')
                    else:
                        header_name = headers[c_idx] if c_idx < len(headers) else ""
                        align = detect_cell_alignment(header_name, val)
                        if "chênh lệch" in header_name.lower() or "(δ)" in header_name.lower() or "(delta)" in header_name.lower():
                            val = sanitize_delta_value(val)
                        val_html = parse_inline_markdown(val)
                        html_body.append(f'<td style="text-align: {align.value};">{val_html}</td>')
                html_body.append("</tr>")
            html_body.append("</table>")

            # Tự động sinh và nhúng biểu đồ trực quan dưới bảng (nếu bảng có số liệu)
            if len(parsed_rows) >= 2:
                data_rows = parsed_rows[1:]
                chart_res = generate_chart_for_table(headers, data_rows, report_title=title)
                if chart_res:
                    _, data_uri = chart_res
                    html_body.append(
                        f'<div style="text-align: center; margin: 20px 0;">'
                        f'<img src="{data_uri}" style="max-width: 100%; height: auto; border-radius: 8px; border: 1px solid #E2E8F0; box-shadow: 0 2px 6px rgba(0,0,0,0.06);" />'
                        f'<p style="font-size: 10pt; color: #64748B; font-style: italic; margin-top: 6px; text-align: center;">Hình: Biểu đồ trực quan hóa dữ liệu thống kê</p>'
                        f'</div>'
                    )

        table_rows = []
        in_table = False

    def parse_inline_markdown(text):
        text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
        text = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", text)
        text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
        return text

    in_list = False
    list_type = None

    def flush_list():
        nonlocal in_list, list_type
        if in_list:
            html_body.append(f"</{list_type}>")
            in_list = False
            list_type = None

    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()

        if line.startswith("|"):
            flush_list()
            in_table = True
            table_rows.append(line)
            idx += 1
            continue
        elif in_table:
            flush_table()

        if not line:
            flush_list()
            idx += 1
            continue

        # Check for center alignment tags (case-insensitive)
        is_center = False
        line_lower = line.lower()
        if (line_lower.startswith("<center>") and line_lower.endswith("</center>")) or (line_lower.startswith('<p align="center">') and line_lower.endswith('</p>')):
            is_center = True
            line = re.sub(r"^<center>", "", line, flags=re.IGNORECASE)
            line = re.sub(r"</center>$", "", line, flags=re.IGNORECASE)
            line = re.sub(r'^<p align="center">', "", line, flags=re.IGNORECASE)
            line = re.sub(r'</p>$', "", line, flags=re.IGNORECASE)
            line = line.strip()

        tag_html = ""

        if line.startswith("# "):
            flush_list()
            tag_html = f"<h2>{parse_inline_markdown(line[2:].strip())}</h2>"
        elif line.startswith("## "):
            flush_list()
            tag_html = f"<h2>{parse_inline_markdown(line[3:].strip())}</h2>"
        elif line.startswith("### "):
            flush_list()
            tag_html = f"<h3>{parse_inline_markdown(line[4:].strip())}</h3>"
        elif re.match(r"^[-*+]\s+", line):
            bullet_text = re.sub(r"^[-*+]\s+", "", line).strip()
            if not in_list or list_type != "ul":
                flush_list()
                html_body.append("<ul>")
                in_list = True
                list_type = "ul"
            tag_html = f"<li>{parse_inline_markdown(bullet_text)}</li>"
        elif re.match(r"^\d+\.\s+", line):
            num_text = re.sub(r"^\d+\.\s+", "", line).strip()
            if not in_list or list_type != "ol":
                flush_list()
                html_body.append("<ol>")
                in_list = True
                list_type = "ol"
            tag_html = f"<li>{parse_inline_markdown(num_text)}</li>"
        else:
            flush_list()
            tag_html = f"<p>{parse_inline_markdown(line)}</p>"

        if is_center and tag_html:
            if tag_html.startswith("<p>"):
                tag_html = tag_html.replace("<p>", '<p style="text-align: center;">', 1)
            elif tag_html.startswith("<h2>"):
                tag_html = tag_html.replace("<h2>", '<h2 style="text-align: center; text-transform: uppercase;">', 1)
            elif tag_html.startswith("<h3>"):
                tag_html = tag_html.replace("<h3>", '<h3 style="text-align: center;">', 1)
            else:
                tag_html = f'<div style="text-align: center;">{tag_html}</div>'

        if tag_html:
            html_body.append(tag_html)

        idx += 1

    flush_table()
    flush_list()

    body_content = "\n".join(html_body)

    html_template = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    body {{
        font-family: 'Times New Roman', Times, serif;
        font-size: 13pt;
        line-height: 1.3;
        color: #000000;
        margin: 0;
        padding: 2cm 2cm 2cm 3cm;
        background-color: #ffffff;
    }}
    .document-title {{
        text-align: center;
        font-size: 14pt;
        font-weight: bold;
        text-transform: uppercase;
        margin-bottom: 24px;
    }}
    p {{
        margin-top: 0;
        margin-bottom: 12px;
        text-align: justify;
    }}
    h2 {{
        font-size: 13pt;
        font-weight: bold;
        margin-top: 18px;
        margin-bottom: 8px;
        text-transform: uppercase;
    }}
    h3 {{
        font-size: 12pt;
        font-weight: bold;
        font-style: italic;
        margin-top: 12px;
        margin-bottom: 6px;
    }}
    .report-table {{
        width: 100%;
        border-collapse: collapse;
        margin-top: 12px;
        margin-bottom: 18px;
    }}
    .report-table th, .report-table td {{
        border: 1px solid #000000;
        padding: 6px 8px;
        font-size: 11pt;
    }}
    .report-table th {{
        background-color: #f2f2f2;
        font-weight: bold;
        text-align: center;
    }}
    .report-table td {{
        text-align: left;
    }}
    ul, ol {{
        margin-top: 0;
        margin-bottom: 12px;
        padding-left: 20px;
    }}
    li {{
        margin-bottom: 4px;
    }}
    code {{
        font-family: Consolas, Monaco, monospace;
        font-size: 10pt;
        background-color: #f4f4f4;
        padding: 2px 4px;
        border-radius: 3px;
    }}
</style>
</head>
<body>
    {body_content}
</body>
</html>
"""
    return html_template


@tool
async def generate_custom_report_docx(title: str, content_markdown: str) -> str:
    """Tạo tệp báo cáo tự do (.docx và .html) từ nội dung Markdown được định nghĩa bởi Agent và trả về liên kết tải xuống/xem trước.

    Args:
        title: Tiêu đề báo cáo (ví dụ: 'Báo cáo Học tập bổ sung lớp 10A1').
        content_markdown: Nội dung báo cáo định dạng Markdown (hỗ trợ các tiêu đề, danh sách, bảng biểu).
    """
    file_uuid = str(uuid.uuid4())[:8]
    filename_docx = f"bao_cao_tu_do_{file_uuid}.docx"
    filename_html = f"bao_cao_tu_do_{file_uuid}.html"
    temp_dir = "temp"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    file_path_docx = os.path.join(temp_dir, filename_docx)
    file_path_html = os.path.join(temp_dir, filename_html)

    try:
        # 1. Generate DOCX
        doc = render_markdown_to_docx(title, content_markdown)
        doc.save(file_path_docx)

        # 2. Generate HTML
        html_content = render_markdown_to_html(title, content_markdown)
        with open(file_path_html, "w", encoding="utf-8") as f:
            f.write(html_content)

        # 3. Generate PDF via Gotenberg to prevent download errors
        filename_pdf = f"bao_cao_tu_do_{file_uuid}.pdf"
        file_path_pdf = os.path.join(temp_dir, filename_pdf)

        gotenberg_url = "https://c2-app-051-gotenberg.up.railway.app/forms/libreoffice/convert"
        try:
            import requests
            with open(file_path_docx, "rb") as f_docx:
                files = {
                    "files": (
                        "report.docx",
                        f_docx.read(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                }
            resp = requests.post(gotenberg_url, files=files, timeout=60)
            if resp.status_code == 200:
                with open(file_path_pdf, "wb") as f_pdf:
                    f_pdf.write(resp.content)
            else:
                print(f"Gotenberg convert failed: {resp.text}")
        except Exception as e:
            print(f"Failed to generate PDF for custom report: {str(e)}")

        from src.config import get_settings

        settings = get_settings()
        download_url = f"{settings.backend_url.rstrip('/')}/api/v1/reports/download/{filename_docx}"
        preview_url = f"{settings.backend_url.rstrip('/')}/api/v1/reports/download/{filename_html}"

        return f"Tệp báo cáo tự do đã được tạo thành công!\n\n👉 [Xem Bản Xem Trước Báo Cáo]({preview_url}) | [Tải Báo Cáo Word (.docx)]({download_url})"
    except Exception as e:
        return f"Lỗi trong quá trình tạo tệp báo cáo tự do: {str(e)}"
