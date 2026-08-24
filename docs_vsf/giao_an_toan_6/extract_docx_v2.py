# -*- coding: utf-8 -*-
"""
Enhanced docx extraction: captures math formulas, tables, textboxes, and images.
"""
import sys, re, json
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8")

_HERE = Path(__file__).resolve().parent
DOCX_FOLDER = _HERE
OUTPUT = _HERE / "docx_extracted_v2.json"

import docx
from lxml import etree

# Namespaces for XML parsing
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'v': 'urn:schemas-microsoft-com:vml',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
}


# ============================================================================
# MATH FORMULA EXTRACTOR (OMML → LaTeX)
# ============================================================================

def _omml_to_latex(elem):
    """Convert OMML (m:oMath) to text representation — captures all math content."""
    if elem is None:
        return ''
    
    # Get local tag name (without namespace)
    tag = etree.QName(elem).localname if elem.tag else ''
    m_ns = NSMAP['m']
    w_ns = NSMAP['w']
    
    # TEXT: m:t or w:t
    if tag == 't':
        return (elem.text or '')
    
    # RUN: m:r — contains m:t
    if tag == 'r':
        t = elem.find(f'{{{m_ns}}}t')
        if t is not None and t.text:
            return t.text
        t = elem.find(f'{{{w_ns}}}t')
        if t is not None and t.text:
            return t.text
        return ''
    
    # FRACTION: m:f → m:num / m:den
    if tag == 'f':
        num = elem.find(f'{{{m_ns}}}num')
        den = elem.find(f'{{{m_ns}}}den')
        # Try direct text extraction as fallback
        all_texts = [c.text for c in elem.iter() if c.text and len(c.text.strip()) > 0]
        num_t = den_t = ''
        if num is not None:
            num_t = _omml_to_latex(num)
            if not num_t.strip():
                # Fallback: extract all text from num
                num_texts = [c.text for c in num.iter() if c.text and len(c.text.strip()) > 0]
                num_t = ''.join(num_texts) if num_texts else ''
        if den is not None:
            den_t = _omml_to_latex(den)
            if not den_t.strip():
                den_texts = [c.text for c in den.iter() if c.text and len(c.text.strip()) > 0]
                den_t = ''.join(den_texts) if den_texts else ''
        if num_t or den_t:
            return f'({num_t})/({den_t})'
        # Last resort: all texts
        if all_texts:
            return f'({" ".join(all_texts)})'
        return ''
    
    # SUBSCRIPT: m:sub
    if tag == 'sub':
        e = elem.find(f'{{{m_ns}}}e')
        # subArg
        sub_arg = elem.find(f'{{{m_ns}}}subArg')
        base = _omml_to_latex(e) if e is not None else ''
        sub = _omml_to_latex(sub_arg) if sub_arg is not None else ''
        return f'{base}_{{{sub}}}' if base or sub else ''
    
    # SUPERSCRIPT: m:sup
    if tag == 'sup':
        e = elem.find(f'{{{m_ns}}}e')
        sup_arg = elem.find(f'{{{m_ns}}}supArg')
        base = _omml_to_latex(e) if e is not None else ''
        sup = _omml_to_latex(sup_arg) if sup_arg is not None else ''
        return f'{base}^{{{sup}}}' if base or sup else ''
    
    # SUBSCRIPT+SUPERSCRIPT: m:sSubSup
    if tag == 'sSubSup':
        e = elem.find(f'{{{m_ns}}}e')
        sub = elem.find(f'{{{m_ns}}}sub')
        sup = elem.find(f'{{{m_ns}}}sup')
        parts = []
        if e is not None:
            parts.append(_omml_to_latex(e))
        if sub is not None:
            parts.append('_' + '{' + _omml_to_latex(sub) + '}')
        if sup is not None:
            parts.append('^' + '{' + _omml_to_latex(sup) + '}')
        return ''.join(parts)
    
    # RADICAL: m:rad (sqrt)
    if tag == 'rad':
        rad_e = elem.find(f'{{{m_ns}}}e')
        deg = elem.find(f'{{{m_ns}}}deg')
        content = _omml_to_latex(rad_e) if rad_e is not None else ''
        if deg is not None:
            d = _omml_to_latex(deg)
            return f'\\sqrt[{d}]{{{content}}}' if d else f'\\sqrt{{{content}}}'
        return f'\\sqrt{{{content}}}' if content else ''
    
    # N-ARY (sum, integral): m:nary
    if tag == 'nary':
        e = elem.find(f'{{{m_ns}}}e')
        return _omml_to_latex(e) if e is not None else ''
    
    # Skip metadata/property elements at all levels
    skip_tags = {'rPr', 'spPr', 'def', 'argPr', 'sty', 'ctrlPr',
                 'accPr', 'radPr', 'sPrePr', 'sSubSupPr', 'subPr', 'supPr', 'naryPr',
                 'eqArrPr', 'funcPr', 'dPr', 'limLowPr', 'limUppPr', 'boxPr', 'barPr',
                 'groupChrPr', 'borderBoxPr', 'mPr', 'm:rPr', 'm:spPr', 'm:def', 'm:argPr',
                 'm:sty', 'm:ctrlPr', 'm:accPr', 'm:radPr', 'm:sPrePr', 'm:sSubSupPr',
                 'm:subPr', 'm:supPr', 'm:naryPr', 'm:eqArrPr', 'm:funcPr', 'm:dPr',
                 'm:limLowPr', 'm:limUppPr', 'm:boxPr', 'm:barPr', 'm:groupChrPr',
                 'm:borderBoxPr', 'm:begChr', 'm:sepChr', 'm:endChr', 'm:grow', 'm:subHide',
                 'm:supHide', 'm:brk', 'm:aln', 'm:type', 'm:val', 'm:maxDist', 'm:objDist',
                 'm:opEmu', 'm:shp', 'm:scr', 'm:sty', 'm:char', 'm:limLoc', 'm:transp',
                 'm:interSp', 'm:intraSp', 'm:plcHide', 'rPr', 'w:rPr', 'w:rPr',  
                 'm:normWs', 'm:manualBreak', 'm:lit', 'm:litTag', 'm:args', 'm:baseJc',
                 'm:defJc', 'm:dispDef', 'm:jc', 'm:smallFrac', 'm:wrapIndent', 'm:rSpRule',
                 'm:rSp', 'm:dist', 'm:lMargin', 'm:rMargin', 'm:defFont', 'm:defFontSize',
                 'm:noBreak', 'm:postSp', 'm:preSp', 'm:script', 'm:show', 'm:showDeg',
                 'm:zeroWid', 'm:pos', 'm:groupChr', 'm:borderBox', 'm:limLow', 'm:limUpp',
                 'tbl', 'tr', 'tc', 'tcPr', 'tblPr', 'trPr', 'gridCol', 'tblGrid', 'tblW',
                 'trHeight', 'vAlign', 'shd', 'tcMar', 'tcW'}
    if tag in skip_tags:
        return ''
    
    # DEFAULT: recurse children
    parts = []
    for child in elem:
        if child.tag:
            child_tag = etree.QName(child).localname
            if child_tag in skip_tags:
                continue
            parts.append(_omml_to_latex(child))
        elif child.text:
            parts.append(child.text)
    return ''.join(parts)


def extract_paragraph_text_with_math(para, xml_root):
    """Extract text from a paragraph including math formulas."""
    result_parts = []
    
    # Get all run text normally
    for run in para.runs:
        result_parts.append(run.text)
    
    # Find oMath and oMathPara elements, convert to LaTeX
    math_spans = []
    for math_elem in xml_root.findall('.//m:oMath', NSMAP):
        latex = _omml_to_latex(math_elem)
        if latex.strip():
            math_spans.append(f'${latex}$')
        else:
            # Fallback: mark as formula
            math_spans.append('$[MATH]$')
    
    # Replace sequences of spaces where math was (math elements don't produce runs)
    # Actually math elements are inline — they should be inserted where they appear
    # But we don't have position info easily, so append at end
    # Better approach: extract via XML walk
    
    # Re-parse via XML for precise ordering: interleave text and math
    full_text = []
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    for child in xml_root:
        local = etree.QName(child).localname if child.tag else ''
        if local == 'r':  # run
            t = child.find(f'{{{ns}}}t')
            if t is not None and t.text:
                full_text.append(t.text)
        elif local.endswith('oMath') or local.endswith('oMathPara'):  # math
            latex = _omml_to_latex(child)
            if latex.strip():
                full_text.append(f' ${latex}$ ')
            else:
                full_text.append(' [MATH] ')
    
    text = ''.join(full_text)
    if text.strip():
        return text
    
    # Fallback
    return ' '.join(result_parts)


def extract_math_text_hybrid(para, para_index=0):
    """Extract text with math formulas converted to LaTeX."""
    normal_text = para.text or ''
    xml_root = para._element
    
    # Find all math elements
    math_elems = xml_root.findall('.//m:oMath', NSMAP)
    if not math_elems:
        return normal_text
    
    # Convert each math element to LaTeX
    math_latex = []
    for me in math_elems:
        latex = _omml_to_latex(me)
        if latex.strip():
            math_latex.append(f'${latex}$')
        else:
            # Fallback: extract any text
            texts = [c.text for c in me.iter() if c.text and len(c.text.strip()) > 0]
            if texts:
                math_latex.append(f'$[{"".join(texts)}]$')
            else:
                # Absolute fallback: convert raw XML
                raw = etree.tostring(me, pretty_print=True).decode()
                # Try to find digits
                import re
                digits = re.findall(r'\d+', raw)
                if digits:
                    math_latex.append(f'$[{"".join(digits)}]$')
                else:
                    math_latex.append('$[MATH]$')
    
    # Build result: interleave text and math
    # Get all text runs (w:t) in order
    parts = []
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    math_idx = 0
    
    # Walk paragraphs children in order
    for child in xml_root:
        local = etree.QName(child).localname if child.tag else ''
        if local == 'r':
            # Regular run — extract text
            t = child.find(f'{{{w_ns}}}t')
            if t is not None and t.text:
                parts.append(t.text)
        elif local.endswith('oMath') or local.endswith('oMathPara'):
            # Math element
            if math_idx < len(math_latex):
                parts.append(math_latex[math_idx])
                math_idx += 1
    
    result = ''.join(parts)
    if result.strip():
        return result
    
    # Debug: no result despite having math
    if para_index < 5:
        debug_inner = etree.tostring(xml_root, pretty_print=True).decode()
        import re
        # Check if any math was found
        omath_count = len(re.findall(r'm:oMath', debug_inner))
        print(f"  [DEBUG] para {para_index}: {omath_count} oMath found, math_latex={math_latex}", file=sys.stderr)
    
    return normal_text


def format_table(table, table_index):
    """Render a table as markdown."""
    rows_list = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' // ') for cell in row.cells]
        rows_list.append(' | '.join(cells))
    
    if not rows_list:
        return f'(table {table_index}: empty)'
    
    header = rows_list[0]
    col_count = len(table.rows[0].cells) if table.rows else 0
    separator = ' | '.join(['---'] * col_count)
    
    body = '\n'.join(rows_list[1:]) if len(rows_list) > 1 else ''
    
    result = f'[TABLE {table_index}]\n{header}\n{separator}\n{body}\n[/TABLE]'
    return result


def extract_textboxes_from_xml(xml_str: str) -> str:
    """Find w:txbxContent elements in paragraph XML."""
    if 'txbxContent' not in xml_str:
        return ''
    
    root = etree.fromstring(xml_str.encode())
    texts = []
    for tbx in root.iter(f'{{{NSMAP["w"]}}}txbxContent'):
        for t in tbx.iter(f'{{{NSMAP["w"]}}}t'):
            if t.text:
                texts.append(t.text)
    
    if texts:
        return '\n'.join(texts)
    return ''


def extract_images_from_xml(xml_str: str) -> list[str]:
    """Find image references in paragraph XML."""
    if 'wp:inline' not in xml_str and 'wp:anchor' not in xml_str and 'v:shape' not in xml_str:
        return []
    
    images = []
    # Just count and mark them
    inline_count = xml_str.count('wp:inline')
    anchor_count = xml_str.count('wp:anchor')
    total = inline_count + anchor_count
    
    if total > 0:
        # Try to find alt text or description
        root = etree.fromstring(xml_str.encode())
        for desc in root.iter(f'{{{NSMAP["wp"]}}}docPr'):
            name = desc.get('name', '') if desc is not None else ''
            descr = desc.get('descr', '') if desc is not None else ''
            if name or descr:
                images.append(f'[HÌNH: {name} - {descr}]')
        
        if not images:
            images.append(f'[HÌNH: {total} hình]')
    
    return images


# ============================================================================
# MAIN EXTRACTION
# ============================================================================

def extract_lessons_enhanced(path):
    """Extract lessons with math, tables, textboxes preserved."""
    doc = docx.Document(path)
    
    # Build a lookup: paragraph index → table
    # In docx, tables are stored between paragraphs in the body
    paragraph_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    table_para_before: list[int] = []
    table_para_after: list[int] = []
    
    # Find which paragraph each table sits between
    # We track this by finding where tables appear in the XML body
    body = doc.element.body
    current_para_idx = -1
    table_idx = 0
    tables_between_paras: dict[int, int] = {}  # after paragraph idx → table index
    
    for child in body:
        local = etree.QName(child).localname if child.tag else ''
        if local == 'p':
            current_para_idx += 1
        elif local == 'tbl':
            if table_idx < table_count:
                tables_between_paras[current_para_idx] = table_idx
                table_idx += 1
    
    # Extract lessons
    lessons = []
    current_lesson = None
    current_content_parts = []
    
    def _flush_lesson():
        nonlocal current_lesson, current_content_parts
        if current_lesson:
            current_lesson['content'] = '\n'.join(current_content_parts)
            # DEBUG: check if content has math
            content_str = current_lesson['content']
            if '$' in content_str or '(4)/(8)' in content_str:
                import sys as _sys
                _sys.stderr.write('[DEBUG] FLUSH content has math: ' + content_str[:100] + '\n')
            lessons.append(current_lesson)
            current_content_parts = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ''
        xml_str = para._element.xml if hasattr(para, '_element') else ''
        
        # Check if a table sits AFTER this paragraph (before next paragraph)
        if i in tables_between_paras:
            tbl_idx = tables_between_paras[i]
            if tbl_idx < table_count:
                table_md = format_table(doc.tables[tbl_idx], tbl_idx + 1)
                current_content_parts.append(f'[Normal] {table_md}')
        
        # Extract math formula from XML
        math_content = extract_math_text_hybrid(para)
        # DEBUG: check para 88 (HK2 fraction)
        if i == 88 and 'phân số' in str(para.text).lower():
            import sys as _sys
            mc = math_content or ''
            _sys.stderr.write('[DEBUG] Para 88 math_content: ' + mc[:200] + '\n')
            _sys.stderr.write('[DEBUG] Para 88 has $: ' + str('$' in mc) + '\n')
        
        # Extract textboxes
        textbox_content = extract_textboxes_from_xml(xml_str)
        
        # Extract images
        images = extract_images_from_xml(xml_str)
        
        if 'heading 1' == style.lower() and text:
            # MERGE CONSECUTIVE HEADING 1s: nếu giữa 2 heading không có content → gộp lại
            if current_lesson and not current_content_parts:
                current_lesson['title'] = current_lesson['title'] + ' | ' + text
            else:
                _flush_lesson()
                current_lesson = {
                    'title': text,
                    'index': i,
                }
                current_content_parts = []
        elif current_lesson and text:
            # Use enhanced math-extracted text
            display_text = math_content or text
            # DEBUG: check display_text for para 88
            if i == 88:
                import sys as _sys
                _sys.stderr.write('[DEBUG] display_text: ' + (display_text or '')[:200] + '\n')
            current_content_parts.append(f'[{style}] {display_text}')
        
        # Append textbox content if found
        if textbox_content:
            current_content_parts.append(f'[Textbox] {textbox_content}')
        
        # Append image markers
        for img in images:
            current_content_parts.append(f'[Normal] {img}')
    
    _flush_lesson()
    
    # Collect remaining tables not assigned to any paragraph
    assigned_tables = set(tables_between_paras.values())
    for ti in range(table_count):
        if ti not in assigned_tables:
            # These tables might be at the end of document or unassigned
            pass  # skip — they're probably empty or metadata
    
    return {
        'file': str(path.name),
        'total_paragraphs': len(doc.paragraphs),
        'lessons': lessons,
        'stats': {
            'tables': table_count,
            'math_paras': len([p for p in doc.paragraphs if 'm:oMath' in (p._element.xml if hasattr(p, '_element') else '')]),
        }
    }


def main():
    files = [
        DOCX_FOLDER / "KHBD Toan 6 (CTST)-Ki.docx",
        DOCX_FOLDER / "KHBD Toan 6 (CTST)-KII.docx",
    ]
    
    result = {}
    for path in files:
        if not path.exists():
            print(f"[ERROR] Not found: {path}")
            continue
        name = "HK1" if "Ki.docx" in path.name else "HK2"
        print(f"[INFO] Processing {name}: {path.name}...")
        result[name] = extract_lessons_enhanced(path)
        print(f"  -> {len(result[name]['lessons'])} lessons / {result[name]['stats']['tables']} tables / {result[name]['stats']['math_paras']} math paras")
    
    OUTPUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n[SUCCESS] Saved to {OUTPUT}")


if __name__ == "__main__":
    main()